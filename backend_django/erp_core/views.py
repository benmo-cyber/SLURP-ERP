from rest_framework import viewsets
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework import status
from django.utils import timezone
from django.db import models
from datetime import datetime
from .models import (
    Item, Lot, ProductionBatch, ProductionBatchInput, ProductionBatchOutput,
    CriticalControlPoint, Formula, FormulaItem,
    PurchaseOrder, PurchaseOrderItem, SalesOrder, SalesOrderItem,
    InventoryTransaction, LotNumberSequence, Vendor, VendorHistory,
    SupplierSurvey, SupplierDocument, TemporaryException, CostMaster, CostMasterHistory, Account,
    FinishedProductSpecification, Customer, CustomerPricing, VendorPricing, SalesOrderLot, Invoice, InvoiceItem,
    ShipToLocation, CustomerContact, SalesCall, CustomerForecast, BatchNumberSequence, LotDepletionLog,
    PurchaseOrderLog, ProductionLog, Shipment, ShipmentItem, LotTransactionLog, ItemPackSize,
    FiscalPeriod, JournalEntry, JournalEntryLine, GeneralLedgerEntry, AccountBalance,
    AccountsPayable, AccountsReceivable, Payment, BankReconciliation
)
from .serializers import (
    ItemSerializer, LotSerializer, ProductionBatchSerializer,
    AccountsPayableSerializer, AccountsReceivableSerializer, PaymentSerializer, BankReconciliationSerializer,
    CriticalControlPointSerializer, FormulaSerializer, FormulaItemSerializer,
    PurchaseOrderSerializer, PurchaseOrderItemSerializer,
    SalesOrderSerializer, SalesOrderItemSerializer,
    InventoryTransactionSerializer, VendorSerializer, VendorHistorySerializer,
    SupplierSurveySerializer, SupplierDocumentSerializer, TemporaryExceptionSerializer,
    CostMasterSerializer, CostMasterHistorySerializer, AccountSerializer,
    FinishedProductSpecificationSerializer, CustomerSerializer, CustomerPricingSerializer, VendorPricingSerializer,
    InvoiceSerializer, InvoiceItemSerializer, ShipmentSerializer, ShipToLocationSerializer,
    CustomerContactSerializer, SalesCallSerializer, CustomerForecastSerializer, LotDepletionLogSerializer,
    LotTransactionLogSerializer, PurchaseOrderLogSerializer, ProductionLogSerializer, ItemPackSizeSerializer,
    FiscalPeriodSerializer, JournalEntrySerializer, JournalEntryLineSerializer, GeneralLedgerEntrySerializer, AccountBalanceSerializer
)
from .email_service import send_invoice_email, send_purchase_order_email, send_sales_order_confirmation_email
from .pdf_generator import generate_invoice_pdf, generate_purchase_order_pdf, generate_sales_order_pdf


def log_lot_transaction(lot, quantity_before, quantity_change, transaction_type, reference_number=None, 
                       reference_type=None, transaction_id=None, batch_id=None, sales_order_id=None, 
                       purchase_order_id=None, notes=None):
    """
    Log ALL lot quantity transactions, not just depletions.
    
    Args:
        lot: The Lot instance
        quantity_before: Quantity remaining before this transaction
        quantity_change: Quantity change (positive for additions, negative for reductions)
        transaction_type: One of 'receipt', 'production_input', 'production_output', 'repack_input', 'repack_output', 'sale', 'adjustment', 'allocation', 'deallocation', 'manual', 'reversal'
        reference_number: Batch number, SO number, PO number, etc.
        reference_type: Type of reference ('batch_number', 'so_number', 'po_number', etc.)
        transaction_id: Related InventoryTransaction ID if applicable
        batch_id: Related ProductionBatch ID if applicable
        sales_order_id: Related SalesOrder ID if applicable
        purchase_order_id: Related PurchaseOrder ID if applicable
        notes: Additional context
    """
    quantity_after = quantity_before + quantity_change
    
    try:
        # Check if the table exists before trying to create a log entry
        from django.db import connection
        table_exists = False
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_lottransactionlog'")
                table_exists = cursor.fetchone() is not None
        except Exception:
            # If we can't check, assume table doesn't exist
            table_exists = False
        
        if not table_exists:
            # Table doesn't exist, skip logging silently
            return
        
        LotTransactionLog.objects.create(
            lot=lot,
            lot_number=lot.lot_number,
            item_sku=lot.item.sku,
            item_name=lot.item.name,
            vendor=lot.item.vendor or '',
            transaction_type=transaction_type,
            quantity_before=quantity_before,
            quantity_change=quantity_change,
            quantity_after=quantity_after,
            unit_of_measure=lot.item.unit_of_measure,
            reference_number=reference_number,
            reference_type=reference_type,
            transaction_id=transaction_id,
            batch_id=batch_id,
            sales_order_id=sales_order_id,
            purchase_order_id=purchase_order_id,
            notes=notes or f'Lot {lot.lot_number} transaction: {transaction_type}'
        )
    except Exception as e:
        # Log error but don't fail the transaction
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f'Failed to log lot transaction for lot {lot.lot_number}: {str(e)}')


def log_lot_depletion(lot, quantity_before, quantity_used, depletion_method, reference_number=None, 
                      reference_type=None, transaction_id=None, batch_id=None, sales_order_id=None, notes=None):
    """
    Log when a lot is depleted to zero or below.
    
    Args:
        lot: The Lot instance
        quantity_before: Quantity remaining before this transaction
        quantity_used: Quantity used in this transaction (positive number)
        depletion_method: One of 'production', 'sales', 'adjustment', 'manual', 'reversal'
        reference_number: Batch number, SO number, etc.
        reference_type: Type of reference ('batch_number', 'so_number', etc.)
        transaction_id: Related InventoryTransaction ID if applicable
        batch_id: Related ProductionBatch ID if applicable
        sales_order_id: Related SalesOrder ID if applicable
        notes: Additional context
    """
    final_quantity = quantity_before - quantity_used
    
    # Also log as a regular transaction
    log_lot_transaction(
        lot=lot,
        quantity_before=quantity_before,
        quantity_change=-quantity_used,
        transaction_type=depletion_method if depletion_method in ['production', 'sales', 'adjustment', 'manual', 'reversal'] else 'adjustment',
        reference_number=reference_number,
        reference_type=reference_type,
        transaction_id=transaction_id,
        batch_id=batch_id,
        sales_order_id=sales_order_id,
        notes=notes
    )
    
    # Only log depletion if the lot is depleted (reaches zero or goes negative)
    if final_quantity <= 0:
        try:
            # Check if the table exists before trying to create a log entry
            from django.db import connection
            table_exists = False
            try:
                with connection.cursor() as cursor:
                    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_lotdepletionlog'")
                    table_exists = cursor.fetchone() is not None
            except Exception:
                # If we can't check, assume table doesn't exist
                table_exists = False
            
            if not table_exists:
                # Table doesn't exist, skip logging silently
                return
            
            LotDepletionLog.objects.create(
                lot=lot,
                lot_number=lot.lot_number,
                item_sku=lot.item.sku,
                item_name=lot.item.name,
                vendor=lot.item.vendor or '',
                initial_quantity=lot.quantity,
                quantity_before=quantity_before,
                quantity_used=quantity_used,
                final_quantity=final_quantity,
                depletion_method=depletion_method,
                reference_number=reference_number,
                reference_type=reference_type,
                transaction_id=transaction_id,
                batch_id=batch_id,
                sales_order_id=sales_order_id,
                notes=notes or f'Lot depleted from {quantity_before} to {final_quantity} via {depletion_method}'
            )
        except Exception as e:
            # Log error but don't fail the transaction
            # Silently ignore if table doesn't exist
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f'Failed to log lot depletion for lot {lot.lot_number}: {str(e)}')


def log_purchase_order_action(po, action, lot=None, notes=None):
    """
    Log purchase order actions (created, updated, check-in, etc.)
    
    Args:
        po: The PurchaseOrder instance
        action: One of 'created', 'updated', 'check_in', 'partial_check_in', 'cancelled', 'completed'
        lot: The Lot instance if this is a check-in
        notes: Additional context
    """
    try:
        # Calculate totals
        total_items = po.items.count()
        total_quantity_ordered = sum(item.quantity_ordered for item in po.items.all())
        total_quantity_received = sum(item.quantity_received for item in po.items.all())
        
        log_data = {
            'purchase_order': po,
            'po_number': po.po_number,
            'action': action,
            'vendor_name': po.vendor_customer_name,
            'vendor_customer_name': po.vendor_customer_name,
            'po_date': po.order_date,
            'required_date': po.required_date,
            'status': po.status,
            'carrier': po.carrier or '',
            'po_received_date': po.received_date,  # PO's received_date
            'total_items': total_items,
            'total_quantity_ordered': total_quantity_ordered,
            'total_quantity_received': total_quantity_received,
            'notes': notes
        }
        
        # If this is a check-in, add lot information
        if lot and action in ['check_in', 'partial_check_in']:
            log_data.update({
                'lot_number': lot.lot_number,
                'item_sku': lot.item.sku,
                'item_name': lot.item.name,
                'quantity_received': lot.quantity,
                'received_date': lot.received_date
            })
        
        PurchaseOrderLog.objects.create(**log_data)
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f'Failed to log purchase order action for PO {po.po_number}: {str(e)}')


def create_ap_entry_from_po(purchase_order, invoice_number=None, invoice_date=None, due_date=None):
    """
    Create an Accounts Payable entry when a purchase order is received.
    
    Args:
        purchase_order: PurchaseOrder instance
        invoice_number: Vendor invoice number (optional)
        invoice_date: Date of vendor invoice (defaults to today)
        due_date: Payment due date (defaults to invoice_date + vendor payment_terms, or 30 days)
    """
    import re
    from django.utils import timezone
    from datetime import timedelta
    
    try:
        # Check if AP entry already exists for this PO
        if AccountsPayable.objects.filter(purchase_order=purchase_order).exists():
            return AccountsPayable.objects.get(purchase_order=purchase_order)
        
        # Get vendor information
        vendor_name = purchase_order.vendor_customer_name or 'Unknown Vendor'
        vendor_id = purchase_order.vendor_customer_id
        
        # Calculate totals
        total_amount = purchase_order.total or 0.0
        if total_amount <= 0:
            # Calculate from items if total is not set
            total_amount = sum(
                item.quantity_ordered * (item.unit_price or 0) 
                for item in purchase_order.items.all()
            )
        
        # Set dates
        if not invoice_date:
            invoice_date = timezone.now().date()
        if not due_date:
            # Use vendor payment terms from Quality > Vendor approval if available
            days = 30
            try:
                vendor = Vendor.objects.filter(name=vendor_name).first()
                if vendor and getattr(vendor, 'payment_terms', None):
                    match = re.search(r'(\d+)', vendor.payment_terms)
                    if match:
                        days = int(match.group(1))
            except Exception:
                pass
            due_date = invoice_date + timedelta(days=days)
        
        # Get or create AP account (typically account number 2000)
        ap_account = None
        try:
            ap_account = Account.objects.filter(account_type='liability', account_number__startswith='2000').first()
            if not ap_account:
                # Create a default AP account if none exists
                ap_account = Account.objects.create(
                    account_number='2000',
                    name='Accounts Payable',
                    account_type='liability',
                    description='Accounts Payable'
                )
        except Exception:
            pass
        
        # Create AP entry
        ap_entry = AccountsPayable.objects.create(
            vendor_name=vendor_name,
            vendor_id=vendor_id,
            purchase_order=purchase_order,
            invoice_number=invoice_number or purchase_order.po_number,
            invoice_date=invoice_date,
            due_date=due_date,
            original_amount=total_amount,
            amount_paid=0.0,
            balance=total_amount,
            status='open',
            account=ap_account,
            notes=f'Created from PO {purchase_order.po_number}'
        )
        
        # Auto-create journal entry for AP
        try:
            journal_entry = create_ap_journal_entry(ap_entry)
            if journal_entry:
                ap_entry.journal_entry = journal_entry
                ap_entry.save()
        except Exception as je_error:
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f'Failed to create journal entry for AP entry {ap_entry.id}: {str(je_error)}')
        
        return ap_entry
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f'Failed to create AP entry for PO {purchase_order.po_number}: {str(e)}')
        return None


def create_ap_journal_entry(ap_entry):
    """
    Create a journal entry when an AP entry is created.
    Debits: Expense/Asset accounts (based on PO items)
    Credits: Accounts Payable
    """
    from django.utils import timezone
    
    try:
        # Get or create fiscal period
        today = timezone.now().date()
        fiscal_period = FiscalPeriod.objects.filter(
            start_date__lte=today,
            end_date__gte=today
        ).first()
        
        if not fiscal_period:
            # Create a default fiscal period if none exists
            from datetime import date
            year = today.year
            fiscal_period, _ = FiscalPeriod.objects.get_or_create(
                period_name=f"{year}-01",
                defaults={
                    'start_date': date(year, 1, 1),
                    'end_date': date(year, 12, 31),
                }
            )
        if fiscal_period.is_closed:
            return None
        # Generate journal entry number
        last_entry = JournalEntry.objects.filter(entry_date=today).order_by('-entry_number').first()
        if last_entry:
            try:
                last_number = int(last_entry.entry_number.split('-')[-1])
                new_number = last_number + 1
            except (ValueError, IndexError):
                new_number = 1
        else:
            new_number = 1
        entry_number = f"JE-{today.strftime('%Y%m%d')}-{new_number:03d}"
        
        # Create journal entry
        journal_entry = JournalEntry.objects.create(
            entry_number=entry_number,
            entry_date=ap_entry.invoice_date,
            description=f'AP Entry: {ap_entry.vendor_name} - {ap_entry.invoice_number}',
            reference_number=ap_entry.invoice_number,
            status='draft',
            fiscal_period=fiscal_period,
            created_by='system'
        )
        
        # Get AP account (liability)
        ap_account = ap_entry.account
        if not ap_account:
            ap_account = Account.objects.filter(account_type='liability', account_number__startswith='2000').first()
            if not ap_account:
                ap_account = Account.objects.create(
                    account_number='2000',
                    name='Accounts Payable',
                    account_type='liability',
                    description='Accounts Payable'
                )
        
        # Use Raw Materials Used account (5000) for purchases
        # In a full implementation, this would be determined by the PO items
        expense_account = Account.objects.filter(account_number='5000').first()
        if not expense_account:
            expense_account = Account.objects.filter(account_type='expense').first()
            if not expense_account:
                expense_account = Account.objects.create(
                    account_number='5000',
                    name='Raw Materials Used',
                    account_type='expense',
                    description='Monthly consumption'
                )
        
        # Create journal entry lines
        # Debit: Expense/Asset account
        JournalEntryLine.objects.create(
            journal_entry=journal_entry,
            account=expense_account,
            debit_credit='debit',
            amount=ap_entry.original_amount,
            description=f'Purchase from {ap_entry.vendor_name}'
        )
        
        # Credit: Accounts Payable
        JournalEntryLine.objects.create(
            journal_entry=journal_entry,
            account=ap_account,
            debit_credit='credit',
            amount=ap_entry.original_amount,
            description=f'AP Entry: {ap_entry.invoice_number}'
        )
        
        # Auto-post the journal entry
        try:
            # Validate debits and credits balance
            total_debits = sum(line.amount for line in journal_entry.lines.filter(debit_credit='debit'))
            total_credits = sum(line.amount for line in journal_entry.lines.filter(debit_credit='credit'))
            
            if abs(total_debits - total_credits) <= 0.01:  # Allow for floating point precision
                # Create GeneralLedgerEntry for each line
                for line in journal_entry.lines.all():
                    GeneralLedgerEntry.objects.create(
                        journal_entry=journal_entry,
                        journal_entry_line=line,
                        account=line.account,
                        fiscal_period=journal_entry.fiscal_period,
                        entry_date=journal_entry.entry_date,
                        description=line.description or journal_entry.description,
                        debit_credit=line.debit_credit,
                        amount=line.amount,
                        reference_number=journal_entry.reference_number,
                        reference_type='ap_entry'
                    )
                
                journal_entry.status = 'posted'
                journal_entry.save()
                
                # Update AccountBalances
                _update_account_balances(journal_entry)
        except Exception as post_error:
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f'Failed to auto-post journal entry {journal_entry.entry_number}: {str(post_error)}')
        
        return journal_entry
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f'Failed to create journal entry for AP entry {ap_entry.id}: {str(e)}')
        return None


def create_ar_journal_entry(ar_entry):
    """
    Create a journal entry when an AR entry is created.
    Debits: Accounts Receivable
    Credits: Revenue account
    """
    from django.utils import timezone
    
    try:
        # Get or create fiscal period
        today = timezone.now().date()
        fiscal_period = FiscalPeriod.objects.filter(
            start_date__lte=today,
            end_date__gte=today
        ).first()
        
        if not fiscal_period:
            # Create a default fiscal period if none exists
            from datetime import date
            year = today.year
            fiscal_period, _ = FiscalPeriod.objects.get_or_create(
                period_name=f"{year}-01",
                defaults={
                    'start_date': date(year, 1, 1),
                    'end_date': date(year, 12, 31),
                }
            )
        if fiscal_period.is_closed:
            return None
        # Generate journal entry number
        last_entry = JournalEntry.objects.filter(entry_date=today).order_by('-entry_number').first()
        if last_entry:
            try:
                last_number = int(last_entry.entry_number.split('-')[-1])
                new_number = last_number + 1
            except (ValueError, IndexError):
                new_number = 1
        else:
            new_number = 1
        entry_number = f"JE-{today.strftime('%Y%m%d')}-{new_number:03d}"
        
        # Create journal entry
        journal_entry = JournalEntry.objects.create(
            entry_number=entry_number,
            entry_date=ar_entry.invoice_date,
            description=f'AR Entry: {ar_entry.customer_name} - Invoice {ar_entry.invoice.invoice_number if ar_entry.invoice else "N/A"}',
            reference_number=ar_entry.invoice.invoice_number if ar_entry.invoice else None,
            status='draft',
            fiscal_period=fiscal_period,
            created_by='system'
        )
        
        # Get AR account (asset) - use Accounts Receivable (1100)
        ar_account = ar_entry.account
        if not ar_account:
            ar_account = Account.objects.filter(account_number='1100').first()  # Accounts Receivable
            if not ar_account:
                ar_account = Account.objects.filter(account_type='asset', account_number__startswith='11').first()
                if not ar_account:
                    ar_account = Account.objects.create(
                        account_number='1100',
                        name='Accounts Receivable',
                        account_type='asset',
                        description=''
                    )
        
        # Get revenue account - use Sales - Finished Goods (4000)
        revenue_account = Account.objects.filter(account_number='4000').first()
        if not revenue_account:
            revenue_account = Account.objects.filter(account_type='revenue').first()
            if not revenue_account:
                revenue_account = Account.objects.create(
                    account_number='4000',
                    name='Sales - Finished Goods',
                    account_type='revenue',
                    description='Parent income account'
                )
        
        # Create journal entry lines
        # Debit: Accounts Receivable
        JournalEntryLine.objects.create(
            journal_entry=journal_entry,
            account=ar_account,
            debit_credit='debit',
            amount=ar_entry.original_amount,
            description=f'AR Entry: Invoice {ar_entry.invoice.invoice_number if ar_entry.invoice else "N/A"}'
        )
        
        # Credit: Revenue
        JournalEntryLine.objects.create(
            journal_entry=journal_entry,
            account=revenue_account,
            debit_credit='credit',
            amount=ar_entry.original_amount,
            description=f'Sale to {ar_entry.customer_name}'
        )
        
        # Auto-post the journal entry
        try:
            # Validate debits and credits balance
            total_debits = sum(line.amount for line in journal_entry.lines.filter(debit_credit='debit'))
            total_credits = sum(line.amount for line in journal_entry.lines.filter(debit_credit='credit'))
            
            if abs(total_debits - total_credits) <= 0.01:  # Allow for floating point precision
                # Create GeneralLedgerEntry for each line
                for line in journal_entry.lines.all():
                    GeneralLedgerEntry.objects.create(
                        journal_entry=journal_entry,
                        journal_entry_line=line,
                        account=line.account,
                        fiscal_period=journal_entry.fiscal_period,
                        entry_date=journal_entry.entry_date,
                        description=line.description or journal_entry.description,
                        debit_credit=line.debit_credit,
                        amount=line.amount,
                        reference_number=journal_entry.reference_number,
                        reference_type='ar_entry'
                    )
                
                journal_entry.status = 'posted'
                journal_entry.save()
                
                # Update AccountBalances
                _update_account_balances(journal_entry)
        except Exception as post_error:
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f'Failed to auto-post journal entry {journal_entry.entry_number}: {str(post_error)}')
        
        return journal_entry
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f'Failed to create journal entry for AR entry {ar_entry.id}: {str(e)}')
        return None


def create_ap_payment_journal_entry(payment, ap_entry):
    """
    Create a journal entry when an AP payment is made.
    Debits: Accounts Payable
    Credits: Cash/Bank account
    """
    from django.utils import timezone
    
    try:
        # Get or create fiscal period
        today = timezone.now().date()
        fiscal_period = FiscalPeriod.objects.filter(
            start_date__lte=today,
            end_date__gte=today
        ).first()
        
        if not fiscal_period:
            from datetime import date
            year = today.year
            fiscal_period, _ = FiscalPeriod.objects.get_or_create(
                period_name=f"{year}-01",
                defaults={
                    'start_date': date(year, 1, 1),
                    'end_date': date(year, 12, 31),
                }
            )
        if fiscal_period.is_closed:
            return None
        # Generate journal entry number
        last_entry = JournalEntry.objects.filter(entry_date=today).order_by('-entry_number').first()
        if last_entry:
            try:
                last_number = int(last_entry.entry_number.split('-')[-1])
                new_number = last_number + 1
            except (ValueError, IndexError):
                new_number = 1
        else:
            new_number = 1
        entry_number = f"JE-{today.strftime('%Y%m%d')}-{new_number:03d}"
        
        # Get AP account
        ap_account = ap_entry.account
        if not ap_account:
            ap_account = Account.objects.filter(account_type='liability', account_number__startswith='2000').first()
        
        # Get cash/bank account from payment - prefer Checking Account (1000) or Savings (1010)
        cash_account = payment.account
        if not cash_account:
            cash_account = Account.objects.filter(account_number='1000').first()  # Checking Account
            if not cash_account:
                cash_account = Account.objects.filter(account_number='1010').first()  # Savings Account
            if not cash_account:
                cash_account = Account.objects.filter(account_type='asset', account_number__startswith='10').first()
                if not cash_account:
                    cash_account = Account.objects.create(
                        account_number='1000',
                        name='Checking Account',
                        account_type='asset',
                        description='Main operating bank account'
                    )
        
        # Create journal entry
        journal_entry = JournalEntry.objects.create(
            entry_number=entry_number,
            entry_date=payment.payment_date,
            description=f'AP Payment: {ap_entry.vendor_name} - {payment.reference_number or payment.payment_method}',
            reference_number=payment.reference_number or f"Payment-{payment.id}",
            status='draft',
            fiscal_period=fiscal_period,
            created_by='system'
        )
        
        # Create journal entry lines
        # Debit: Accounts Payable
        JournalEntryLine.objects.create(
            journal_entry=journal_entry,
            account=ap_account,
            debit_credit='debit',
            amount=payment.amount,
            description=f'Payment to {ap_entry.vendor_name}'
        )
        
        # Credit: Cash/Bank
        JournalEntryLine.objects.create(
            journal_entry=journal_entry,
            account=cash_account,
            debit_credit='credit',
            amount=payment.amount,
            description=f'Payment: {payment.reference_number or payment.payment_method}'
        )
        
        # Auto-post the journal entry
        total_debits = sum(line.amount for line in journal_entry.lines.filter(debit_credit='debit'))
        total_credits = sum(line.amount for line in journal_entry.lines.filter(debit_credit='credit'))
        
        if abs(total_debits - total_credits) <= 0.01:
            for line in journal_entry.lines.all():
                GeneralLedgerEntry.objects.create(
                    journal_entry=journal_entry,
                    account=line.account,
                    fiscal_period=journal_entry.fiscal_period,
                    entry_date=journal_entry.entry_date,
                    description=line.description or journal_entry.description,
                    debit=line.amount if line.debit_credit == 'debit' else 0.0,
                    credit=line.amount if line.debit_credit == 'credit' else 0.0
                )
            
            journal_entry.status = 'posted'
            journal_entry.save()
            _update_account_balances(journal_entry)
            
            # Link payment to journal entry
            payment.journal_entry = journal_entry
            payment.save()
        
        return journal_entry
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f'Failed to create journal entry for AP payment {payment.id}: {str(e)}')
        return None


def create_ar_payment_journal_entry(payment, ar_entry):
    """
    Create a journal entry when an AR payment is received.
    Debits: Cash/Bank account
    Credits: Accounts Receivable
    """
    from django.utils import timezone
    
    try:
        # Get or create fiscal period
        today = timezone.now().date()
        fiscal_period = FiscalPeriod.objects.filter(
            start_date__lte=today,
            end_date__gte=today
        ).first()
        
        if not fiscal_period:
            from datetime import date
            year = today.year
            fiscal_period, _ = FiscalPeriod.objects.get_or_create(
                period_name=f"{year}-01",
                defaults={
                    'start_date': date(year, 1, 1),
                    'end_date': date(year, 12, 31),
                }
            )
        if fiscal_period.is_closed:
            return None
        # Generate journal entry number
        last_entry = JournalEntry.objects.filter(entry_date=today).order_by('-entry_number').first()
        if last_entry:
            try:
                last_number = int(last_entry.entry_number.split('-')[-1])
                new_number = last_number + 1
            except (ValueError, IndexError):
                new_number = 1
        else:
            new_number = 1
        entry_number = f"JE-{today.strftime('%Y%m%d')}-{new_number:03d}"
        
        # Get AR account - use Accounts Receivable (1100)
        ar_account = ar_entry.account
        if not ar_account:
            ar_account = Account.objects.filter(account_number='1100').first()  # Accounts Receivable
            if not ar_account:
                ar_account = Account.objects.filter(account_type='asset', account_number__startswith='11').first()
        
        # Get cash/bank account from payment - prefer Checking Account (1000) or Savings (1010)
        cash_account = payment.account
        if not cash_account:
            cash_account = Account.objects.filter(account_number='1000').first()  # Checking Account
            if not cash_account:
                cash_account = Account.objects.filter(account_number='1010').first()  # Savings Account
            if not cash_account:
                cash_account = Account.objects.filter(account_type='asset', account_number__startswith='10').first()
                if not cash_account:
                    cash_account = Account.objects.create(
                        account_number='1000',
                        name='Checking Account',
                        account_type='asset',
                        description='Main operating bank account'
                    )
        
        # Create journal entry
        journal_entry = JournalEntry.objects.create(
            entry_number=entry_number,
            entry_date=payment.payment_date,
            description=f'AR Payment: {ar_entry.customer_name} - {payment.reference_number or payment.payment_method}',
            reference_number=payment.reference_number or f"Payment-{payment.id}",
            status='draft',
            fiscal_period=fiscal_period,
            created_by='system'
        )
        
        # Create journal entry lines
        # Debit: Cash/Bank
        JournalEntryLine.objects.create(
            journal_entry=journal_entry,
            account=cash_account,
            debit_credit='debit',
            amount=payment.amount,
            description=f'Payment from {ar_entry.customer_name}'
        )
        
        # Credit: Accounts Receivable
        JournalEntryLine.objects.create(
            journal_entry=journal_entry,
            account=ar_account,
            debit_credit='credit',
            amount=payment.amount,
            description=f'Payment: {payment.reference_number or payment.payment_method}'
        )
        
        # Auto-post the journal entry
        total_debits = sum(line.amount for line in journal_entry.lines.filter(debit_credit='debit'))
        total_credits = sum(line.amount for line in journal_entry.lines.filter(debit_credit='credit'))
        
        if abs(total_debits - total_credits) <= 0.01:
            for line in journal_entry.lines.all():
                GeneralLedgerEntry.objects.create(
                    journal_entry=journal_entry,
                    journal_entry_line=line,
                    account=line.account,
                    fiscal_period=journal_entry.fiscal_period,
                    entry_date=journal_entry.entry_date,
                    description=line.description or journal_entry.description,
                    debit_credit=line.debit_credit,
                    amount=line.amount,
                    reference_number=journal_entry.reference_number,
                    reference_type='ar_payment'
                )
            
            journal_entry.status = 'posted'
            journal_entry.save()
            _update_account_balances(journal_entry)
            
            # Link payment to journal entry
            payment.journal_entry = journal_entry
            payment.save()
        
        return journal_entry
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f'Failed to create journal entry for AR payment {payment.id}: {str(e)}')
        return None


def _update_account_balances(journal_entry):
    """Helper to update account balances after posting a journal entry."""
    if not journal_entry.fiscal_period:
        return
    
    for line in journal_entry.lines.all():
        account_balance, created = AccountBalance.objects.get_or_create(
            account=line.account,
            fiscal_period=journal_entry.fiscal_period,
            defaults={'opening_balance': 0.0, 'closing_balance': 0.0, 'period_debits': 0.0, 'period_credits': 0.0}
        )
        
        if line.debit_credit == 'debit':
            account_balance.period_debits += line.amount
        else:
            account_balance.period_credits += line.amount
        
        # Calculate closing balance based on account type
        if line.account.account_type in ['asset', 'expense']:
            account_balance.closing_balance = account_balance.opening_balance + account_balance.period_debits - account_balance.period_credits
        elif line.account.account_type in ['liability', 'equity', 'revenue']:
            account_balance.closing_balance = account_balance.opening_balance + account_balance.period_credits - account_balance.period_debits
        
        account_balance.save()


def create_ar_entry_from_invoice(invoice):
    """
    Create an Accounts Receivable entry when an invoice is created.
    
    Args:
        invoice: Invoice instance
    """
    from django.utils import timezone
    
    try:
        # Check if AR entry already exists for this invoice
        if AccountsReceivable.objects.filter(invoice=invoice).exists():
            return AccountsReceivable.objects.get(invoice=invoice)
        
        # Get customer information
        customer_name = None
        customer_id = None
        sales_order = None
        
        # Try to get customer from sales_order
        if hasattr(invoice, 'sales_order') and invoice.sales_order:
            sales_order = invoice.sales_order
            if sales_order.customer:
                customer_name = sales_order.customer.name
                customer_id = str(sales_order.customer.id)
            else:
                customer_name = sales_order.customer_name
                customer_id = sales_order.customer_legacy_id
        
        # Fallback to invoice customer fields
        if not customer_name:
            customer_name = getattr(invoice, 'customer_vendor_name', None) or 'Unknown Customer'
            customer_id = getattr(invoice, 'customer_vendor_id', None)
        
        # Calculate totals
        total_amount = getattr(invoice, 'grand_total', None) or getattr(invoice, 'total_amount', None) or 0.0
        if total_amount <= 0:
            # Calculate from items if total is not set
            total_amount = getattr(invoice, 'subtotal', 0.0) or 0.0
            total_amount += getattr(invoice, 'freight', 0.0) or 0.0
            total_amount += getattr(invoice, 'tax', 0.0) or getattr(invoice, 'tax_amount', 0.0) or 0.0
            total_amount -= getattr(invoice, 'discount', 0.0) or 0.0
        
        # Get dates
        invoice_date = getattr(invoice, 'invoice_date', None) or timezone.now().date()
        due_date = getattr(invoice, 'due_date', None) or invoice_date
        
        # Get or create AR account - use Accounts Receivable (1100)
        ar_account = None
        try:
            ar_account = Account.objects.filter(account_number='1100').first()  # Accounts Receivable
            if not ar_account:
                ar_account = Account.objects.filter(account_type='asset', account_number__startswith='11').first()
            if not ar_account:
                # Create a default AR account if none exists
                ar_account = Account.objects.create(
                    account_number='1100',
                    name='Accounts Receivable',
                    account_type='asset',
                    description=''
                )
        except Exception:
            pass
        
        # Create AR entry
        ar_entry = AccountsReceivable.objects.create(
            customer_name=customer_name,
            customer_id=customer_id,
            invoice=invoice,
            sales_order=sales_order,
            invoice_date=invoice_date,
            due_date=due_date,
            original_amount=total_amount,
            amount_paid=0.0,
            balance=total_amount,
            status='open',
            account=ar_account,
            notes=f'Created from invoice {invoice.invoice_number}'
        )
        
        # Auto-create journal entry for AR
        try:
            journal_entry = create_ar_journal_entry(ar_entry)
            if journal_entry:
                ar_entry.journal_entry = journal_entry
                ar_entry.save()
        except Exception as je_error:
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f'Failed to create journal entry for AR entry {ar_entry.id}: {str(je_error)}')
        
        return ar_entry
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f'Failed to create AR entry for invoice {invoice.invoice_number}: {str(e)}')
        return None


def log_production_batch_closure(batch, notes=None):
    """
    Log when a production batch is closed.
    
    Args:
        batch: The ProductionBatch instance
        notes: Additional context
    """
    try:
        import json
        
        # Get input materials information
        input_materials = []
        input_lots = []
        for input_item in batch.inputs.all():
            input_materials.append({
                'item_sku': input_item.lot.item.sku,
                'item_name': input_item.lot.item.name,
                'quantity_used': input_item.quantity_used
            })
            input_lots.append(input_item.lot.lot_number)
        
        # Get output lot information
        output_lot_number = None
        output_quantity = None
        output = batch.outputs.first()
        if output:
            output_lot_number = output.lot.lot_number
            output_quantity = output.quantity_produced
        
        # Extract QC information from batch notes
        qc_parameters = None
        qc_actual = None
        qc_initials = None
        if batch.notes:
            import re
            # Look for QC Parameters: ... pattern
            qc_params_match = re.search(r'QC Parameters:\s*(.+?)(?:\n|$)', batch.notes)
            if qc_params_match:
                qc_parameters = qc_params_match.group(1).strip()
            
            # Look for QC Actual: ... pattern
            qc_actual_match = re.search(r'QC Actual:\s*(.+?)(?:\n|$)', batch.notes)
            if qc_actual_match:
                qc_actual = qc_actual_match.group(1).strip()
            
            # Look for QC Initials: ... pattern
            qc_initials_match = re.search(r'QC Initials:\s*(.+?)(?:\n|$)', batch.notes)
            if qc_initials_match:
                qc_initials = qc_initials_match.group(1).strip()
        
        # Create notes that clarify if this is a repack of a distributed item
        item_type_label = 'Distributed Item' if batch.batch_type == 'repack' and batch.finished_good_item.item_type == 'distributed_item' else 'Finished Good'
        closure_notes = notes or f'Batch {batch.batch_number} closed'
        if batch.batch_type == 'repack' and batch.finished_good_item.item_type == 'distributed_item':
            closure_notes = f'Repack/Relabel of distributed item (not production). {closure_notes}'
        
        ProductionLog.objects.create(
            batch=batch,
            batch_number=batch.batch_number,
            batch_type=batch.batch_type,
            finished_good_sku=batch.finished_good_item.sku,
            finished_good_name=batch.finished_good_item.name,
            quantity_produced=batch.quantity_produced,
            quantity_actual=batch.quantity_actual,
            variance=batch.variance,
            wastes=batch.wastes,
            spills=batch.spills,
            unit_of_measure=batch.finished_good_item.unit_of_measure,
            production_date=batch.production_date,
            closed_date=batch.closed_date or timezone.now(),
            input_materials=json.dumps(input_materials),
            input_lots=json.dumps(input_lots),
            output_lot_number=output_lot_number,
            output_quantity=output_quantity,
            qc_parameters=qc_parameters,
            qc_actual=qc_actual,
            qc_initials=qc_initials,
            notes=closure_notes,
            recipe_snapshot=getattr(batch, 'recipe_snapshot', None) or None
        )
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f'Failed to log production batch closure for batch {batch.batch_number}: {str(e)}')


def generate_lot_number():
    """Generate a unique lot number in format 1yy00000 (7 digits: 1 + year + 5-digit sequence)"""
    from django.db import transaction
    from .models import LotNumberSequence
    
    today = timezone.now()
    year_prefix = today.strftime('%y')  # 2-digit year
    
    # Use select_for_update to lock the row and prevent race conditions
    with transaction.atomic():
        # Get or create sequence for this year with lock
        sequence, created = LotNumberSequence.objects.select_for_update().get_or_create(
            year_prefix=year_prefix,
            defaults={'sequence_number': 0}
        )
        
        # Increment sequence
        sequence.sequence_number += 1
        sequence.save()
        
        # Format: 1 + yy + 5-digit sequence (1yy00000)
        lot_number = f"1{year_prefix}{sequence.sequence_number:05d}"
        
        # Double-check uniqueness (in case of any edge case)
        max_retries = 10
        retry_count = 0
        while Lot.objects.filter(lot_number=lot_number).exists() and retry_count < max_retries:
            sequence.sequence_number += 1
            sequence.save()
            lot_number = f"1{year_prefix}{sequence.sequence_number:05d}"
            retry_count += 1
        
        if retry_count >= max_retries:
            # Fallback: add timestamp milliseconds for uniqueness
            import time
            lot_number = f"1{year_prefix}{sequence.sequence_number:05d}{int(time.time() * 1000) % 1000:03d}"
    
    return lot_number

def generate_po_number():
    """Generate a unique PO number in format 2yy000 (6 digits: 2 + year + 3-digit sequence)"""
    from django.db import transaction
    from .models import PONumberSequence
    
    today = timezone.now()
    year_prefix = today.strftime('%y')  # 2-digit year
    
    with transaction.atomic():
        sequence, created = PONumberSequence.objects.select_for_update().get_or_create(
            year_prefix=year_prefix,
            defaults={'sequence_number': 0}
        )
        
        sequence.sequence_number += 1
        sequence.save()
        
        # Format: 2 + yy + 3-digit sequence (2yy000)
        po_number = f"2{year_prefix}{sequence.sequence_number:03d}"
        
        # Double-check uniqueness
        max_retries = 10
        retry_count = 0
        from .models import PurchaseOrder
        while PurchaseOrder.objects.filter(po_number=po_number).exists() and retry_count < max_retries:
            sequence.sequence_number += 1
            sequence.save()
            po_number = f"2{year_prefix}{sequence.sequence_number:03d}"
            retry_count += 1
    
    return po_number

def generate_sales_order_number():
    """Generate a unique sales order number in format 3yy0000 (7 digits: 3 + year + 4-digit sequence)"""
    from django.db import transaction
    from .models import SalesOrderNumberSequence
    
    today = timezone.now()
    year_prefix = today.strftime('%y')  # 2-digit year
    
    with transaction.atomic():
        sequence, created = SalesOrderNumberSequence.objects.select_for_update().get_or_create(
            year_prefix=year_prefix,
            defaults={'sequence_number': 0}
        )
        
        sequence.sequence_number += 1
        sequence.save()
        
        # Format: 3 + yy + 4-digit sequence (3yy0000)
        so_number = f"3{year_prefix}{sequence.sequence_number:04d}"
        
        # Double-check uniqueness
        max_retries = 10
        retry_count = 0
        from .models import SalesOrder
        while SalesOrder.objects.filter(so_number=so_number).exists() and retry_count < max_retries:
            sequence.sequence_number += 1
            sequence.save()
            so_number = f"3{year_prefix}{sequence.sequence_number:04d}"
            retry_count += 1
    
    return so_number

def generate_customer_id():
    """Generate a unique customer ID in format 001, 002, etc. (3-digit sequence)"""
    from django.db import transaction
    from django.db.utils import OperationalError
    from .models import Customer, CustomerNumberSequence
    
    try:
        # Try to use the sequence table if it exists
        with transaction.atomic():
            # Get or create the single sequence record with lock
            sequence, created = CustomerNumberSequence.objects.select_for_update().get_or_create(
                id=1,  # Single sequence record
                defaults={'sequence_number': 0}
            )
            
            sequence.sequence_number += 1
            sequence.save()
            
            # Format: 3-digit sequence (001, 002, etc.)
            customer_id = f"{sequence.sequence_number:03d}"
            
            # Double-check uniqueness
            max_retries = 10
            retry_count = 0
            while Customer.objects.filter(customer_id=customer_id).exists() and retry_count < max_retries:
                sequence.sequence_number += 1
                sequence.save()
                customer_id = f"{sequence.sequence_number:03d}"
                retry_count += 1
            
            return customer_id
    except (OperationalError, Exception) as e:
        # If sequence table doesn't exist yet, fall back to finding max customer_id
        # This handles the case where migrations haven't been run yet
        try:
            # Get all existing customer IDs and find the highest numeric one
            existing_customers = Customer.objects.exclude(customer_id__isnull=True).exclude(customer_id='')
            max_id = 0
            
            for customer in existing_customers:
                try:
                    # Try to parse customer_id as integer
                    customer_num = int(customer.customer_id)
                    if customer_num > max_id:
                        max_id = customer_num
                except (ValueError, TypeError):
                    # If customer_id is not numeric, skip it
                    continue
            
            # Generate next ID
            next_id = max_id + 1
            customer_id = f"{next_id:03d}"
            
            # Double-check uniqueness
            max_retries = 10
            retry_count = 0
            while Customer.objects.filter(customer_id=customer_id).exists() and retry_count < max_retries:
                next_id += 1
                customer_id = f"{next_id:03d}"
                retry_count += 1
            
            return customer_id
        except Exception:
            # Ultimate fallback: start from 001
            return "001"

def generate_invoice_number():
    """Generate a unique invoice number in format 4yy0000 (7 digits: 4 + year + 4-digit sequence)"""
    from django.db import transaction
    from .models import InvoiceNumberSequence
    
    today = timezone.now()
    year_prefix = today.strftime('%y')  # 2-digit year
    
    with transaction.atomic():
        sequence, created = InvoiceNumberSequence.objects.select_for_update().get_or_create(
            year_prefix=year_prefix,
            defaults={'sequence_number': 0}
        )
        
        sequence.sequence_number += 1
        sequence.save()
        
        # Format: 4 + yy + 4-digit sequence (4yy0000)
        invoice_number = f"4{year_prefix}{sequence.sequence_number:04d}"
        
        # Double-check uniqueness (if Invoice model exists)
        try:
            from .models import Invoice
            max_retries = 10
            retry_count = 0
            while Invoice.objects.filter(invoice_number=invoice_number).exists() and retry_count < max_retries:
                sequence.sequence_number += 1
                sequence.save()
                invoice_number = f"4{year_prefix}{sequence.sequence_number:04d}"
                retry_count += 1
        except ImportError:
            # Invoice model doesn't exist yet, that's okay
            pass
    
    return invoice_number

def generate_batch_number(batch_type='production'):
    """Generate a unique batch number in format BT-YYYYMMDD-001 or R-YYYYMMDD-001"""
    from django.db import transaction
    from .models import BatchNumberSequence
    
    today = timezone.now()
    date_prefix = today.strftime('%Y%m%d')  # YYYYMMDD
    prefix = 'R' if batch_type == 'repack' else 'BT'
    
    # Use select_for_update to lock the row and prevent race conditions
    with transaction.atomic():
        # Get or create sequence for this date with lock
        sequence, created = BatchNumberSequence.objects.select_for_update().get_or_create(
            date_prefix=date_prefix,
            defaults={'sequence_number': 0}
        )
        
        # Increment sequence
        sequence.sequence_number += 1
        sequence.save()
        
        # Format: PREFIX-YYYYMMDD-001
        batch_number = f"{prefix}-{date_prefix}-{sequence.sequence_number:03d}"
        
        # Double-check uniqueness (in case of any edge case)
        max_retries = 10
        retry_count = 0
        while ProductionBatch.objects.filter(batch_number=batch_number).exists() and retry_count < max_retries:
            sequence.sequence_number += 1
            sequence.save()
            batch_number = f"{prefix}-{date_prefix}-{sequence.sequence_number:03d}"
            retry_count += 1
        
        if retry_count >= max_retries:
            # Fallback: add timestamp milliseconds for uniqueness
            import time
            batch_number = f"{prefix}-{date_prefix}-{sequence.sequence_number:03d}-{int(time.time() * 1000) % 1000:03d}"
    
    return batch_number


class ItemViewSet(viewsets.ModelViewSet):
    queryset = Item.objects.all()
    serializer_class = ItemSerializer
    
    def get_queryset(self):
        queryset = Item.objects.all()
        # Filter by approved vendors if requested
        approved_only = self.request.query_params.get('approved_vendors_only', None)
        if approved_only == 'true':
            approved_vendor_names = Vendor.objects.filter(approval_status='approved').values_list('name', flat=True)
            queryset = queryset.filter(vendor__in=approved_vendor_names)
        return queryset
    
    def create(self, request, *args, **kwargs):
        # Clean up the data - convert empty strings to None for optional fields
        data = request.data.copy()
        if 'description' in data and data['description'] == '':
            data['description'] = None
        if 'vendor' in data and data['vendor'] == '':
            data['vendor'] = None
        if 'pack_size' in data and (data['pack_size'] == '' or data['pack_size'] is None):
            data.pop('pack_size', None)
        if 'price' in data and (data['price'] == '' or data['price'] is None):
            data.pop('price', None)
        
        # Ensure on_order is set to 0 if not provided
        if 'on_order' not in data:
            data['on_order'] = 0
        
        # Check for duplicate SKU + vendor combination
        if data.get('sku') and data.get('vendor'):
            existing = Item.objects.filter(sku=data['sku'], vendor=data['vendor']).first()
            if existing:
                return Response(
                    {'error': f'Item with SKU "{data["sku"]}" already exists for vendor "{data["vendor"]}". Each vendor can only have one item per SKU.'},
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        serializer = self.get_serializer(data=data)
        serializer.is_valid(raise_exception=True)
        item = serializer.save()
        
        # Auto-create CostMaster entry when item is created
        if item.vendor:
            # Convert price based on unit of measure
            price_per_kg = None
            price_per_lb = None
            freight_per_kg = 0.0
            
            if data.get('price'):
                price = float(data['price'])
                unit_of_measure = data.get('unit_of_measure', 'lbs')
                
                if unit_of_measure == 'lbs':
                    price_per_lb = price
                    price_per_kg = price * 2.20462  # Convert lb to kg
                elif unit_of_measure == 'kg':
                    price_per_kg = price
                    price_per_lb = price / 2.20462  # Convert kg to lb
                else:
                    # For 'ea', use price as-is for both (or handle differently)
                    price_per_kg = price
                    price_per_lb = price
            
            # Calculate landed cost using Excel formula: (Price per kg * (1 + Tariff)) + Freight per kg
            # Tariff defaults to 0 if not provided
            tariff = float(data.get('tariff', 0)) if data.get('tariff') else 0.0
            landed_cost_per_kg = None
            landed_cost_per_lb = None
            if price_per_kg is not None:
                landed_cost_per_kg = (price_per_kg * (1 + tariff)) + freight_per_kg
            if price_per_lb is not None:
                # Convert freight to lb and apply same formula
                landed_cost_per_lb = (price_per_lb * (1 + tariff)) + (freight_per_kg / 2.20462)
            
            # Get or create CostMaster entry (one per SKU + vendor combination)
            # Tariff is manual only (HTS/country of origin kept for reference; no external API)
            tariff = float(data.get('tariff', 0)) if data.get('tariff') else 0.0
            item.tariff = tariff
            item.save()
            
            # Get vendor name - handle both Vendor object and string
            vendor_name = None
            if item.vendor:
                if hasattr(item.vendor, 'name'):
                    vendor_name = item.vendor.name
                else:
                    vendor_name = str(item.vendor)
            
            cost_master, created = CostMaster.objects.get_or_create(
                wwi_product_code=item.sku,
                vendor=vendor_name,
                defaults={
                    'vendor_material': item.name,
                    'price_per_kg': price_per_kg,
                    'price_per_lb': price_per_lb,
                    'freight_per_kg': freight_per_kg,
                    'tariff': tariff,
                    'hts_code': data.get('hts_code') or None,
                    'origin': data.get('country_of_origin') or None,
                    # Landed cost will be calculated automatically in save()
                }
            )
            # Update if it already existed
            if not created:
                cost_master.vendor_material = item.name
                if price_per_kg is not None:
                    cost_master.price_per_kg = price_per_kg
                if price_per_lb is not None:
                    cost_master.price_per_lb = price_per_lb
                cost_master.freight_per_kg = freight_per_kg
                cost_master.tariff = tariff
                # Update HTS code and origin if provided
                if data.get('hts_code'):
                    cost_master.hts_code = data['hts_code']
                if data.get('country_of_origin'):
                    cost_master.origin = data['country_of_origin']
                # Recalculate landed cost with the tariff
                # Formula: (Price per kg * (1 + Tariff)) + Freight per kg
                cost_master.calculate_landed_cost()
                cost_master.save()
        
        headers = self.get_success_headers(serializer.data)
        
        # Check for orphaned inventory that can be reassigned (matching SKU and vendor)
        self._reassign_orphaned_inventory(item)
        
        return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)
    
    def _reassign_orphaned_inventory(self, item):
        """Reassign orphaned inventory to a newly created item if SKU matches"""
        from .models import OrphanedInventory, OrphanedPurchaseOrderItem, Lot, PurchaseOrderItem
        from django.utils import timezone
        
        # Find orphaned inventory with matching SKU and vendor
        orphaned_lots = OrphanedInventory.objects.filter(
            original_item_sku=item.sku,
            original_item_vendor=item.vendor or '',
            reassigned_item__isnull=True
        )
        
        for orphaned_lot in orphaned_lots:
            # Create a new lot with the orphaned inventory data
            new_lot = Lot.objects.create(
                lot_number=orphaned_lot.lot_number,
                vendor_lot_number=orphaned_lot.vendor_lot_number,
                item=item,
                quantity=orphaned_lot.quantity,
                quantity_remaining=orphaned_lot.quantity_remaining,
                received_date=orphaned_lot.received_date,
                expiration_date=orphaned_lot.expiration_date,
                status=orphaned_lot.status,
                po_number=orphaned_lot.po_number,
                freight_actual=orphaned_lot.freight_actual,
                short_reason=orphaned_lot.short_reason
            )
            
            # Mark as reassigned
            orphaned_lot.reassigned_item = item
            orphaned_lot.reassigned_at = timezone.now()
            orphaned_lot.save()
        
        # Find orphaned PO items with matching SKU and vendor
        orphaned_po_items = OrphanedPurchaseOrderItem.objects.filter(
            original_item_sku=item.sku,
            original_item_vendor=item.vendor or '',
            reassigned_item__isnull=True
        )
        
        for orphaned_po_item in orphaned_po_items:
            # Find PO items without an item reference (orphaned)
            po_items = PurchaseOrderItem.objects.filter(
                purchase_order=orphaned_po_item.purchase_order,
                item__isnull=True
            )
            
            # Try to match by quantity and unit price
            matching_po_item = po_items.filter(
                quantity_ordered=orphaned_po_item.quantity_ordered,
                unit_price=orphaned_po_item.unit_price
            ).first()
            
            if matching_po_item:
                matching_po_item.item = item
                matching_po_item.save()
            else:
                # Create new PO item if no match found
                PurchaseOrderItem.objects.create(
                    purchase_order=orphaned_po_item.purchase_order,
                    item=item,
                    quantity_ordered=orphaned_po_item.quantity_ordered,
                    quantity_received=orphaned_po_item.quantity_received,
                    unit_price=orphaned_po_item.unit_price,
                    notes=orphaned_po_item.notes
                )
            
            # Mark as reassigned
            orphaned_po_item.reassigned_item = item
            orphaned_po_item.reassigned_at = timezone.now()
            orphaned_po_item.save()
    
    @action(detail=True, methods=['post'], url_path='reassign-orphaned-inventory')
    def reassign_orphaned_inventory(self, request, pk=None):
        """Manually reassign orphaned inventory to this item"""
        item = self.get_object()
        
        # Get counts before reassignment
        from .models import OrphanedInventory, OrphanedPurchaseOrderItem
        orphaned_lots_count = OrphanedInventory.objects.filter(
            original_item_sku=item.sku,
            original_item_vendor=item.vendor or '',
            reassigned_item__isnull=True
        ).count()
        orphaned_po_items_count = OrphanedPurchaseOrderItem.objects.filter(
            original_item_sku=item.sku,
            original_item_vendor=item.vendor or '',
            reassigned_item__isnull=True
        ).count()
        
        # Perform reassignment
        self._reassign_orphaned_inventory(item)
        
        # Get counts after reassignment
        reassigned_lots = OrphanedInventory.objects.filter(
            original_item_sku=item.sku,
            original_item_vendor=item.vendor or '',
            reassigned_item=item
        ).count()
        reassigned_po_items = OrphanedPurchaseOrderItem.objects.filter(
            original_item_sku=item.sku,
            original_item_vendor=item.vendor or '',
            reassigned_item=item
        ).count()
        
        return Response({
            'message': 'Orphaned inventory reassigned successfully',
            'lots_reassigned': reassigned_lots,
            'po_items_reassigned': reassigned_po_items,
            'lots_found': orphaned_lots_count,
            'po_items_found': orphaned_po_items_count
        }, status=status.HTTP_200_OK)
    
    @action(detail=False, methods=['get'], url_path='orphaned-inventory')
    def list_orphaned_inventory(self, request):
        """List all orphaned inventory that hasn't been reassigned"""
        from .models import OrphanedInventory, OrphanedPurchaseOrderItem
        
        # Get query parameters
        sku = request.query_params.get('sku', None)
        vendor = request.query_params.get('vendor', None)
        
        # Filter orphaned lots
        orphaned_lots = OrphanedInventory.objects.filter(reassigned_item__isnull=True)
        if sku:
            orphaned_lots = orphaned_lots.filter(original_item_sku=sku)
        if vendor:
            orphaned_lots = orphaned_lots.filter(original_item_vendor=vendor or '')
        
        # Filter orphaned PO items
        orphaned_po_items = OrphanedPurchaseOrderItem.objects.filter(reassigned_item__isnull=True)
        if sku:
            orphaned_po_items = orphaned_po_items.filter(original_item_sku=sku)
        if vendor:
            orphaned_po_items = orphaned_po_items.filter(original_item_vendor=vendor or '')
        
        # Serialize results
        lots_data = [{
            'id': lot.id,
            'original_item_sku': lot.original_item_sku,
            'original_item_name': lot.original_item_name,
            'original_item_vendor': lot.original_item_vendor,
            'lot_number': lot.lot_number,
            'quantity_remaining': lot.quantity_remaining,
            'received_date': lot.received_date,
            'expiration_date': lot.expiration_date,
            'status': lot.status,
            'po_number': lot.po_number,
            'created_at': lot.created_at,
            'notes': lot.notes
        } for lot in orphaned_lots]
        
        po_items_data = [{
            'id': po_item.id,
            'original_item_sku': po_item.original_item_sku,
            'original_item_name': po_item.original_item_name,
            'original_item_vendor': po_item.original_item_vendor,
            'purchase_order_id': po_item.purchase_order_id,
            'purchase_order_number': po_item.purchase_order.po_number,
            'quantity_ordered': po_item.quantity_ordered,
            'quantity_received': po_item.quantity_received,
            'unit_price': po_item.unit_price,
            'created_at': po_item.created_at,
            'notes': po_item.notes
        } for po_item in orphaned_po_items]
        
        return Response({
            'orphaned_lots': lots_data,
            'orphaned_po_items': po_items_data,
            'total_orphaned_lots': len(lots_data),
            'total_orphaned_po_items': len(po_items_data)
        }, status=status.HTTP_200_OK)
    
    def update(self, request, *args, **kwargs):
        """Update item and sync to CostMaster, creating history if price changed"""
        instance = self.get_object()
        old_price = instance.price
        old_pack_size = instance.pack_size
        old_unit = instance.unit_of_measure
        
        # Clean up the data
        data = request.data.copy()
        if 'description' in data and data['description'] == '':
            data['description'] = None
        if 'vendor' in data and data['vendor'] == '':
            data['vendor'] = None
        if 'pack_size' in data and (data['pack_size'] == '' or data['pack_size'] is None):
            data.pop('pack_size', None)
        if 'price' in data and (data['price'] == '' or data['price'] is None):
            data.pop('price', None)
        
        # Trim SKU and name if provided
        if 'sku' in data and isinstance(data['sku'], str):
            data['sku'] = data['sku'].strip()
        if 'name' in data and isinstance(data['name'], str):
            data['name'] = data['name'].strip()
        
        # Check for unique constraint violation before updating
        if 'sku' in data or 'vendor' in data:
            new_sku = data.get('sku', instance.sku)
            new_vendor = data.get('vendor', instance.vendor) or None
            # Check if another item with same SKU + vendor exists (excluding current instance)
            existing_item = Item.objects.filter(sku=new_sku, vendor=new_vendor).exclude(id=instance.id).first()
            if existing_item:
                return Response(
                    {'error': f'An item with SKU "{new_sku}" and vendor "{new_vendor or "None"}" already exists.'},
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        # Update the item
        partial = kwargs.pop('partial', False)
        serializer = self.get_serializer(instance, data=data, partial=partial)
        if not serializer.is_valid():
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f'Serializer validation failed: {serializer.errors}')
            logger.error(f'Data being validated: {data}')
            return Response(
                serializer.errors,
                status=status.HTTP_400_BAD_REQUEST
            )
        item = serializer.save()
        
        # Sync to CostMaster if vendor exists
        if item.vendor and (data.get('price') is not None or data.get('pack_size') is not None or data.get('hts_code') is not None or data.get('country_of_origin') is not None):
            # Find or create CostMaster entry
            # Get vendor name - handle both Vendor object and string
            vendor_name = None
            if item.vendor:
                if hasattr(item.vendor, 'name'):
                    vendor_name = item.vendor.name
                else:
                    vendor_name = str(item.vendor)
            
            cost_master, created = CostMaster.objects.get_or_create(
                wwi_product_code=item.sku,
                vendor=vendor_name,
                defaults={
                    'vendor_material': item.name,
                    'vendor': vendor_name,
                }
            )
            
            # Update HTS code and origin if provided
            if data.get('hts_code'):
                cost_master.hts_code = data['hts_code']
            if data.get('country_of_origin'):
                cost_master.origin = data['country_of_origin']
            
            # Tariff is manual only (no external API)
            tariff_updated = False
            if data.get('tariff') is not None:
                tariff = float(data.get('tariff', 0)) if data.get('tariff') else 0.0
                cost_master.tariff = tariff
                item.tariff = tariff
                item.save()
                tariff_updated = True
            
            # Check if price changed
            price_changed = False
            new_price_per_kg = None
            new_price_per_lb = None
            
            if data.get('price') is not None:
                price = float(data['price'])
                unit_of_measure = data.get('unit_of_measure', item.unit_of_measure)
                
                if unit_of_measure == 'lbs':
                    new_price_per_lb = price
                    new_price_per_kg = price * 2.20462
                elif unit_of_measure == 'kg':
                    new_price_per_kg = price
                    new_price_per_lb = price / 2.20462
                else:
                    new_price_per_kg = price
                    new_price_per_lb = price
                
                # Check if price actually changed
                if (cost_master.price_per_kg != new_price_per_kg or 
                    cost_master.price_per_lb != new_price_per_lb):
                    price_changed = True
            
            # Update CostMaster
            if new_price_per_kg is not None:
                cost_master.price_per_kg = new_price_per_kg
            if new_price_per_lb is not None:
                cost_master.price_per_lb = new_price_per_lb
            if not created:
                cost_master.vendor_material = item.name
                # Get vendor name - handle both Vendor object and string
                vendor_name = None
                if item.vendor:
                    if hasattr(item.vendor, 'name'):
                        vendor_name = item.vendor.name
                    else:
                        vendor_name = str(item.vendor)
                cost_master.vendor = vendor_name
            
            # Update HTS code and origin if provided in update
            if data.get('hts_code') is not None:
                cost_master.hts_code = data['hts_code']
            if data.get('country_of_origin') is not None:
                cost_master.origin = data['country_of_origin']
            
            # Tariff is manual only (no external API)
            if data.get('tariff') is not None:
                tariff = float(data.get('tariff', 0)) if data.get('tariff') else 0.0
                cost_master.tariff = tariff
                item.tariff = tariff
                item.save()
            
            # Recalculate landed cost using Excel formula: (Price per kg * (1 + Tariff)) + Freight per kg
            cost_master.calculate_landed_cost()
            cost_master.save()
            
            # Create history record if price changed
            if price_changed:
                CostMasterHistory.objects.create(
                    cost_master=cost_master,
                    price_per_kg=cost_master.price_per_kg,
                    price_per_lb=cost_master.price_per_lb,
                    changed_by=request.user.username if hasattr(request.user, 'username') else 'system',
                    notes=f'Updated from Item {item.sku} - Price: {old_price} -> {item.price}, Pack Size: {old_pack_size} -> {item.pack_size}'
                )
        
        return Response(serializer.data)
    
    def destroy(self, request, *args, **kwargs):
        """Delete an item, but preserve inventory by creating orphaned inventory entries"""
        from .models import OrphanedInventory, OrphanedPurchaseOrderItem
        from django.utils import timezone
        
        item = self.get_object()
        
        # Store item details before deletion
        item_sku = item.sku
        item_name = item.name
        item_vendor = item.vendor or ''
        item_type = item.item_type
        item_unit = item.unit_of_measure
        
        # Check for active inventory
        has_lots = item.lots.filter(quantity_remaining__gt=0).exists()
        has_po_items = item.purchase_order_items.filter(purchase_order__status__in=['draft', 'issued']).exists()
        
        # Create orphaned inventory entries for lots with remaining quantity
        if has_lots:
            for lot in item.lots.filter(quantity_remaining__gt=0):
                OrphanedInventory.objects.create(
                    original_item_sku=item_sku,
                    original_item_name=item_name,
                    original_item_vendor=item_vendor,
                    original_item_type=item_type,
                    original_item_unit=item_unit,
                    lot_number=lot.lot_number,
                    vendor_lot_number=lot.vendor_lot_number,
                    quantity=lot.quantity,
                    quantity_remaining=lot.quantity_remaining,
                    received_date=lot.received_date,
                    expiration_date=lot.expiration_date,
                    status=lot.status,
                    po_number=lot.po_number,
                    freight_actual=lot.freight_actual,
                    short_reason=lot.short_reason,
                    notes=f'Item deleted on {timezone.now().strftime("%Y-%m-%d %H:%M:%S")}'
                )
        
        # Create orphaned PO item entries for active purchase orders
        if has_po_items:
            for po_item in item.purchase_order_items.filter(purchase_order__status__in=['draft', 'issued']):
                OrphanedPurchaseOrderItem.objects.create(
                    original_item_sku=item_sku,
                    original_item_name=item_name,
                    original_item_vendor=item_vendor,
                    original_item_unit=item_unit,
                    purchase_order=po_item.purchase_order,
                    quantity_ordered=po_item.quantity_ordered,
                    quantity_received=po_item.quantity_received,
                    unit_price=po_item.unit_price,
                    notes=po_item.notes or f'Item deleted on {timezone.now().strftime("%Y-%m-%d %H:%M:%S")}'
                )
                # Clear the item reference but keep the PO item
                po_item.item = None
                po_item.save()
        
        # Now delete the item (which will cascade delete lots and other references)
        # But we've already preserved the important data in orphaned inventory
        return super().destroy(request, *args, **kwargs)


class LotViewSet(viewsets.ModelViewSet):
    queryset = Lot.objects.select_related('item').all()
    serializer_class = LotSerializer
    
    @action(detail=False, methods=['get'])
    def lots_by_sku_vendor(self, request):
        """Get lot details for a specific SKU and vendor"""
        sku = request.query_params.get('sku', None)
        vendor = request.query_params.get('vendor', None)
        
        if not sku:
            return Response({'error': 'SKU parameter is required'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Get items with this SKU
        items = Item.objects.filter(sku=sku)
        if vendor:
            # Handle "Unknown" vendor - this means items with null/empty vendor
            if vendor == "Unknown":
                from django.db.models import Q
                items = items.filter(Q(vendor__isnull=True) | Q(vendor=''))
            else:
                items = items.filter(vendor=vendor)
        
        if not items.exists():
            return Response({'error': 'No items found for this SKU/vendor combination'}, status=status.HTTP_404_NOT_FOUND)
        
        # Get all lots for these items (accepted and on_hold; visible = qty>0 or depleted <24h)
        item_ids = items.values_list('id', flat=True)
        from django.utils import timezone
        from datetime import timedelta
        from django.db.models import Q
        now = timezone.now()
        depleted_cutoff = now - timedelta(hours=24)
        try:
            lot_visible_q = Q(quantity_remaining__gt=0) | Q(depleted_at__gte=depleted_cutoff)
            lots = Lot.objects.filter(
                item_id__in=item_ids
            ).filter(status__in=['accepted', 'on_hold']).filter(lot_visible_q).select_related('item').order_by('-received_date')
        except Exception:
            lots = Lot.objects.filter(
                item_id__in=item_ids,
                quantity_remaining__gt=0
            ).filter(status__in=['accepted', 'on_hold']).select_related('item').order_by('-received_date')
        
        # Get PO tracking info for lots with PO numbers
        from .models import PurchaseOrder
        po_tracking_map = {}
        po_numbers = [lot.po_number for lot in lots if lot.po_number]
        if po_numbers:
            pos = PurchaseOrder.objects.filter(po_number__in=po_numbers)
            for po in pos:
                po_tracking_map[po.po_number] = {
                    'tracking_number': po.tracking_number,
                    'carrier': po.carrier
                }
        
        serializer = self.get_serializer(lots, many=True)
        data = serializer.data
        
        # Add PO tracking info and commitment info to each lot
        from .models import SalesOrderLot, ProductionBatchInput
        from django.db.models import Sum
        
        for i, lot_data in enumerate(data):
            lot = lots[i]  # Get the actual lot object
            po_number = lot_data.get('po_number')
            if po_number and po_number in po_tracking_map:
                tracking_info = po_tracking_map[po_number]
                lot_data['po_tracking_number'] = tracking_info['tracking_number']
                lot_data['po_carrier'] = tracking_info['carrier']
            
            # Check if lot is committed to sales (allocated but not shipped)
            committed_to_sales_qty = 0.0
            try:
                sales_allocations = SalesOrderLot.objects.filter(
                    lot=lot,
                    sales_order_item__sales_order__status__in=[
                        'draft', 'allocated', 'issued', 'ready_for_shipment'
                    ]
                ).aggregate(total=Sum('quantity_allocated'))
                committed_to_sales_qty = sales_allocations['total'] or 0.0
            except Exception:
                pass
            
            # Check if lot is committed to production (in_progress batches)
            committed_to_production_qty = 0.0
            try:
                production_inputs = ProductionBatchInput.objects.filter(
                    lot=lot,
                    batch__status='in_progress'
                ).aggregate(total=Sum('quantity_used'))
                committed_to_production_qty = production_inputs['total'] or 0.0
            except Exception:
                pass
            
            # Add commitment quantities
            lot_data['committed_to_sales_qty'] = committed_to_sales_qty
            lot_data['committed_to_production_qty'] = committed_to_production_qty
        
        return Response(data)
    
    @action(detail=True, methods=['post'])
    def checkout_indirect_material(self, request, pk=None):
        """Checkout indirect material from a lot (simple checkout without batch)"""
        from .models import InventoryTransaction
        
        lot = self.get_object()
        
        # Verify it's an indirect material
        if lot.item.item_type != 'indirect_material':
            return Response(
                {'error': 'This lot is not an indirect material'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        quantity = request.data.get('quantity')
        notes = request.data.get('notes', '')
        reference_number = request.data.get('reference_number', '')
        
        if not quantity or float(quantity) <= 0:
            return Response(
                {'error': 'Valid quantity is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        quantity = round(float(quantity), 2)
        
        # Check available quantity
        if lot.quantity_remaining < quantity:
            return Response(
                {'error': f'Insufficient quantity. Available: {lot.quantity_remaining}, Requested: {quantity}'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Create inventory transaction
        quantity_before = lot.quantity_remaining
        transaction = InventoryTransaction.objects.create(
            transaction_type='indirect_material_checkout',
            lot=lot,
            quantity=round(-quantity, 2),
            notes=notes or f'Indirect material checkout - {lot.item.name}',
            reference_number=reference_number
        )
        
        # Log the transaction
        log_lot_transaction(
            lot=lot,
            quantity_before=quantity_before,
            quantity_change=-quantity,
            transaction_type='indirect_material_checkout',
            reference_number=reference_number,
            reference_type='checkout',
            transaction_id=transaction.id,
            notes=notes or f'Indirect material checkout - {lot.item.name}'
        )
        
        # Update lot quantity_remaining
        lot.quantity_remaining = round(lot.quantity_remaining - quantity, 2)
        lot.save()
        
        serializer = self.get_serializer(lot)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=['post'])
    def put_on_hold(self, request, pk=None):
        """Put a partial or full amount of this lot on hold. Body: { quantity }."""
        lot = self.get_object()
        try:
            quantity = round(float(request.data.get('quantity', 0)), 2)
        except (TypeError, ValueError):
            return Response({'error': 'Valid quantity is required'}, status=status.HTTP_400_BAD_REQUEST)
        if quantity <= 0:
            return Response({'error': 'Quantity must be positive'}, status=status.HTTP_400_BAD_REQUEST)
        from django.db.models import Sum
        from .models import SalesOrderLot, ProductionBatchInput
        sales_alloc = SalesOrderLot.objects.filter(
            lot=lot,
            sales_order_item__sales_order__status__in=[
                'draft', 'allocated', 'issued', 'ready_for_shipment', 'shipped'
            ]
        ).aggregate(total=Sum('quantity_allocated'))['total'] or 0.0
        prod_alloc = ProductionBatchInput.objects.filter(
            lot=lot, batch__status='in_progress'
        ).aggregate(total=Sum('quantity_used'))['total'] or 0.0
        available = max(0.0, lot.quantity_remaining - sales_alloc - prod_alloc - getattr(lot, 'quantity_on_hold', 0.0))
        if quantity > available:
            return Response(
                {'error': f'Only {available} available to put on hold (remaining minus sales/prod allocations and current on hold)'},
                status=status.HTTP_400_BAD_REQUEST
            )
        current_hold = getattr(lot, 'quantity_on_hold', 0.0)
        lot.quantity_on_hold = round(current_hold + quantity, 2)
        lot.on_hold = True
        if lot.quantity_on_hold >= lot.quantity_remaining:
            lot.status = 'on_hold'
        lot.save(update_fields=['quantity_on_hold', 'on_hold', 'status'])
        serializer = self.get_serializer(lot)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=['post'])
    def release_from_hold(self, request, pk=None):
        """Release a quantity from hold. Body: { quantity }."""
        lot = self.get_object()
        try:
            quantity = round(float(request.data.get('quantity', 0)), 2)
        except (TypeError, ValueError):
            return Response({'error': 'Valid quantity is required'}, status=status.HTTP_400_BAD_REQUEST)
        if quantity <= 0:
            return Response({'error': 'Quantity must be positive'}, status=status.HTTP_400_BAD_REQUEST)
        current_hold = getattr(lot, 'quantity_on_hold', 0.0)
        if quantity > current_hold:
            return Response(
                {'error': f'Only {current_hold} on hold; cannot release more'},
                status=status.HTTP_400_BAD_REQUEST
            )
        lot.quantity_on_hold = round(current_hold - quantity, 2)
        if lot.quantity_on_hold <= 0:
            lot.quantity_on_hold = 0.0
            lot.on_hold = False
            lot.status = 'accepted'
        else:
            lot.on_hold = True
            lot.status = 'on_hold'
        lot.save(update_fields=['quantity_on_hold', 'on_hold', 'status'])
        serializer = self.get_serializer(lot)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=['post'])
    def reconcile(self, request, pk=None):
        """Admin override: set quantity_remaining to match reality. Requires staff or superuser."""
        if not getattr(request.user, 'is_authenticated', True) or (
            not getattr(request.user, 'is_staff', False) and not getattr(request.user, 'is_superuser', False)
        ):
            return Response({'error': 'Admin override requires staff or superuser.'}, status=status.HTTP_403_FORBIDDEN)
        lot = self.get_object()
        try:
            new_remaining = round(float(request.data.get('quantity_remaining')), 2)
        except (TypeError, ValueError):
            return Response({'error': 'Valid quantity_remaining is required'}, status=status.HTTP_400_BAD_REQUEST)
        if new_remaining < 0:
            return Response({'error': 'quantity_remaining cannot be negative'}, status=status.HTTP_400_BAD_REQUEST)
        reason = (request.data.get('reason') or '').strip() or 'Admin reconcile'
        quantity_before = lot.quantity_remaining
        quantity_change = new_remaining - quantity_before
        if quantity_change == 0:
            serializer = self.get_serializer(lot)
            return Response(serializer.data, status=status.HTTP_200_OK)
        lot.quantity_remaining = new_remaining
        if new_remaining > 0:
            lot.depleted_at = None
        lot.save(update_fields=['quantity_remaining', 'depleted_at'])
        try:
            from .models import LotTransactionLog
            LotTransactionLog.objects.create(
                lot=lot,
                lot_number=lot.lot_number or '',
                item_sku=lot.item.sku,
                item_name=lot.item.name,
                vendor=lot.item.vendor or '',
                transaction_type='adjustment',
                quantity_before=quantity_before,
                quantity_change=quantity_change,
                quantity_after=new_remaining,
                unit_of_measure=lot.item.unit_of_measure,
                reference_number=None,
                reference_type='admin_reconcile',
                notes=reason,
                logged_by=getattr(request.user, 'username', None) or getattr(request.user, 'email', None) or 'admin'
            )
        except Exception as e:
            import logging
            logging.getLogger(__name__).warning(f'Failed to log reconcile: {e}')
        serializer = self.get_serializer(lot)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=False, methods=['get'])
    def debug_lot(self, request):
        """Debug endpoint to check a specific lot number"""
        lot_number = request.query_params.get('lot_number', None)
        
        if not lot_number:
            return Response({'error': 'lot_number parameter is required'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            lot = Lot.objects.select_related('item').get(lot_number=lot_number)
            from django.db.models import Sum
            from .models import SalesOrderLot
            
            # Get allocations
            allocations = SalesOrderLot.objects.filter(lot=lot).select_related(
                'sales_order_item__sales_order'
            )
            
            allocation_details = []
            total_allocated = 0.0
            for alloc in allocations:
                allocation_details.append({
                    'sales_order': alloc.sales_order_item.sales_order.so_number,
                    'sales_order_status': alloc.sales_order_item.sales_order.status,
                    'quantity_allocated': alloc.quantity_allocated,
                    'sales_order_item_id': alloc.sales_order_item.id
                })
                total_allocated += alloc.quantity_allocated
            
            return Response({
                'lot_number': lot.lot_number,
                'item_sku': lot.item.sku,
                'item_name': lot.item.name,
                'item_id': lot.item.id,
                'vendor': lot.item.vendor,
                'status': lot.status,
                'quantity': lot.quantity,
                'quantity_remaining': lot.quantity_remaining,
                'allocations': allocation_details,
                'total_allocated': total_allocated,
                'calculated_remaining': max(0.0, lot.quantity_remaining - total_allocated),
                'received_date': lot.received_date,
                'expiration_date': lot.expiration_date
            })
        except Lot.DoesNotExist:
            return Response({'error': f'Lot {lot_number} not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            import traceback
            return Response({
                'error': str(e),
                'traceback': traceback.format_exc()
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=False, methods=['get'])
    def inventory_details(self, request):
        """Get inventory details grouped hierarchically: SKU -> Vendor -> Lots"""
        from django.db.models import Sum, Q
        from django.db import connection
        from django.db.utils import OperationalError
        from .models import SalesOrderItem, ProductionBatchInput, Item, CostMaster, PurchaseOrder
        
        # Check if required tables exist
        try:
            with connection.cursor() as cursor:
                # Check Item table
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_item'")
                if not cursor.fetchone():
                    return Response([])
                
                # Check Lot table
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_lot'")
                if not cursor.fetchone():
                    return Response([])
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response([])
        
        try:
            from django.utils import timezone
            from datetime import timedelta
            # inventory_table: finished_good | raw_material | indirect_material (split view); or item_type for legacy filter
            inventory_table = request.query_params.get('inventory_table', None)
            item_type_filter = request.query_params.get('item_type', None)
            
            if inventory_table:
                # Separate tables: FG, Raw, Indirect. Distributed: raw until repack closed, then FG.
                if inventory_table == 'finished_good':
                    items = Item.objects.filter(item_type__in=['finished_good', 'distributed_item'])
                elif inventory_table == 'raw_material':
                    items = Item.objects.filter(item_type__in=['raw_material', 'distributed_item'])
                elif inventory_table == 'indirect_material':
                    items = Item.objects.filter(item_type='indirect_material')
                else:
                    items = Item.objects.all()
            else:
                items = Item.objects.all()
                if item_type_filter:
                    items = items.filter(item_type=item_type_filter)
        except (OperationalError, Exception) as e:
            import traceback
            traceback.print_exc()
            return Response([])
        
        inventory_data = []
        sku_master_data = {}  # Store master SKU aggregations
        
        # Pre-calculate item-level sales allocations using SalesOrderLot
        item_sales_allocations = {}
        try:
            # Check if SalesOrderLot table exists
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_salesorderlot'")
                salesorderlot_exists = cursor.fetchone() is not None
            
            if salesorderlot_exists:
                for item_id in items.values_list('id', flat=True):
                    try:
                        # Sum allocations from SalesOrderLot for this item
                        # Include all statuses where material is allocated to sales orders
                        total_allocated = SalesOrderLot.objects.filter(
                            sales_order_item__item_id=item_id,
                            sales_order_item__sales_order__status__in=[
                                'draft', 'allocated', 'issued', 'ready_for_shipment', 'shipped'
                            ]
                        ).aggregate(
                            total=Sum('quantity_allocated')
                        )['total'] or 0.0
                        item_sales_allocations[item_id] = total_allocated
                    except Exception:
                        item_sales_allocations[item_id] = 0.0
            else:
                # Table doesn't exist, initialize all to 0
                for item_id in items.values_list('id', flat=True):
                    item_sales_allocations[item_id] = 0.0
        except Exception as e:
            import traceback
            traceback.print_exc()
            # Initialize all to 0 if there's an error
            for item_id in items.values_list('id', flat=True):
                item_sales_allocations[item_id] = 0.0
        
        # Detect if Lot table has quantity_on_hold (migration 0061); if not, use legacy on_hold logic
        use_legacy_on_hold = False
        try:
            Lot.objects.filter(quantity_remaining__gt=0).first()
        except Exception as e:
            err = str(e).lower()
            if 'quantity_on_hold' in err or 'no such column' in err:
                use_legacy_on_hold = True
        
        # Lots visible in inventory: qty > 0 OR depleted within last 24h (so can unfk quickly)
        from datetime import timedelta
        now = timezone.now()
        depleted_cutoff = now - timedelta(hours=24)
        try:
            lot_visible_q = Q(quantity_remaining__gt=0) | Q(depleted_at__gte=depleted_cutoff)
        except Exception:
            lot_visible_q = Q(quantity_remaining__gt=0)
        
        # For inventory_table split: which lot IDs are repack outputs (distributed -> FG)
        repack_output_lot_ids = set()
        if inventory_table and inventory_table in ('finished_good', 'raw_material'):
            try:
                from .models import ProductionBatchOutput
                repack_output_lot_ids = set(
                    ProductionBatchOutput.objects.filter(
                        batch__batch_type='repack', batch__status='closed'
                    ).values_list('lot_id', flat=True)
                )
            except Exception:
                pass
        
        # Group items by SKU to avoid duplicates
        items_by_sku = {}
        for item in items:
            if item.sku not in items_by_sku:
                items_by_sku[item.sku] = []
            items_by_sku[item.sku].append(item)
        
        # For each unique SKU, get all vendors
        # CRITICAL: Process every SKU, even if it has no lots or vendors
        for sku, sku_items in items_by_sku.items():
            if not sku_items:
                continue  # Skip if no items (shouldn't happen)
            
            # DEBUG: Ensure we're processing this SKU
            print(f"Processing SKU: {sku} with {len(sku_items)} items")
            # Use the first item as representative for SKU-level data
            item = sku_items[0]
            # Get vendors for this SKU from CostMaster
            cost_masters = []
            try:
                with connection.cursor() as cursor:
                    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_costmaster'")
                    if cursor.fetchone():
                        cost_masters = list(CostMaster.objects.filter(wwi_product_code=sku))
            except Exception:
                pass
            
            # Get all lots for all items with this SKU (accepted and on_hold; visible = qty>0 or depleted <24h)
            item_ids = [i.id for i in sku_items]
            item_types = {i.id: i.item_type for i in sku_items}
            try:
                if use_legacy_on_hold:
                    qs = Lot.objects.filter(
                        item_id__in=item_ids,
                        status__in=['accepted', 'on_hold']
                    ).filter(lot_visible_q).only(
                        'id', 'lot_number', 'vendor_lot_number', 'item_id', 'pack_size_id', 'quantity', 'quantity_remaining',
                        'received_date', 'expiration_date', 'status', 'on_hold', 'freight_actual', 'po_number', 'short_reason', 'created_at'
                    ).select_related('item')
                else:
                    qs = Lot.objects.filter(
                        item_id__in=item_ids,
                        status__in=['accepted', 'on_hold']
                    ).filter(lot_visible_q).select_related('item')
                item_lots = list(qs)
            except Exception:
                item_lots = []
            
            # For split table: distributed_item lots go to FG only if from repack output, to Raw only if not
            if inventory_table == 'finished_good' and repack_output_lot_ids is not None:
                item_lots = [lot for lot in item_lots if item_types.get(lot.item_id) != 'distributed_item' or lot.id in repack_output_lot_ids]
            elif inventory_table == 'raw_material' and repack_output_lot_ids is not None:
                item_lots = [lot for lot in item_lots if item_types.get(lot.item_id) != 'distributed_item' or lot.id not in repack_output_lot_ids]
            
            # Build a map of vendor -> lots by checking PO vendor
            vendor_lots_map = {}
            lots_without_vendor = []
            
            # Check if PurchaseOrder table exists
            po_table_exists = False
            try:
                with connection.cursor() as cursor:
                    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_purchaseorder'")
                    po_table_exists = cursor.fetchone() is not None
            except Exception:
                pass
            
            # Include lot (already filtered by lot_visible_q: qty>0 or depleted within 24h)
            for lot in item_lots:
                
                vendor_name = None
                if lot.po_number and po_table_exists:
                    try:
                        po = PurchaseOrder.objects.get(po_number=lot.po_number)
                        vendor_name = po.vendor_customer_name
                    except (PurchaseOrder.DoesNotExist, Exception):
                        pass
                
                if vendor_name:
                    if vendor_name not in vendor_lots_map:
                        vendor_lots_map[vendor_name] = []
                    vendor_lots_map[vendor_name].append(lot)
                else:
                    lots_without_vendor.append(lot)
            
            # Get unique vendors from all items with this SKU (this ensures all items are represented)
            # IMPORTANT: Include items with null/empty vendor as "Unknown"
            item_vendors = set()
            for sku_item in sku_items:
                vendor_name = sku_item.vendor if sku_item.vendor else "Unknown"
                item_vendors.add(vendor_name)
            
            # ALWAYS ensure we have at least one vendor entry (even if empty)
            if not item_vendors:
                item_vendors.add("Unknown")
            
            # Also get vendors from CostMaster and lots
            cost_master_vendors = set()
            for cm in cost_masters:
                if cm.vendor:
                    item_exists = Item.objects.filter(sku=sku, vendor=cm.vendor).exists()
                    if item_exists:
                        cost_master_vendors.add(cm.vendor)
            
            # Combine all vendor sources - items, cost master, and lots
            all_vendors = item_vendors.union(cost_master_vendors).union(set(vendor_lots_map.keys()))
            
            # If we have lots without vendor info, ensure "Unknown" is in the list
            if lots_without_vendor and "Unknown" not in all_vendors:
                all_vendors.add("Unknown")
            
            # CRITICAL: Ensure we always have at least one vendor entry for every SKU
            # This guarantees all items are shown, even if they have no vendor and no lots
            if not all_vendors:
                all_vendors = {"Unknown"}
            
            # Create an inventory entry for each vendor
            for vendor_name in sorted(all_vendors):
                # Get lots for this vendor
                vendor_lots = vendor_lots_map.get(vendor_name, [])
                
                # If this is the first vendor and there are lots without vendor, assign them here
                if vendor_name == sorted(all_vendors)[0] and lots_without_vendor:
                    vendor_lots.extend(lots_without_vendor)
                
                # Find the vendor-specific item - prioritize exact match
                vendor_item = Item.objects.filter(sku=sku, vendor=vendor_name).first()
                
                # If no exact match and vendor is "Unknown", use any item with this SKU
                if not vendor_item:
                    if vendor_name == "Unknown":
                        vendor_item = Item.objects.filter(sku=sku).first()
                    else:
                        # Try to find item with null/empty vendor for this SKU
                        vendor_item = Item.objects.filter(sku=sku, vendor__isnull=True).first() or \
                                     Item.objects.filter(sku=sku, vendor='').first() or \
                                     Item.objects.filter(sku=sku).first()
                
                # Always create entry - we know we have items for this SKU
                # If vendor_item is still None, use the first item from sku_items
                if not vendor_item:
                    vendor_item = sku_items[0]  # Use first item as fallback
                
                # Use the vendor-specific item if it exists, otherwise use the first item
                display_item = vendor_item if vendor_item else item
                
                # Aggregate quantities from lots
                total_quantity = sum(lot.quantity for lot in vendor_lots)
                
                # Calculate quantity_remaining for each lot, excluding sales allocations
                # Pre-calculate lot-level sales allocations
                lot_sales_allocations = {}
                try:
                    with connection.cursor() as cursor:
                        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_salesorderlot'")
                        if cursor.fetchone():
                            for lot in vendor_lots:
                                allocated = SalesOrderLot.objects.filter(
                                    lot=lot,
                                    sales_order_item__sales_order__status__in=[
                                        'draft', 'allocated', 'issued', 'ready_for_shipment', 'shipped'
                                    ]
                                ).aggregate(
                                    total=Sum('quantity_allocated')
                                )['total'] or 0.0
                                lot_sales_allocations[lot.id] = allocated
                except Exception:
                    pass
                
                # Calculate quantity_remaining for each lot, subtracting sales allocations
                quantity_remaining = sum(
                    max(0.0, lot.quantity_remaining - lot_sales_allocations.get(lot.id, 0.0))
                    for lot in vendor_lots
                )
                
                # Get on_order from the vendor-specific item
                on_order = vendor_item.on_order if vendor_item else 0.0
                
                # Calculate raw materials used in closed production batches (these should be deducted from TQ)
                consumed_in_production = 0.0
                try:
                    with connection.cursor() as cursor:
                        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_productionbatchinput'")
                        if cursor.fetchone():
                            consumed_in_production = ProductionBatchInput.objects.filter(
                                lot__in=vendor_lots,
                                batch__status='closed'  # Only count closed batches (materials already consumed)
                            ).aggregate(
                                total=Sum('quantity_used')
                            )['total'] or 0.0
                except Exception:
                    pass
                
                # TQ = lot quantities + on_order - materials consumed in closed batches
                total_quantity = total_quantity + on_order - consumed_in_production
                
                # Calculate allocated to production (sum across all vendor lots)
                # Only allocate materials when batch status is 'in_progress' (not 'scheduled' or 'closed')
                # 'scheduled' batches haven't started yet, so material isn't actually allocated
                # 'closed' batches have already consumed the material (quantity_remaining already reduced)
                allocated_to_production = 0.0
                try:
                    with connection.cursor() as cursor:
                        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_productionbatchinput'")
                        if cursor.fetchone():
                            allocated_to_production = ProductionBatchInput.objects.filter(
                                lot__in=vendor_lots,
                                batch__status='in_progress'  # Only count in_progress batches (material actively allocated)
                            ).aggregate(
                                total=Sum('quantity_used')
                            )['total'] or 0.0
                except Exception:
                    pass
                
                # Calculate on hold (quantity_on_hold is source of truth when column exists; else use status/on_hold)
                if use_legacy_on_hold:
                    on_hold = sum(lot.quantity_remaining for lot in vendor_lots if lot.status == 'on_hold' or lot.on_hold)
                else:
                    on_hold = sum(getattr(lot, 'quantity_on_hold', 0.0) for lot in vendor_lots)
                
                # Get sales allocation - sum across all items with this SKU for this vendor
                allocated_to_sales = 0.0
                if vendor_item:
                    allocated_to_sales = item_sales_allocations.get(vendor_item.id, 0.0)
                else:
                    # Sum across all items with this SKU
                    for sku_item in sku_items:
                        allocated_to_sales += item_sales_allocations.get(sku_item.id, 0.0)
                
                # Available = quantity_remaining - on_hold - allocated_to_production.
                # We do NOT reduce lot.quantity_remaining when adding to a batch; we reduce when the batch closes.
                # So "available" = what's left minus what's committed to in-progress batches.
                available = max(0.0, quantity_remaining - on_hold - allocated_to_production)
                
                # Create vendor-level inventory entry (nested under SKU)
                vendor_entry = {
                    'id': f"{sku}_{vendor_name}",  # Composite ID based on SKU+vendor
                    'item_id': display_item.id,
                    'item_sku': sku,
                    'description': display_item.name,
                    'vendor': vendor_name,
                    'pack_size': display_item.pack_size,  # Legacy field
                    'pack_size_unit': display_item.unit_of_measure,  # Legacy field
                    'pack_sizes': [ItemPackSizeSerializer(ps).data for ps in display_item.pack_sizes.filter(is_active=True)],
                    'total_quantity': total_quantity,
                    'allocated_to_sales': allocated_to_sales,
                    'allocated_to_production': allocated_to_production,
                    'on_hold': on_hold,
                    'on_order': on_order,
                    'available': available,
                    'quantity_remaining': quantity_remaining,
                    'lot_count': len(vendor_lots),
                    'item_type': display_item.item_type,
                    'level': 'vendor',  # Mark as vendor level
                }
                
                # Initialize or update master SKU aggregation
                if sku not in sku_master_data:
                    sku_master_data[sku] = {
                        'id': f"SKU_{sku}",  # Master SKU ID
                        'item_sku': sku,
                        'description': display_item.name,
                        'item_id': display_item.id,  # Use first item's ID for FPS lookup
                        'item_type': display_item.item_type,
                        'pack_size_unit': display_item.unit_of_measure,
                        'total_quantity': 0.0,
                        'allocated_to_sales': 0.0,
                        'allocated_to_production': 0.0,
                        'on_hold': 0.0,
                        'on_order': 0.0,
                        'available': 0.0,
                        'quantity_remaining': 0.0,
                        'lot_count': 0,
                        'vendor_count': 0,
                        'level': 'sku',  # Mark as SKU master level
                        'vendors': [],  # Store vendor entries
                    }
                
                # Aggregate to master SKU totals
                master = sku_master_data[sku]
                master['total_quantity'] += total_quantity
                master['allocated_to_sales'] += allocated_to_sales
                master['allocated_to_production'] += allocated_to_production
                master['on_hold'] += on_hold
                master['on_order'] += on_order
                master['available'] += available
                master['quantity_remaining'] += quantity_remaining
                master['lot_count'] += len(vendor_lots)
                master['vendor_count'] += 1
                master['vendors'].append(vendor_entry)
            
            # CRITICAL FALLBACK: If no vendor entries were created for this SKU, 
            # create a default entry to ensure the SKU appears in inventory
            # This handles edge cases where vendor matching fails
            if sku not in sku_master_data:
                # This ensures all SKUs are represented even if vendor logic fails
                first_item = sku_items[0]
                total_sales_alloc = sum(item_sales_allocations.get(i.id, 0.0) for i in sku_items)
                total_on_order = sum(getattr(i, 'on_order', 0.0) or 0.0 for i in sku_items)
                
                sku_master_data[sku] = {
                    'id': f"SKU_{sku}",
                    'item_sku': sku,
                    'description': first_item.name,
                    'item_id': first_item.id,
                    'item_type': first_item.item_type,
                    'pack_size_unit': first_item.unit_of_measure,
                    'total_quantity': total_on_order,  # Include on_order in total_quantity
                    'allocated_to_sales': total_sales_alloc,
                    'allocated_to_production': 0.0,
                    'on_hold': 0.0,
                    'on_order': total_on_order,
                    'available': 0.0,
                    'quantity_remaining': 0.0,
                    'lot_count': 0,
                    'vendor_count': 0,
                    'level': 'sku',
                    'vendors': [],
                }
        
        # Convert master SKU data to list and return
        for sku, master_data in sku_master_data.items():
            inventory_data.append(master_data)
        
        # DEBUG: Log what we're returning
        try:
            print(f"[INVENTORY DEBUG] Returning {len(inventory_data)} SKU entries")
            if inventory_data:
                print(f"[INVENTORY DEBUG] First entry: SKU={inventory_data[0].get('item_sku')}, level={inventory_data[0].get('level')}")
            else:
                print(f"[INVENTORY DEBUG] WARNING: No inventory data to return!")
                print(f"[INVENTORY DEBUG] Items in DB: {items.count()}, SKUs grouped: {len(items_by_sku)}")
        except Exception:
            pass
        
        try:
            return Response(inventory_data)
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response([])
    
    def create(self, request, *args, **kwargs):
        # Get the item to check if it's a raw material
        item_id = request.data.get('item_id')
        item = None
        if item_id:
            try:
                item = Item.objects.get(id=item_id)
            except Item.DoesNotExist:
                pass
        
        # Lot numbers should NEVER be generated on receipt/check-in
        # Lot numbers are ONLY generated when:
        # 1. Production batch closure (output lots)
        # 2. Repack batch closure (output lots)
        # For raw materials, use vendor_lot_number as the lot_number
        # For other items checked in, use vendor_lot_number if provided, otherwise leave blank (will be set on batch closure)
        lot_number = None
        
        # For raw materials, vendor_lot_number is required and used as lot_number
        if item and item.item_type == 'raw_material':
            vendor_lot_number = request.data.get('vendor_lot_number')
            # Handle None, empty string, or whitespace-only strings
            if vendor_lot_number is None:
                vendor_lot_number = ''
            elif isinstance(vendor_lot_number, str):
                vendor_lot_number = vendor_lot_number.strip()
            else:
                vendor_lot_number = str(vendor_lot_number).strip()
            
            if not vendor_lot_number:
                return Response(
                    {'error': 'Vendor lot number is required for raw materials'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            # Use vendor lot number as the lot number for raw materials
            lot_number = vendor_lot_number
        else:
            # For non-raw materials, use vendor_lot_number if provided, otherwise leave blank
            # Lot number will be generated when batch is closed
            vendor_lot_number = request.data.get('vendor_lot_number')
            if vendor_lot_number:
                if isinstance(vendor_lot_number, str):
                    vendor_lot_number = vendor_lot_number.strip()
                else:
                    vendor_lot_number = str(vendor_lot_number).strip()
                if vendor_lot_number:
                    lot_number = vendor_lot_number
        
        # Get lot_status from request data (renamed to avoid shadowing status module)
        lot_status = request.data.get('status', 'accepted')
        
        # Create the lot
        serializer_data = {
            **request.data,
            'status': lot_status,
        }
        
        # Only set lot_number if we have one (vendor_lot_number for raw materials or provided vendor_lot_number for others)
        if lot_number:
            serializer_data['lot_number'] = lot_number
        else:
            # Database requires lot_number to be NOT NULL, so generate one if not provided
            # This will be used for the lot and can be replaced on batch closure if needed
            serializer_data['lot_number'] = generate_lot_number()
        
        serializer = self.get_serializer(data=serializer_data)
        serializer.is_valid(raise_exception=True)
        lot = serializer.save()
        
        # Handle pack_size_id if provided
        pack_size_id = request.data.get('pack_size_id')
        if pack_size_id:
            try:
                pack_size = ItemPackSize.objects.get(id=pack_size_id, item=lot.item, is_active=True)
                lot.pack_size = pack_size
                lot.save()
            except ItemPackSize.DoesNotExist:
                # If pack_size_id doesn't exist or doesn't match item, try to find default
                default_pack_size = ItemPackSize.objects.filter(item=lot.item, is_default=True, is_active=True).first()
                if default_pack_size:
                    lot.pack_size = default_pack_size
                    lot.save()
        else:
            # If no pack_size_id provided, try to set default
            default_pack_size = ItemPackSize.objects.filter(item=lot.item, is_default=True, is_active=True).first()
            if default_pack_size:
                lot.pack_size = default_pack_size
                lot.save()
        
        # For raw materials, ensure vendor_lot_number is saved
        if item and item.item_type == 'raw_material':
            vendor_lot_number = request.data.get('vendor_lot_number')
            # Handle None, empty string, or whitespace-only strings
            if vendor_lot_number is None:
                vendor_lot_number = ''
            elif isinstance(vendor_lot_number, str):
                vendor_lot_number = vendor_lot_number.strip()
            else:
                vendor_lot_number = str(vendor_lot_number).strip()
            
            # Always save vendor_lot_number for raw materials (it should match lot_number)
            if vendor_lot_number:
                lot.vendor_lot_number = vendor_lot_number
                # Ensure lot_number matches vendor_lot_number for raw materials
                if lot.lot_number != vendor_lot_number:
                    lot.lot_number = vendor_lot_number
                lot.save()
        
        # Set quantity_remaining and on_hold based on lot_status
        if lot_status == 'accepted':
            lot.quantity_remaining = lot.quantity
            lot.on_hold = False
        elif lot_status == 'rejected':
            lot.quantity_remaining = 0
            lot.on_hold = False
        elif lot_status == 'on_hold':
            lot.quantity_remaining = lot.quantity  # Physical qty in house; not available until released
            lot.on_hold = True
        lot.save()
        
        # Only create inventory transaction for accepted lots
        if lot_status == 'accepted':
            transaction = InventoryTransaction.objects.create(
                transaction_type='receipt',
                lot=lot,
                quantity=lot.quantity,
            )
            
            # Get PO if exists for logging
            po = None
            if lot.po_number:
                try:
                    from .models import PurchaseOrder, PurchaseOrderItem
                    po = PurchaseOrder.objects.get(po_number=lot.po_number)
                except PurchaseOrder.DoesNotExist:
                    pass
            
            # Log the transaction
            log_lot_transaction(
                lot=lot,
                quantity_before=0.0,  # New lot, no previous quantity
                quantity_change=lot.quantity,
                transaction_type='receipt',
                reference_number=lot.po_number,
                reference_type='po_number',
                transaction_id=transaction.id,
                purchase_order_id=po.id if po else None,
                notes=f'Lot received - PO: {lot.po_number}' if lot.po_number else 'Lot received'
            )
            
            # Update on_order and PO item if PO number is provided
            if lot.po_number and po:
                try:
                    for po_item in po.items.all():
                        if po_item.item == lot.item:
                            # Update quantity received
                            po_item.quantity_received += lot.quantity
                            po_item.save()
                            
                            # Reduce on_order by the quantity received
                            item = lot.item
                            item.on_order = max(0, (item.on_order or 0) - lot.quantity)
                            item.save()
                            
                            # If quantity received is less than ordered, keep remaining on_order
                            # The outstanding balance stays on_order until fully received
                            remaining_ordered = po_item.quantity_ordered - po_item.quantity_received
                            if remaining_ordered > 0:
                                # Outstanding quantity remains on_order
                                pass
                            break
                    
                    # Check if all items are fully received
                    # Use a small tolerance for floating point comparison
                    all_received = True
                    for po_item in po.items.all():
                        # Allow small tolerance (0.01) for floating point precision
                        if po_item.quantity_received < (po_item.quantity_ordered - 0.01):
                            all_received = False
                            break
                    
                    # If all items are fully received, update PO status to 'received'
                    if all_received and po.status == 'issued':
                        po.status = 'received'
                        po.save()
                        log_purchase_order_action(po, 'completed', lot=lot, notes='All items fully received')
                    else:
                        # Partial check-in
                        log_purchase_order_action(po, 'partial_check_in', lot=lot, notes=f'Partial check-in: {lot.quantity} received')
                except PurchaseOrder.DoesNotExist:
                    pass
        
        # Log check-in to CheckInLog with all form data
        try:
            from .models import CheckInLog
            
            # Get all check-in form fields from request
            coa = request.data.get('coa', False)
            if isinstance(coa, str):
                coa = coa.lower() in ('true', '1', 'yes')
            
            prod_free_pests = request.data.get('prod_free_pests', False)
            if isinstance(prod_free_pests, str):
                prod_free_pests = prod_free_pests.lower() in ('true', '1', 'yes')
            
            carrier_free_pests = request.data.get('carrier_free_pests', False)
            if isinstance(carrier_free_pests, str):
                carrier_free_pests = carrier_free_pests.lower() in ('true', '1', 'yes')
            
            shipment_accepted = request.data.get('shipment_accepted', False)
            if isinstance(shipment_accepted, str):
                shipment_accepted = shipment_accepted.lower() in ('true', '1', 'yes')
            
            CheckInLog.objects.create(
                lot=lot,
                lot_number=lot.lot_number or '',
                item_id=lot.item.id,
                item_sku=lot.item.sku,
                item_name=lot.item.name,
                item_type=lot.item.item_type,
                item_unit_of_measure=lot.item.unit_of_measure,
                po_number=lot.po_number,
                vendor_name=lot.item.vendor if hasattr(lot.item, 'vendor') and lot.item.vendor else None,
                received_date=lot.received_date,
                vendor_lot_number=lot.vendor_lot_number,
                quantity=lot.quantity,
                quantity_unit=lot.item.unit_of_measure,  # Use item's native unit
                status=lot_status,
                short_reason=lot.short_reason,
                coa=coa,
                prod_free_pests=prod_free_pests,
                carrier_free_pests=carrier_free_pests,
                shipment_accepted=shipment_accepted,
                initials=request.data.get('initials') or '',
                carrier=request.data.get('carrier') or '',
                freight_actual=lot.freight_actual,
                notes=request.data.get('notes') or '',
                checked_in_by=request.user.username if hasattr(request, 'user') and hasattr(request.user, 'username') else 'system'
            )
        except Exception as e:
            # Log error but don't fail the check-in
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f'Failed to log check-in for lot {lot.lot_number}: {str(e)}')
            import traceback
            logger.warning(traceback.format_exc())
        
        headers = self.get_success_headers(serializer.data)
        return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)
    
    @action(detail=True, methods=['post'], url_path='reverse-check-in', url_name='reverse-check-in')
    def reverse_check_in(self, request, pk=None):
        try:
            lot = self.get_object()
        except Exception as e:
            return Response(
                {'error': f'Lot not found: {str(e)}'},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Store lot info before any operations that might trigger foreign key checks
        lot_id = lot.id
        lot_number = lot.lot_number
        try:
            item_sku = lot.item.sku if lot.item else None
        except Exception:
            item_sku = None
        quantity = lot.quantity
        
        # Validate that lot hasn't been used
        if lot.quantity_remaining < lot.quantity:
            return Response(
                {'error': 'Cannot reverse check-in: lot has been partially or fully used'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Get item_id before any operations that might trigger foreign key checks
        try:
            item_id = lot.item_id if hasattr(lot, 'item_id') else None
            po_number = lot.po_number if hasattr(lot, 'po_number') else None
        except Exception:
            item_id = None
            po_number = None
        
        # Reverse the inventory transaction (if one exists) using raw SQL
        from django.db import connection
        from django.utils import timezone
        
        try:
            with connection.cursor() as cursor:
                # Check if receipt transaction exists
                cursor.execute("""
                    SELECT id FROM erp_core_inventorytransaction 
                    WHERE lot_id = ? AND transaction_type = 'receipt'
                    LIMIT 1
                """, [lot_id])
                receipt_transaction = cursor.fetchone()
                
                if receipt_transaction:
                    # Create reverse transaction using raw SQL
                    cursor.execute("""
                        INSERT INTO erp_core_inventorytransaction 
                        (transaction_type, lot_id, quantity, notes, created_at)
                        VALUES (?, ?, ?, ?, ?)
                    """, ['adjustment', lot_id, -quantity, 'Reverse check-in', timezone.now()])
            
            # Update PO if it exists using raw SQL
            if po_number:
                try:
                    with connection.cursor() as cursor:
                        # Get PO ID
                        cursor.execute("SELECT id FROM erp_core_purchaseorder WHERE po_number = ?", [po_number])
                        po_result = cursor.fetchone()
                        if po_result:
                            po_id = po_result[0]
                            
                            # Update PO item quantity_received
                            # First, get the current quantity_received to know how much to restore
                            if item_id:
                                cursor.execute("""
                                    SELECT quantity_received 
                                    FROM erp_core_purchaseorderitem 
                                    WHERE purchase_order_id = ? AND item_id = ?
                                """, [po_id, item_id])
                                result = cursor.fetchone()
                                current_received = result[0] if result else 0.0
                                
                                # Calculate how much will actually be reversed (can't go below 0)
                                amount_to_reverse = min(quantity, current_received)
                                
                                # Update quantity_received
                                cursor.execute("""
                                    UPDATE erp_core_purchaseorderitem 
                                    SET quantity_received = MAX(0, quantity_received - ?)
                                    WHERE purchase_order_id = ? AND item_id = ?
                                """, [quantity, po_id, item_id])
                                
                                # Only restore on_order by the amount that was actually reversed
                                if amount_to_reverse > 0:
                                    cursor.execute("""
                                        UPDATE erp_core_item 
                                        SET on_order = on_order + ?
                                        WHERE id = ?
                                    """, [amount_to_reverse, item_id])
                            
                            # Update PO status back to 'issued' if it was 'received'
                            cursor.execute("""
                                UPDATE erp_core_purchaseorder 
                                SET status = 'issued'
                                WHERE id = ? AND status = 'received'
                            """, [po_id])
                except Exception:
                    pass  # PO might not exist, that's okay
        except Exception as e:
            import traceback
            traceback.print_exc()
            # Don't return error here - continue with deletion
        
        # Store lot info for response (already stored above, but create response object)
        lot_info = {
            'lot_number': lot_number,
            'item_sku': item_sku,
            'quantity': quantity
        }
        
        # Check for related objects that might prevent deletion using raw SQL only
        from django.db import connection
        from django.db.utils import OperationalError
        
        # Check if lot is used in sales orders using raw SQL
        has_sales_allocations = False
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_salesorderlot'")
                if cursor.fetchone():
                    cursor.execute("SELECT COUNT(*) FROM erp_core_salesorderlot WHERE lot_id = ?", [lot_id])
                    result = cursor.fetchone()
                    has_sales_allocations = result[0] > 0 if result else False
        except Exception:
            pass
        
        # Check if lot is used in production batches using raw SQL
        has_production_usage = False
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_productionbatchinput'")
                if cursor.fetchone():
                    cursor.execute("SELECT COUNT(*) FROM erp_core_productionbatchinput WHERE lot_id = ?", [lot_id])
                    result = cursor.fetchone()
                    has_production_usage = result[0] > 0 if result else False
        except Exception:
            pass
        
        if has_sales_allocations:
            return Response(
                {'error': 'Cannot reverse check-in: lot is allocated to sales orders. Please remove allocations first.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        if has_production_usage:
            return Response(
                {'error': 'Cannot reverse check-in: lot is used in production batches. Cannot reverse lots that have been used in production.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Delete related inventory transactions first using raw SQL to avoid ORM foreign key checks
        try:
            with connection.cursor() as cursor:
                cursor.execute("DELETE FROM erp_core_inventorytransaction WHERE lot_id = ?", [lot_id])
        except Exception:
            pass  # Table might not exist, that's okay
        
        # Delete related SalesOrderLot entries if table exists
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_salesorderlot'")
                if cursor.fetchone():
                    cursor.execute("DELETE FROM erp_core_salesorderlot WHERE lot_id = ?", [lot_id])
        except Exception:
            pass  # Table doesn't exist or error, that's okay
        
        # Delete related ProductionBatchInput entries if table exists
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_productionbatchinput'")
                if cursor.fetchone():
                    cursor.execute("DELETE FROM erp_core_productionbatchinput WHERE lot_id = ?", [lot_id])
        except Exception:
            pass  # Table doesn't exist or error, that's okay
        
        # Now delete the lot using raw SQL ONLY to avoid Django ORM foreign key checks
        # We never use lot.delete() because it will check foreign keys to tables that may not exist
        try:
            with connection.cursor() as cursor:
                # Temporarily disable foreign key checks
                cursor.execute("PRAGMA foreign_keys = OFF")
                # Delete the lot
                cursor.execute("DELETE FROM erp_core_lot WHERE id = ?", [lot_id])
                # Re-enable foreign key checks
                cursor.execute("PRAGMA foreign_keys = ON")
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response(
                {'error': f'Error deleting lot: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        
        return Response({
            'message': 'Check-in reversed successfully',
            'lot': lot_info
        }, status=status.HTTP_200_OK)


class ItemPackSizeViewSet(viewsets.ModelViewSet):
    """ViewSet for managing item pack sizes"""
    queryset = ItemPackSize.objects.select_related('item').all()
    serializer_class = ItemPackSizeSerializer
    
    def get_queryset(self):
        queryset = ItemPackSize.objects.select_related('item').all()
        
        # Filter by item_id if provided
        item_id = self.request.query_params.get('item_id', None)
        if item_id:
            queryset = queryset.filter(item_id=item_id)
        
        # Filter by is_active if provided
        is_active = self.request.query_params.get('is_active', None)
        if is_active is not None:
            is_active_bool = is_active.lower() == 'true'
            queryset = queryset.filter(is_active=is_active_bool)
        
        return queryset.order_by('item', 'pack_size', 'pack_size_unit')


class ProductionBatchViewSet(viewsets.ModelViewSet):
    queryset = ProductionBatch.objects.select_related('finished_good_item').prefetch_related(
        'inputs__lot__item', 'outputs__lot__item'
    ).all()
    serializer_class = ProductionBatchSerializer
    
    def create(self, request, *args, **kwargs):
        """Handle batch creation for both production and repack batches"""
        data = request.data.copy()
        inputs_data = data.pop('inputs', [])
        outputs_data = data.pop('outputs', [])
        indirect_materials_data = data.pop('indirect_materials', [])  # Separate indirect materials
        work_in_partials_data = data.pop('work_in_partials', [])  # Partial lots to work into this batch
        batch_type = data.get('batch_type', 'production')
        
        # Convert production_date from date string to datetime if needed
        if 'production_date' in data and isinstance(data['production_date'], str):
            from django.utils.dateparse import parse_date, parse_datetime
            # Try parsing as datetime first
            parsed = parse_datetime(data['production_date'])
            if not parsed:
                # Try parsing as date and convert to datetime
                date_obj = parse_date(data['production_date'])
                if date_obj:
                    # Create datetime at noon in local timezone (CST) to avoid timezone conversion issues
                    # This ensures the date stays the same regardless of timezone
                    # We use noon instead of midnight to avoid edge cases with DST
                    from datetime import time as dt_time
                    local_tz = timezone.get_current_timezone()
                    local_midday = local_tz.localize(datetime.combine(date_obj, dt_time(12, 0, 0)))
                    parsed = local_midday
            if parsed:
                data['production_date'] = parsed
            else:
                # If parsing fails, use current time
                data['production_date'] = timezone.now()
        
        # Auto-generate batch number if not provided or if it already exists
        if 'batch_number' not in data or not data.get('batch_number'):
            data['batch_number'] = generate_batch_number(batch_type)
        else:
            # Check if the provided batch number already exists
            existing_batch = ProductionBatch.objects.filter(batch_number=data['batch_number']).first()
            if existing_batch:
                # Generate a new unique batch number
                data['batch_number'] = generate_batch_number(batch_type)
        
        # Ensure finished_good_item_id is set (required for both production and repack)
        if 'finished_good_item_id' not in data or not data.get('finished_good_item_id'):
            if batch_type == 'production':
                return Response(
                    {'error': 'finished_good_item_id is required for production batches'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            # For repack, we'll set it from the input lots below
        
        # For repack batches, ensure we have the item and inputs
        if batch_type == 'repack':
            if not inputs_data:
                return Response(
                    {'error': 'Repack batches require at least one input lot'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Get the item from the first input lot
            try:
                first_lot = Lot.objects.get(id=inputs_data[0]['lot_id'])
                item = first_lot.item
                
                # Validate that all input lots are for the same item
                for input_data in inputs_data:
                    lot = Lot.objects.get(id=input_data['lot_id'])
                    if lot.item.id != item.id:
                        return Response(
                            {'error': 'All input lots must be for the same item in a repack batch'},
                            status=status.HTTP_400_BAD_REQUEST
                        )
                
                # Set the finished_good_item to the item (even though it's not a finished good)
                data['finished_good_item_id'] = item.id
            except Lot.DoesNotExist:
                return Response(
                    {'error': 'Invalid lot ID in inputs'},
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        # Validate inputs before creating batch
        from .models import ProductionBatchInput, ProductionBatchOutput, InventoryTransaction
        quantity_produced_from_request = round(float(data.get('quantity_produced', 0)), 2)  # Round to 2 decimal places
        total_input_quantity_in_lbs = 0.0  # Track total in lbs for validation
        total_input_quantity_native = 0.0  # Track total in item's native unit (for repack batches)
        
        # First pass: validate all inputs and calculate total
        for input_data in inputs_data:
            lot_id = input_data.get('lot_id')
            # Preserve exact values - if it's a whole number, keep it exact; otherwise round to 2 decimals
            raw_quantity = float(input_data.get('quantity_used', 0))
            # Check if it's effectively a whole number (within floating point tolerance)
            # Use tolerance of 0.01 to catch floating point errors (e.g., 615.99 -> 616)
            rounded_to_int = round(raw_quantity)
            if abs(raw_quantity - rounded_to_int) <= 0.01:
                quantity_used = rounded_to_int  # Preserve exact whole numbers
            else:
                quantity_used = round(raw_quantity, 2)  # Round to 2 decimal places
            
            if not lot_id or quantity_used <= 0:
                return Response(
                    {'error': 'Invalid input data: lot_id and quantity_used are required'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            try:
                lot = Lot.objects.get(id=lot_id)
                # Available = quantity_remaining minus what's already committed to in-progress batches (we don't reduce lot until batch closes)
                from django.db.models import Sum
                already_allocated = ProductionBatchInput.objects.filter(
                    lot=lot, batch__status='in_progress'
                ).aggregate(total=Sum('quantity_used'))['total'] or 0.0
                available = lot.quantity_remaining - already_allocated
                if available < quantity_used:
                    return Response(
                        {'error': f'Insufficient quantity in lot {lot.lot_number}. Available: {available}, Requested: {quantity_used}'},
                        status=status.HTTP_400_BAD_REQUEST
                    )
                
                # Convert quantity to lbs for validation
                quantity_used_in_lbs = quantity_used
                if lot.item.unit_of_measure == 'kg':
                    quantity_used_in_lbs = quantity_used * 2.20462  # Convert kg to lbs
                elif lot.item.unit_of_measure == 'ea':
                    # For "each" items, assume 1:1 ratio (may need adjustment based on business rules)
                    quantity_used_in_lbs = quantity_used
                
                total_input_quantity_in_lbs += quantity_used_in_lbs
                total_input_quantity_native += quantity_used  # Keep in native unit
                
            except Lot.DoesNotExist:
                return Response(
                    {'error': f'Lot with id {lot_id} not found'},
                    status=status.HTTP_404_NOT_FOUND
                )
        
        # For repack batches, quantity_produced should be in the item's native unit (not lbs)
        # For production batches, quantity_produced is in lbs
        tolerance = 0.02  # Allow for kg/lbs conversion rounding (e.g. 700.01 vs 700.00)
        if batch_type == 'repack':
            quantity_produced = round(total_input_quantity_native, 2)
            if abs(total_input_quantity_native - quantity_produced_from_request) > tolerance:
                return Response(
                    {
                        'error': f'Quantity mismatch: Total quantity used ({total_input_quantity_native:.2f} {lot.item.unit_of_measure}) must equal quantity to produce ({quantity_produced_from_request:.2f} {lot.item.unit_of_measure})'
                    },
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            # Use actual total input (rounded) as quantity_produced so inventory never drifts from conversion rounding
            total_rounded = round(total_input_quantity_in_lbs, 2)
            if abs(total_input_quantity_in_lbs - quantity_produced_from_request) > tolerance:
                return Response(
                    {
                        'error': f'Quantity mismatch: Total quantity used ({total_input_quantity_in_lbs:.2f} lbs) must equal quantity to produce ({quantity_produced_from_request:.2f} lbs)'
                    },
                    status=status.HTTP_400_BAD_REQUEST
                )
            quantity_produced = total_rounded  # Store exact sum of inputs so no drift over time
        data['quantity_produced'] = quantity_produced
        
        # Store work_in_partials in notes if provided (will be processed when closing)
        if work_in_partials_data:
            import json
            partials_json = json.dumps(work_in_partials_data)
            if data.get('notes'):
                data['notes'] = f"{data['notes']}\n[WORK_IN_PARTIALS:{partials_json}]"
            else:
                data['notes'] = f"[WORK_IN_PARTIALS:{partials_json}]"
        
        # Create the batch
        serializer = self.get_serializer(data=data)
        if not serializer.is_valid():
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f'Serializer validation failed: {serializer.errors}')
            logger.error(f'Data being validated: {data}')
            return Response(
                serializer.errors,
                status=status.HTTP_400_BAD_REQUEST
            )
        batch = serializer.save()
        
        # If batch is created as 'closed', work_in_partials will be processed in the update method
        # when the status changes to closed. For now, we'll just ensure the batch gets saved with the notes.
        
        # Create inputs (second pass: actually create them - validation already done above)
        for input_data in inputs_data:
            lot_id = input_data.get('lot_id')
            # Preserve exact values - if it's a whole number, keep it exact; otherwise round to 2 decimals
            raw_quantity = float(input_data.get('quantity_used', 0))
            # Check if it's effectively a whole number (within floating point tolerance)
            if abs(raw_quantity - round(raw_quantity)) < 0.001:
                quantity_used = round(raw_quantity)  # Preserve exact whole numbers
            else:
                quantity_used = round(raw_quantity, 2)  # Round to 2 decimal places
            
            # Get lot (already validated in first pass)
            lot = Lot.objects.get(id=lot_id)
            
            ProductionBatchInput.objects.create(
                batch=batch,
                lot=lot,
                quantity_used=quantity_used
            )
            # Do NOT reduce lot.quantity_remaining here. Lots are reduced when the batch is closed.
            # This way "Available" in inventory = quantity_remaining - allocated_to_production shows correctly
            # (e.g. 5000 received - 140 committed = 4860 available).
        
        # Validate that total input quantity equals quantity to produce (with tolerance for conversion rounding)
        # This is a second validation check after batch creation
        tolerance = 0.02  # Allow for kg/lbs conversion and floating point (e.g. 700.01 vs 700.00)
        if batch_type == 'repack':
            # For repack, we need to get the item to know the unit
            if inputs_data:
                first_lot = Lot.objects.get(id=inputs_data[0]['lot_id'])
                item_unit = first_lot.item.unit_of_measure
                if abs(total_input_quantity_native - quantity_produced) > tolerance:
                    batch.delete()
                    return Response(
                        {
                            'error': f'Quantity mismatch: Total quantity used ({total_input_quantity_native:.2f} {item_unit}) must equal quantity to produce ({quantity_produced:.2f} {item_unit})'
                        },
                        status=status.HTTP_400_BAD_REQUEST
                    )
        else:
            if abs(total_input_quantity_in_lbs - quantity_produced) > tolerance:
                batch.delete()
                return Response(
                    {
                        'error': f'Quantity mismatch: Total quantity used ({total_input_quantity_in_lbs:.2f} lbs) must equal quantity to produce ({quantity_produced:.2f} lbs)'
                    },
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        # Create outputs
        # NOTE: For both production and repack batches, output lots are created when the batch is closed, not on creation
        # This ensures internal lot numbers are only generated when batches are completed
        if batch_type == 'repack':
            # For repack, don't create output lot yet - it will be created when batch is closed
            # Just validate that we have the necessary data
            if not outputs_data:
                # Calculate total_input_quantity for validation
                total_input_quantity = 0.0
                for input_data in inputs_data:
                    lot_id = input_data.get('lot_id')
                    # Preserve exact values - if it's a whole number, keep it exact; otherwise round to 2 decimals
                    raw_quantity = float(input_data.get('quantity_used', 0))
                    # Check if it's effectively a whole number (within floating point tolerance)
                    # Use tolerance of 0.01 to catch floating point errors (e.g., 615.99 -> 616)
                    rounded_to_int = round(raw_quantity)
                    if abs(raw_quantity - rounded_to_int) <= 0.01:
                        quantity_used = rounded_to_int  # Preserve exact whole numbers
                    else:
                        quantity_used = round(raw_quantity, 2)  # Round to 2 decimal places
                    lot = Lot.objects.get(id=lot_id)
                    total_input_quantity += quantity_used  # quantity_used is already in lot's native unit
                
                # Output lot will be created when batch is closed with lot number generated at that time
            else:
                # Use provided outputs
                for output_data in outputs_data:
                    lot_id = output_data.get('lot_id')
                    quantity_produced = float(output_data.get('quantity_produced', 0))
                    
                    if not lot_id or quantity_produced <= 0:
                        batch.delete()
                        return Response(
                            {'error': 'Invalid output data: lot_id and quantity_produced are required'},
                            status=status.HTTP_400_BAD_REQUEST
                        )
                    
                    try:
                        lot = Lot.objects.get(id=lot_id)
                        ProductionBatchOutput.objects.create(
                            batch=batch,
                            lot=lot,
                            quantity_produced=quantity_produced
                        )
                        
                        # Create inventory transaction
                        transaction = InventoryTransaction.objects.create(
                            transaction_type='repack_output' if batch_type == 'repack' else 'production_output',
                            lot=lot,
                            quantity=quantity_produced,
                            notes=f'{"Repack" if batch_type == "repack" else "Production"} batch {batch.batch_number} output',
                            reference_number=batch.batch_number
                        )
                        
                        # Log the transaction for repacks
                        if batch_type == 'repack':
                            log_lot_transaction(
                                lot=lot,
                                quantity_before=0.0,  # New lot
                                quantity_change=quantity_produced,
                                transaction_type='repack_output',
                                reference_number=batch.batch_number,
                                reference_type='batch_number',
                                transaction_id=transaction.id,
                                batch_id=batch.id,
                                notes=f'Repack batch {batch.batch_number} output - Distributed item relabeled/repacked'
                            )
                    except Lot.DoesNotExist:
                        batch.delete()
                        return Response(
                            {'error': f'Lot with id {lot_id} not found'},
                            status=status.HTTP_404_NOT_FOUND
                        )
        else:
            # For production batches, use provided outputs (existing behavior)
            for output_data in outputs_data:
                lot_id = output_data.get('lot_id')
                quantity_produced = round(float(output_data.get('quantity_produced', 0)), 2)  # Round to 2 decimal places
                
                if lot_id and quantity_produced > 0:
                    try:
                        lot = Lot.objects.get(id=lot_id)
                        ProductionBatchOutput.objects.create(
                            batch=batch,
                            lot=lot,
                            quantity_produced=quantity_produced
                        )
                        
                        InventoryTransaction.objects.create(
                            transaction_type='production_output',
                            lot=lot,
                            quantity=round(quantity_produced, 2),
                            notes=f'Batch {batch.batch_number} output',
                            reference_number=batch.batch_number
                        )
                    except Lot.DoesNotExist:
                        pass  # Allow production batches to be created without outputs initially
        
        # Process indirect materials consumption
        for indirect_data in indirect_materials_data:
            lot_id = indirect_data.get('lot_id')
            quantity_used = round(float(indirect_data.get('quantity_used', 0)), 2)
            
            if not lot_id or quantity_used <= 0:
                continue  # Skip invalid entries
            
            try:
                lot = Lot.objects.get(id=lot_id)
                
                # Verify it's an indirect material
                if lot.item.item_type != 'indirect_material':
                    continue  # Skip if not an indirect material
                
                # Check available quantity
                if lot.quantity_remaining < quantity_used:
                    batch.delete()
                    return Response(
                        {'error': f'Insufficient quantity in indirect material lot {lot.lot_number}. Available: {lot.quantity_remaining}, Requested: {quantity_used}'},
                        status=status.HTTP_400_BAD_REQUEST
                    )
                
                # Create ProductionBatchInput for indirect material (reusing the same model)
                ProductionBatchInput.objects.create(
                    batch=batch,
                    lot=lot,
                    quantity_used=quantity_used
                )
                
                # Create inventory transaction
                quantity_before = lot.quantity_remaining
                transaction = InventoryTransaction.objects.create(
                    transaction_type='indirect_material_consumption',
                    lot=lot,
                    quantity=round(-quantity_used, 2),
                    notes=f'{batch.get_batch_type_display()} batch {batch.batch_number} - indirect material consumption',
                    reference_number=batch.batch_number
                )
                
                # Log the transaction
                log_lot_transaction(
                    lot=lot,
                    quantity_before=quantity_before,
                    quantity_change=-quantity_used,
                    transaction_type='indirect_material_consumption',
                    reference_number=batch.batch_number,
                    reference_type='batch_number',
                    transaction_id=transaction.id,
                    batch_id=batch.id,
                    notes=f'Indirect material consumed in {batch.get_batch_type_display()} batch {batch.batch_number}'
                )
                
                # Update lot quantity_remaining
                lot.quantity_remaining = round(lot.quantity_remaining - quantity_used, 2)
                lot.save()
                
            except Lot.DoesNotExist:
                batch.delete()
                return Response(
                    {'error': f'Indirect material lot with id {lot_id} not found'},
                    status=status.HTTP_404_NOT_FOUND
                )
        
        # Return the created batch
        serializer = self.get_serializer(batch)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    def update(self, request, *args, **kwargs):
        """Update batch and handle status changes"""
        from .models import ProductionBatchInput, ProductionBatchOutput, InventoryTransaction
        
        partial = kwargs.pop('partial', False)
        instance = self.get_object()
        old_status = instance.status
        data = request.data.copy()
        new_status = data.get('status', old_status)
        
        # Handle work_in_partials (partial lots to work into this batch)
        work_in_partials_data = data.pop('work_in_partials', [])
        
        # Handle inputs update if provided
        inputs_data = data.pop('inputs', None)
        if inputs_data is not None:
            # Round quantity_produced if provided
            if 'quantity_produced' in data:
                data['quantity_produced'] = round(float(data['quantity_produced']), 2)
            
            # Delete existing inputs; restore lot only if we had reduced (legacy: transaction exists)
            for existing_input in instance.inputs.all():
                lot = existing_input.lot
                # Only restore if there's a consumption transaction (we now only reduce at batch close)
                old_transactions = InventoryTransaction.objects.filter(
                    lot=lot,
                    reference_number=instance.batch_number,
                    transaction_type__in=['production_input', 'repack_input'],
                    quantity__lt=0
                ).order_by('-transaction_date')
                for old_txn in old_transactions:
                    if abs(abs(old_txn.quantity) - existing_input.quantity_used) < 0.01:
                        lot.quantity_remaining = round(lot.quantity_remaining + existing_input.quantity_used, 2)
                        lot.save()
                        break
                for old_txn in list(old_transactions):
                    if abs(abs(old_txn.quantity) - existing_input.quantity_used) < 0.01:
                        old_txn.delete()
                from .models import LotTransactionLog
                LotTransactionLog.objects.filter(
                    lot=lot,
                    reference_number=instance.batch_number,
                    transaction_type__in=['production_input', 'repack_input'],
                    batch_id=instance.id
                ).delete()
                existing_input.delete()
            
            # Create new inputs with rounded quantities
            for input_data in inputs_data:
                lot_id = input_data.get('lot_id')
                # Preserve exact values - if it's a whole number, keep it exact; otherwise round to 2 decimals
                raw_quantity = float(input_data.get('quantity_used', 0))
                # Check if it's effectively a whole number (within floating point tolerance)
                # Use tolerance of 0.01 to catch floating point errors (e.g., 615.99 -> 616)
                rounded_to_int = round(raw_quantity)
                if abs(raw_quantity - rounded_to_int) <= 0.01:
                    quantity_used = rounded_to_int  # Preserve exact whole numbers
                else:
                    quantity_used = round(raw_quantity, 2)  # Round to 2 decimal places
                
                if lot_id and quantity_used > 0:
                    try:
                        from django.db.models import Sum
                        lot = Lot.objects.get(id=lot_id)
                        already_allocated = ProductionBatchInput.objects.filter(
                            lot=lot, batch__status='in_progress'
                        ).aggregate(total=Sum('quantity_used'))['total'] or 0.0
                        available = lot.quantity_remaining - already_allocated
                        if available < quantity_used:
                            return Response(
                                {'error': f'Insufficient quantity in lot {lot.lot_number}. Available: {available}, Requested: {quantity_used}'},
                                status=status.HTTP_400_BAD_REQUEST
                            )
                        ProductionBatchInput.objects.create(
                            batch=instance,
                            lot=lot,
                            quantity_used=quantity_used
                        )
                        # Do NOT reduce lot here; we reduce when the batch is closed.
                    except Lot.DoesNotExist:
                        return Response(
                            {'error': f'Lot with id {lot_id} not found'},
                            status=status.HTTP_404_NOT_FOUND
                        )
        
        # Update the batch
        serializer = self.get_serializer(instance, data=data, partial=partial)
        serializer.is_valid(raise_exception=True)
        batch = serializer.save()
        
        # If status changed to 'closed', create output lots if they don't exist and log the closure
        if old_status != 'closed' and new_status == 'closed':
            # Set closed_date if not already set
            if not batch.closed_date:
                batch.closed_date = timezone.now()
                batch.save()
            
            # Extract work_in_partials from notes if they were stored there (from earlier update or creation)
            work_in_partials_from_notes = []
            if batch.notes:
                import json
                import re
                # Look for [WORK_IN_PARTIALS:...] in notes
                match = re.search(r'\[WORK_IN_PARTIALS:(.*?)\]', batch.notes)
                if match:
                    try:
                        work_in_partials_from_notes = json.loads(match.group(1))
                    except json.JSONDecodeError:
                        pass
            
            # Use work_in_partials_data from request if provided, otherwise from notes
            final_work_in_partials = work_in_partials_data if work_in_partials_data else work_in_partials_from_notes
            
            # Log production batch closure
            log_production_batch_closure(batch, notes=f'Batch {batch.batch_number} closed')
            
            # Reduce input lot quantities now that batch is closed (we don't reduce at batch create)
            batch_type_label = 'repack' if batch.batch_type == 'repack' else 'production'
            input_transaction_type = 'repack_input' if batch.batch_type == 'repack' else 'production_input'
            for batch_input in batch.inputs.all():
                lot = batch_input.lot
                qty = batch_input.quantity_used
                rounded_qty = round(qty) if abs(qty - round(qty)) <= 0.01 else round(qty, 2)
                quantity_before = lot.quantity_remaining
                if lot.quantity_remaining < rounded_qty:
                    continue  # Should not happen if create validation was correct
                lot.quantity_remaining = round(lot.quantity_remaining - rounded_qty, 2)
                lot.save()
                InventoryTransaction.objects.create(
                    transaction_type=input_transaction_type,
                    lot=lot,
                    quantity=-rounded_qty,
                    notes=f'{batch_type_label.capitalize()} batch {batch.batch_number} input (closed)',
                    reference_number=batch.batch_number
                )
                log_lot_transaction(
                    lot=lot,
                    quantity_before=quantity_before,
                    quantity_change=-rounded_qty,
                    transaction_type=input_transaction_type,
                    reference_number=batch.batch_number,
                    reference_type='batch_number',
                    transaction_id=None,
                    batch_id=batch.id,
                    notes=f'Used in {batch_type_label} batch {batch.batch_number} (closed)'
                )
            
            # Check if batch has outputs
            if batch.outputs.exists():
                # Log that outputs already exist (batch may have been closed before)
                import logging
                logger = logging.getLogger(__name__)
                existing_output = batch.outputs.first()
                logger.info(f'Batch {batch.batch_number} already has output lot {existing_output.lot.lot_number if existing_output else "unknown"}')
            else:
                # Create output lot for production batches
                if batch.batch_type == 'production':
                    # Use quantity_actual if available, otherwise quantity_produced
                    # Subtract spills from the actual quantity to get the net sellable amount
                    base_quantity = batch.quantity_actual if batch.quantity_actual and batch.quantity_actual > 0 else batch.quantity_produced
                    spills_amount = batch.spills or 0
                    # Calculate net output quantity (actual minus spills)
                    output_quantity = round(max(0, base_quantity - spills_amount), 2)
                    item = batch.finished_good_item
                    
                    # Calculate total from partials worked in (if any)
                    partial_quantities = []
                    partial_lots_to_delete = []
                    if final_work_in_partials:
                        for partial_data in final_work_in_partials:
                            partial_lot_id = partial_data.get('lot_id')
                            if partial_lot_id:
                                try:
                                    partial_lot = Lot.objects.get(id=partial_lot_id, item=item, status='accepted')
                                    if partial_lot.quantity_remaining > 0:
                                        partial_qty = partial_lot.quantity_remaining
                                        partial_quantities.append(partial_qty)
                                        partial_lots_to_delete.append(partial_lot)
                                except Lot.DoesNotExist:
                                    pass
                    
                    # Add partial quantities to output quantity
                    total_partial_qty = sum(partial_quantities)
                    combined_output_quantity = round(output_quantity + total_partial_qty, 2)
                    
                    # Generate new lot number
                    lot_number = generate_lot_number()
                    
                    # Get pack size for the lot (use default pack size for this item)
                    pack_size = ItemPackSize.objects.filter(item=item, is_default=True, is_active=True).first()
                    
                    # Create new combined lot (on hold until micro testing results)
                    new_lot = Lot.objects.create(
                        lot_number=lot_number,
                        item=item,
                        pack_size=pack_size,
                        quantity=combined_output_quantity,
                        quantity_remaining=combined_output_quantity,
                        quantity_on_hold=combined_output_quantity,
                        received_date=timezone.now(),
                        status='on_hold',
                        on_hold=True
                    )
                    
                    # Create output record
                    ProductionBatchOutput.objects.create(
                        batch=batch,
                        lot=new_lot,
                        quantity_produced=combined_output_quantity
                    )
                    
                    # Create inventory transaction for new production output
                    spills_note = f', spills: {spills_amount} lbs' if spills_amount > 0 else ''
                    InventoryTransaction.objects.create(
                        transaction_type='production_output',
                        lot=new_lot,
                        quantity=round(output_quantity, 2),  # New production quantity (after spills)
                        notes=f'Production batch {batch.batch_number} output (actual: {base_quantity} lbs, net: {output_quantity} lbs{spills_note})',
                        reference_number=batch.batch_number
                    )
                    
                    # Create transactions for partials worked in and delete partial lots
                    for partial_lot, partial_qty in zip(partial_lots_to_delete, partial_quantities):
                        # Transaction for partial being worked in
                        InventoryTransaction.objects.create(
                            transaction_type='production_output',
                            lot=new_lot,
                            quantity=round(partial_qty, 2),
                            notes=f'Production batch {batch.batch_number} - worked in partial from lot {partial_lot.lot_number}',
                            reference_number=batch.batch_number
                        )
                        
                        # Transaction to remove partial lot
                        InventoryTransaction.objects.create(
                            transaction_type='adjustment',
                            lot=partial_lot,
                            quantity=round(-partial_qty, 2),
                            notes=f'Worked into batch {batch.batch_number}',
                            reference_number=batch.batch_number
                        )
                        
                        # Delete the partial lot
                        partial_lot.delete()
                    
                    # Log the created lot for debugging
                    import logging
                    logger = logging.getLogger(__name__)
                    work_in_note = f', worked in {total_partial_qty} lbs from partials' if total_partial_qty > 0 else ''
                    spills_note = f' (actual: {base_quantity} lbs, spills: {spills_amount} lbs, net: {output_quantity} lbs)' if spills_amount > 0 else ''
                    logger.info(f'Created output lot {new_lot.lot_number} for batch {batch.batch_number}: item={item.sku}, quantity={combined_output_quantity} lbs (new: {output_quantity}{work_in_note}{spills_note}), status={new_lot.status}')
                
                # For repack batches, create output lot when batch is closed (similar to production batches)
                elif batch.batch_type == 'repack':
                    # Check if batch has outputs
                    if not batch.outputs.exists():
                        # Calculate output quantity from inputs
                        total_input_quantity = 0.0
                        for input_item in batch.inputs.all():
                            total_input_quantity += input_item.quantity_used
                        
                        output_quantity = round(total_input_quantity, 2)
                        item = batch.finished_good_item
                        
                        # Get pack_size from first input or use default
                        pack_size = None
                        first_input = batch.inputs.first()
                        if first_input and first_input.lot and first_input.lot.pack_size:
                            pack_size = first_input.lot.pack_size
                        else:
                            pack_size = ItemPackSize.objects.filter(item=item, is_default=True, is_active=True).first()
                        
                        # Generate new lot number (ONLY on batch closure)
                        lot_number = generate_lot_number()
                        
                        # Create new lot (on hold until micro testing results)
                        new_lot = Lot.objects.create(
                            lot_number=lot_number,
                            item=item,
                            pack_size=pack_size,
                            quantity=output_quantity,
                            quantity_remaining=output_quantity,
                            quantity_on_hold=output_quantity,
                            received_date=timezone.now(),
                            status='on_hold',
                            on_hold=True
                        )
                        
                        # Create output record
                        ProductionBatchOutput.objects.create(
                            batch=batch,
                            lot=new_lot,
                            quantity_produced=output_quantity
                        )
                        
                        # Create inventory transaction for output (add quantity)
                        transaction = InventoryTransaction.objects.create(
                            transaction_type='repack_output',
                            lot=new_lot,
                            quantity=output_quantity,
                            notes=f'Repack batch {batch.batch_number} output',
                            reference_number=batch.batch_number
                        )
                        
                        # Log the transaction
                        log_lot_transaction(
                            lot=new_lot,
                            quantity_before=0.0,  # New lot
                            quantity_change=output_quantity,
                            transaction_type='repack_output',
                            reference_number=batch.batch_number,
                            reference_type='batch_number',
                            transaction_id=transaction.id,
                            batch_id=batch.id,
                            notes=f'Repack batch {batch.batch_number} output - Distributed item relabeled/repacked'
                        )
        
        # Return the updated batch
        serializer = self.get_serializer(batch)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'], url_path='partials', url_name='partials')
    def get_partial_lots(self, request):
        """Get partial lots (finished good lots with quantity < pack size) for a specific finished good"""
        finished_good_item_id = request.query_params.get('finished_good_item_id')
        
        if not finished_good_item_id:
            return Response(
                {'error': 'finished_good_item_id parameter is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            finished_good = Item.objects.get(id=finished_good_item_id, item_type='finished_good')
        except Item.DoesNotExist:
            return Response(
                {'error': 'Finished good item not found'},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Get all lots for this finished good that are accepted and have remaining quantity
        lots = Lot.objects.filter(
            item=finished_good,
            status='accepted',
            quantity_remaining__gt=0
        ).select_related('item', 'pack_size').order_by('-received_date')
        
        # Filter to only include partials (quantity < pack size)
        partial_lots = []
        for lot in lots:
            if lot.pack_size:
                # Get pack size value in the same unit as lot quantity
                pack_size_value = lot.pack_size.pack_size
                pack_size_unit = lot.pack_size.pack_size_unit
                
                # Convert pack size to lbs if needed for comparison
                if pack_size_unit == 'kg':
                    pack_size_in_lbs = pack_size_value * 2.20462
                elif pack_size_unit == 'lbs':
                    pack_size_in_lbs = pack_size_value
                else:
                    # For 'ea' or other units, compare directly
                    pack_size_in_lbs = pack_size_value
                
                # Check if lot quantity is less than pack size
                if lot.quantity_remaining < pack_size_in_lbs:
                    partial_lots.append(lot)
            else:
                # No pack size defined, skip (we need pack size to determine if it's a partial)
                pass
        
        # Serialize the partial lots
        serializer = LotSerializer(partial_lots, many=True)
        return Response(serializer.data)
    
    @action(detail=True, methods=['get'], url_path='pdf', url_name='pdf')
    def generate_pdf(self, request, pk=None):
        """Generate a PDF batch ticket. Tries HTML→PDF first; falls back to current template/flowable."""
        from django.http import HttpResponse, JsonResponse

        batch = self.get_object()
        pdf_bytes, filename = None, None
        try:
            from .batch_ticket_pdf_html import generate_batch_ticket_pdf_from_html
            pdf_bytes, filename = generate_batch_ticket_pdf_from_html(batch)
        except Exception:
            pass
        if not pdf_bytes:
            try:
                from .batch_ticket_pdf import build_batch_ticket_pdf
                pdf_bytes, filename = build_batch_ticket_pdf(batch)
            except ImportError as e:
                return JsonResponse(
                    {'error': 'Batch ticket PDF module not available: ' + str(e)},
                    status=500
                )
            except Exception as e:
                import traceback
                traceback.print_exc()
                return JsonResponse(
                    {'error': 'Failed to generate batch ticket PDF: ' + str(e)},
                    status=500
                )
        if not pdf_bytes:
            return JsonResponse(
                {'error': 'Failed to generate batch ticket PDF'},
                status=500
            )
        response = HttpResponse(pdf_bytes, content_type='application/pdf')
        response['Content-Disposition'] = f'inline; filename="{filename}"'
        response['X-Content-Type-Options'] = 'nosniff'
        return response
    
    @action(detail=True, methods=['post'])
    def reverse(self, request, pk=None):
        batch = self.get_object()
        
        try:
            # Store batch number for error messages
            batch_number = batch.batch_number
            
            # BEFORE reversing, create a log entry for the unclose action
            # This preserves the audit trail - the original closure log stays, and we add a new entry
            try:
                import json
                from django.utils import timezone
                
                # Get input materials information
                input_materials = []
                input_lots = []
                for input_item in batch.inputs.all():
                    try:
                        input_materials.append({
                            'item_sku': input_item.lot.item.sku,
                            'item_name': input_item.lot.item.name,
                            'quantity_used': input_item.quantity_used
                        })
                        input_lots.append(input_item.lot.lot_number)
                    except Exception:
                        pass
                
                # Get output lot information
                output_lot_number = None
                output_quantity = None
                output = batch.outputs.first()
                if output:
                    try:
                        output_lot_number = output.lot.lot_number
                        output_quantity = output.quantity_produced
                    except Exception:
                        pass
                
                # Create log entry for unclose action BEFORE batch is deleted
                # Since ProductionLog.batch now allows NULL (SET_NULL on delete),
                # the log entry will be preserved even after batch deletion
                ProductionLog.objects.create(
                    batch=batch,  # Will be set to NULL when batch is deleted (SET_NULL)
                    batch_number=batch.batch_number,
                    batch_type=batch.batch_type,
                    finished_good_sku=batch.finished_good_item.sku,
                    finished_good_name=batch.finished_good_item.name,
                    quantity_produced=batch.quantity_produced,
                    quantity_actual=batch.quantity_actual,
                    variance=batch.variance,
                    wastes=batch.wastes,
                    spills=batch.spills,
                    production_date=batch.production_date,
                    closed_date=batch.closed_date or timezone.now(),  # Keep original closed_date
                    input_materials=json.dumps(input_materials),
                    input_lots=json.dumps(input_lots),
                    output_lot_number=output_lot_number,
                    output_quantity=output_quantity,
                    notes=f"BATCH UNCLOSED/REVERSED - {batch.notes or ''}",
                    recipe_snapshot=getattr(batch, 'recipe_snapshot', None) or None,
                    closed_by=None,  # Clear closed_by since it's being unclosed
                    logged_at=timezone.now()
                )
            except Exception as log_error:
                # Log error but don't fail the reversal
                import logging
                logger = logging.getLogger(__name__)
                logger.warning(f'Failed to create unclose log entry for batch {batch_number}: {str(log_error)}')
            
            # CRITICAL: Get all lot IDs BEFORE doing anything else
            # Get output lot IDs and lot objects before deleting anything
            output_lot_ids = []
            output_lots_to_delete = []
            try:
                for batch_output in batch.outputs.all():
                    lot_id = batch_output.lot_id
                    if lot_id:
                        output_lot_ids.append(lot_id)
                        try:
                            lot = Lot.objects.get(id=lot_id)
                            output_lots_to_delete.append(lot)
                        except Lot.DoesNotExist:
                            pass
            except Exception:
                pass
            
            # Get input lot IDs before processing (for reversing transactions)
            input_lot_ids = list(batch.inputs.values_list('lot_id', flat=True))
            
            # Restore input lot quantities only if batch was closed (we reduce lots at close, not at create)
            if batch.status == 'closed':
                for batch_input in batch.inputs.all():
                    try:
                        lot_id = batch_input.lot_id
                        if lot_id:
                            lot = Lot.objects.get(id=lot_id)
                            lot.quantity_remaining = round(lot.quantity_remaining + batch_input.quantity_used, 2)
                            lot.save()
                    except (Lot.DoesNotExist, AttributeError, ValueError):
                        pass
            
            # Delete output lots FIRST (these were created by the batch)
            # This must happen before deleting the batch to avoid foreign key issues
            from .models import InventoryTransaction
            
            # Delete all transactions for output lots first
            if output_lot_ids:
                InventoryTransaction.objects.filter(lot_id__in=output_lot_ids).delete()
            
            # Then delete the output lots themselves
            for lot in output_lots_to_delete:
                try:
                    lot.delete()
                except Exception:
                    pass
            
            # Reverse all inventory transactions for this batch
            # Find transactions by looking at lots associated with batch inputs/outputs
            
            # Find and reverse transactions for input lots using reference_number (more reliable than notes)
            # Note: output lots are already deleted, so we only need to reverse input lot transactions
            if input_lot_ids:
                # Find transactions by reference_number and transaction type
                transactions = InventoryTransaction.objects.filter(
                    lot_id__in=input_lot_ids,
                    reference_number=batch_number,
                    transaction_type__in=['production_input', 'repack_input']
                )
                
                for transaction in transactions:
                    try:
                        lot_id = transaction.lot_id
                        if not lot_id:
                            continue
                        
                        # Check if lot still exists before creating reverse transaction
                        try:
                            lot = Lot.objects.get(id=lot_id)
                            # Create reversal transaction - positive quantity to add back to inventory
                            reverse_quantity = abs(transaction.quantity)  # Make it positive
                            reverse_transaction = InventoryTransaction.objects.create(
                                transaction_type='adjustment',
                                lot=lot,
                                quantity=reverse_quantity,
                                reference_number=batch_number,
                                notes=f'UNFK: Reverse batch {batch_number} - Return input to inventory'
                            )
                            
                            # Log the reversal transaction
                            from .models import LotTransactionLog
                            log_lot_transaction(
                                lot=lot,
                                quantity_before=lot.quantity_remaining - reverse_quantity,
                                quantity_change=reverse_quantity,
                                transaction_type='adjustment',
                                reference_number=batch_number,
                                reference_type='batch_number',
                                transaction_id=reverse_transaction.id,
                                batch_id=batch.id if hasattr(batch, 'id') else None,
                                notes=f'UNFK: Reversed batch {batch_number} - Input returned to inventory'
                            )
                        except Lot.DoesNotExist:
                            # Lot doesn't exist anymore, skip
                            pass
                    except (AttributeError, ValueError, TypeError) as e:
                        # Handle any field access errors
                        import traceback
                        print(f"Error processing transaction {transaction.id}: {e}")
                        traceback.print_exc()
                        pass
            
            # Delete the batch (this will cascade delete ProductionBatchInput and ProductionBatchOutput records)
            batch.delete()
            
            return Response({'message': 'Batch ticket reversed successfully'}, status=status.HTTP_200_OK)
        except Exception as e:
            import traceback
            error_msg = str(e)
            traceback.print_exc()
            return Response(
                {'error': f'Failed to reverse batch: {error_msg}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class FormulaViewSet(viewsets.ModelViewSet):
    queryset = Formula.objects.select_related('finished_good', 'critical_control_point').prefetch_related('ingredients__item').all()
    serializer_class = FormulaSerializer
    
    def create(self, request, *args, **kwargs):
        """Create formula with ingredients"""
        data = request.data.copy()
        ingredients_data = data.pop('ingredients', [])
        
        # Create the formula
        serializer = self.get_serializer(data=data)
        serializer.is_valid(raise_exception=True)
        formula = serializer.save()
        
        # Create formula items
        from .models import FormulaItem
        for ingredient_data in ingredients_data:
            FormulaItem.objects.create(
                formula=formula,
                item_id=ingredient_data.get('item_id'),
                percentage=ingredient_data.get('percentage', 0),
                notes=ingredient_data.get('notes')
            )
        
        # Return the formula with ingredients
        serializer = self.get_serializer(formula)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    def update(self, request, *args, **kwargs):
        """Update formula and its ingredients"""
        partial = kwargs.pop('partial', False)
        instance = self.get_object()
        data = request.data.copy()
        ingredients_data = data.pop('ingredients', None)
        
        # Update the formula
        serializer = self.get_serializer(instance, data=data, partial=partial)
        serializer.is_valid(raise_exception=True)
        formula = serializer.save()
        
        # Update ingredients if provided
        if ingredients_data is not None:
            from .models import FormulaItem
            # Delete existing ingredients
            FormulaItem.objects.filter(formula=formula).delete()
            # Create new ingredients
            for ingredient_data in ingredients_data:
                FormulaItem.objects.create(
                    formula=formula,
                    item_id=ingredient_data.get('item_id'),
                    percentage=ingredient_data.get('percentage', 0),
                    notes=ingredient_data.get('notes')
                )
        
        # Return the formula with ingredients
        serializer = self.get_serializer(formula)
        return Response(serializer.data)


class CriticalControlPointViewSet(viewsets.ModelViewSet):
    queryset = CriticalControlPoint.objects.all()
    serializer_class = CriticalControlPointSerializer


class PurchaseOrderViewSet(viewsets.ModelViewSet):
    queryset = PurchaseOrder.objects.prefetch_related('items__item').all()
    serializer_class = PurchaseOrderSerializer
    
    def get_queryset(self):
        queryset = PurchaseOrder.objects.prefetch_related('items__item').all()
        
        # Filter by status if provided
        status = self.request.query_params.get('status', None)
        if status:
            queryset = queryset.filter(status=status)
        
        return queryset
    
    def list(self, request, *args, **kwargs):
        try:
            return super().list(request, *args, **kwargs)
        except Exception as e:
            import traceback
            print(f"Error in PurchaseOrderViewSet.list: {e}")
            traceback.print_exc()
            raise
    
    def create(self, request, *args, **kwargs):
        # Make a mutable copy of request.data
        data = request.data.copy()
        items_data = data.pop('items', [])
        
        # Handle vendor_id -> vendor_customer_name mapping
        if 'vendor_id' in data:
            vendor_id = data.pop('vendor_id')
            try:
                vendor = Vendor.objects.get(id=vendor_id)
                data['vendor_customer_name'] = vendor.name
                data['vendor_customer_id'] = str(vendor.id)
            except Vendor.DoesNotExist:
                return Response(
                    {'error': f'Vendor with id {vendor_id} not found'},
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        # Generate PO number if not provided (format: 2yy000)
        if not data.get('po_number'):
            data['po_number'] = generate_po_number()
        
        # Ensure po_type is set
        if 'po_type' not in data:
            data['po_type'] = 'vendor'
        
        # Set required_date from expected_delivery_date if not provided
        if not data.get('required_date') and data.get('expected_delivery_date'):
            data['required_date'] = data['expected_delivery_date']
        
        # Set expected_delivery_date to match required_date if not provided
        if not data.get('expected_delivery_date') and data.get('required_date'):
            data['expected_delivery_date'] = data['required_date']
        
        # Create the purchase order
        try:
            serializer = self.get_serializer(data=data)
            serializer.is_valid(raise_exception=True)
            purchase_order = serializer.save()
        except Exception as e:
            import traceback
            print(f"Error creating PurchaseOrder: {e}")
            traceback.print_exc()
            return Response(
                {'error': f'Failed to create purchase order: {str(e)}', 'details': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        
        # Create purchase order items
        for item_data in items_data:
            try:
                item_id = item_data.get('item_id')
                if not item_id:
                    raise ValueError(f'item_id is required for item: {item_data}')
                
                # Map unit_cost to unit_price (database has unit_price, not unit_cost)
                unit_price = item_data.get('unit_cost', item_data.get('unit_price', 0))
                quantity_ordered = item_data.get('quantity', item_data.get('quantity_ordered', 0))
                
                if quantity_ordered <= 0:
                    raise ValueError(f'quantity must be greater than 0, got: {quantity_ordered}')
                
                PurchaseOrderItem.objects.create(
                    purchase_order=purchase_order,
                    item_id=item_id,
                    quantity_ordered=quantity_ordered,
                    unit_price=unit_price,  # Use unit_price as that's what exists in DB
                    notes=item_data.get('notes', ''),
                )
            except Exception as e:
                import traceback
                print(f"Error creating PurchaseOrderItem: {e}")
                print(f"Item data: {item_data}")
                traceback.print_exc()
                # Delete the purchase order if item creation fails
                purchase_order.delete()
                return Response(
                    {'error': f'Failed to create purchase order item: {str(e)}', 'item_data': item_data},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
        
        # Calculate totals if method exists
        if hasattr(purchase_order, 'calculate_totals'):
            purchase_order.calculate_totals()
            purchase_order.save()
        
        # Log purchase order creation
        log_purchase_order_action(purchase_order, 'created', notes='Purchase order created')
        
        # Return the created purchase order with items
        response_serializer = self.get_serializer(purchase_order)
        headers = self.get_success_headers(response_serializer.data)
        return Response(response_serializer.data, status=status.HTTP_201_CREATED, headers=headers)
    
    @action(detail=True, methods=['post'])
    def issue(self, request, pk=None):
        """Issue a purchase order - changes status to 'issued' and updates inventory"""
        purchase_order = self.get_object()
        
        if purchase_order.status != 'draft':
            return Response(
                {'error': f'Purchase order must be in draft status to issue. Current status: {purchase_order.status}'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Update status to issued
        purchase_order.status = 'issued'
        purchase_order.save()
        
        # Log status change
        log_purchase_order_action(purchase_order, 'updated', notes='Purchase order issued')
        
        # Increment on_order for each item
        for po_item in purchase_order.items.all():
            if po_item.item:
                item = po_item.item
                item.on_order = (item.on_order or 0) + po_item.quantity_ordered
                item.save()
        
        # Generate PDF and send email
        try:
            pdf_content = generate_purchase_order_pdf(purchase_order)
            send_purchase_order_email(purchase_order, pdf_content)
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Failed to send purchase order email: {str(e)}")
            # Don't fail the request if email fails
        
        # Return updated purchase order
        serializer = self.get_serializer(purchase_order)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=['post'])
    def receive(self, request, pk=None):
        """Mark purchase order as received"""
        purchase_order = self.get_object()
        
        if purchase_order.status != 'issued':
            return Response(
                {'error': f'Purchase order must be in issued status to receive. Current status: {purchase_order.status}'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        purchase_order.status = 'received'
        purchase_order.received_date = timezone.now()
        purchase_order.save()
        
        # Log status change
        log_purchase_order_action(purchase_order, 'updated', notes='Purchase order marked as received')
        
        # Create AP entry when PO is received
        create_ap_entry_from_po(purchase_order)
        
        serializer = self.get_serializer(purchase_order)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=['post'])
    def revise(self, request, pk=None):
        """Create a revision of a purchase order"""
        original_po = self.get_object()
        
        # Create a copy of the PO
        new_po_data = {
            'po_number': original_po.po_number,  # Same PO number
            'po_type': original_po.po_type,
            'vendor_customer_name': original_po.vendor_customer_name,
            'vendor_customer_id': original_po.vendor_customer_id,
            'status': 'draft',  # New revision starts as draft
            'revision_number': (original_po.revision_number or 0) + 1,
            'original_po': original_po,
            'order_number': original_po.order_number,
            'expected_delivery_date': original_po.expected_delivery_date,
            'required_date': original_po.required_date,
            'shipping_terms': original_po.shipping_terms,
            'shipping_method': original_po.shipping_method,
            'ship_to_name': original_po.ship_to_name,
            'ship_to_address': original_po.ship_to_address,
            'ship_to_city': original_po.ship_to_city,
            'ship_to_state': original_po.ship_to_state,
            'ship_to_zip': original_po.ship_to_zip,
            'ship_to_country': original_po.ship_to_country,
            'vendor_address': original_po.vendor_address,
            'vendor_city': original_po.vendor_city,
            'vendor_state': original_po.vendor_state,
            'vendor_zip': original_po.vendor_zip,
            'vendor_country': original_po.vendor_country,
            'subtotal': original_po.subtotal,
            'discount': original_po.discount,
            'shipping_cost': original_po.shipping_cost,
            'total': original_po.total,
            'coa_sds_email': original_po.coa_sds_email,
            'notes': original_po.notes,
        }
        
        new_po = PurchaseOrder.objects.create(**new_po_data)
        
        # Copy items
        for original_item in original_po.items.all():
            PurchaseOrderItem.objects.create(
                purchase_order=new_po,
                item=original_item.item,
                quantity_ordered=original_item.quantity_ordered,
                unit_price=original_item.unit_price,
                notes=original_item.notes,
            )
        
        # If original PO was issued, reverse its inventory impact and mark as superseded
        if original_po.status == 'issued':
            # Reverse on_order for each item
            for po_item in original_po.items.all():
                if po_item.item:
                    item = po_item.item
                    item.on_order = max(0, (item.on_order or 0) - po_item.quantity_ordered)
                    item.save()
            
            original_po.status = 'superseded'
            original_po.save()
        
        serializer = self.get_serializer(new_po)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    def update(self, request, *args, **kwargs):
        """Update purchase order and log the change"""
        partial = kwargs.pop('partial', False)
        instance = self.get_object()
        old_status = instance.status
        
        # Call parent update method
        response = super().update(request, *args, partial=partial, **kwargs)
        
        # Reload instance to get updated values
        instance.refresh_from_db()
        # Recalculate subtotal and total from items (and discount/shipping_cost)
        if hasattr(instance, 'calculate_totals'):
            instance.calculate_totals()
            instance.refresh_from_db()
        
        # Log the update
        notes = []
        if old_status != instance.status:
            notes.append(f'Status changed from {old_status} to {instance.status}')
        
        # Check if items were modified
        if 'items' in request.data:
            notes.append('Items updated')
        
        # Check if other fields were modified
        modified_fields = []
        for field in ['vendor_customer_name', 'order_date', 'expected_delivery_date', 'required_date', 
                     'carrier', 'tracking_number', 'shipping_cost', 'total', 'notes']:
            if field in request.data:
                modified_fields.append(field)
        
        if modified_fields:
            notes.append(f'Fields updated: {", ".join(modified_fields)}')
        
        log_notes = '; '.join(notes) if notes else 'Purchase order updated'
        log_purchase_order_action(instance, 'updated', notes=log_notes)
        
        return response
    
    @action(detail=True, methods=['post'])
    def cancel(self, request, pk=None):
        """Cancel a purchase order"""
        purchase_order = self.get_object()
        
        if purchase_order.status == 'completed':
            return Response(
                {'error': 'Cannot cancel a completed purchase order'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # If PO was issued, reverse inventory impact
        if purchase_order.status == 'issued':
            for po_item in purchase_order.items.all():
                if po_item.item:
                    item = po_item.item
                    item.on_order = max(0, (item.on_order or 0) - po_item.quantity_ordered)
                    item.save()
        
        purchase_order.status = 'cancelled'
        purchase_order.save()
        
        # Log cancellation
        log_purchase_order_action(purchase_order, 'cancelled', notes='Purchase order cancelled')
        
        serializer = self.get_serializer(purchase_order)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=['get'], url_path='pdf', url_name='pdf')
    def generate_pdf(self, request, pk=None):
        """Generate a PDF purchase order. Prefer Jinja2 HTML→PDF; fallback to template fill then ReportLab."""
        import logging
        from pathlib import Path
        from django.http import HttpResponse
        from .pdf_generator import generate_purchase_order_pdf
        from .po_pdf_template import get_po_template_pdf_path, fill_po_template_pdf
        from .po_pdf_html import generate_po_pdf_from_html

        logger = logging.getLogger(__name__)
        purchase_order = self.get_object()

        try:
            pdf_content = generate_po_pdf_from_html(purchase_order)
            if pdf_content:
                logger.info("PO PDF: using HTML path for PO %s", getattr(purchase_order, "po_number", pk))
            if not pdf_content:
                project_root = Path(__file__).resolve().parent.parent.parent
                template_path = get_po_template_pdf_path(project_root)
                if template_path:
                    pdf_content = fill_po_template_pdf(template_path, purchase_order)
                    if pdf_content:
                        logger.info("PO PDF: using template-fill path for PO %s", getattr(purchase_order, "po_number", pk))
            if not pdf_content:
                pdf_content = generate_purchase_order_pdf(purchase_order)
                logger.info("PO PDF: using ReportLab fallback for PO %s", getattr(purchase_order, "po_number", pk))

            response = HttpResponse(pdf_content, content_type='application/pdf')
            response['Content-Disposition'] = f'inline; filename="Purchase_Order_{purchase_order.po_number}.pdf"'
            return response
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Failed to generate purchase order PDF: {str(e)}")
            from rest_framework.response import Response
            from rest_framework import status as http_status
            return Response(
                {'error': f'Failed to generate PDF: {str(e)}'},
                status=http_status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class SalesOrderViewSet(viewsets.ModelViewSet):
    queryset = SalesOrder.objects.prefetch_related('items__item').all()
    serializer_class = SalesOrderSerializer
    
    def create(self, request, *args, **kwargs):
        # Make a mutable copy of request.data
        data = request.data.copy()
        items_data = data.pop('items', [])
        
        # Generate sales order number if not provided (format: 3yy0000)
        if not data.get('so_number'):
            data['so_number'] = generate_sales_order_number()
        
        # Handle customer - if customer or customer_id is provided, set customer FK
        customer_id = data.get('customer') or data.get('customer_id')
        if customer_id:
            try:
                customer = Customer.objects.get(id=customer_id)
                data['customer'] = customer.id
                # Auto-populate customer_name if not provided
                if not data.get('customer_name'):
                    data['customer_name'] = customer.name
                if not data.get('customer_reference_number') and data.get('customer_id'):
                    data['customer_reference_number'] = data.get('customer_id')
            except Customer.DoesNotExist:
                pass
        
        # Handle ship_to_location - if ship_to_location is provided, validate it belongs to the customer
        ship_to_location_id = data.get('ship_to_location')
        if ship_to_location_id:
            try:
                from .models import ShipToLocation
                ship_to_location = ShipToLocation.objects.get(id=ship_to_location_id)
                # Validate that ship-to location belongs to the selected customer
                if customer_id and ship_to_location.customer.id != customer_id:
                    return Response(
                        {'error': 'Ship-to location does not belong to the selected customer'},
                        status=status.HTTP_400_BAD_REQUEST
                    )
                # Auto-populate address fields from ship-to location if not provided
                if not data.get('customer_address'):
                    data['customer_address'] = ship_to_location.address
                if not data.get('customer_city'):
                    data['customer_city'] = ship_to_location.city
                if not data.get('customer_state'):
                    data['customer_state'] = ship_to_location.state or ''
                if not data.get('customer_zip'):
                    data['customer_zip'] = ship_to_location.zip_code
                if not data.get('customer_country'):
                    data['customer_country'] = ship_to_location.country
                if not data.get('customer_phone'):
                    data['customer_phone'] = ship_to_location.phone or ''
            except ShipToLocation.DoesNotExist:
                return Response(
                    {'error': 'Ship-to location not found'},
                    status=status.HTTP_404_NOT_FOUND
                )
        
        # Set default values for required numeric fields that exist in DB but not in model
        # These fields have NOT NULL constraints in the database
        numeric_defaults = {
            'subtotal': float(data.get('subtotal', 0.0)),
            'freight': float(data.get('freight', 0.0)),
            'misc': float(data.get('misc', 0.0)),
            'prepaid': float(data.get('prepaid', 0.0)),
            'discount': float(data.get('discount', 0.0)),
            'grand_total': float(data.get('grand_total', 0.0)),
        }
        
        # Validate the data first (but don't save yet)
        serializer = self.get_serializer(data=data)
        serializer.is_valid(raise_exception=True)
        validated_data = serializer.validated_data
        
        # Extract values safely
        customer_obj = validated_data.get('customer')
        customer_id = customer_obj.id if customer_obj else None
        
        ship_to_obj = validated_data.get('ship_to_location')
        ship_to_id = ship_to_obj.id if ship_to_obj else None
        
        # Create the sales order using raw SQL to include fields not in the model
        from django.db import connection
        from django.utils import timezone
        
        now = timezone.now()
        
        # Prepare all values in correct order
        values = [
            validated_data.get('so_number'),
            customer_id,
            ship_to_id,
            validated_data.get('customer_name', ''),
            validated_data.get('customer_legacy_id'),
            validated_data.get('customer_reference_number'),
            validated_data.get('customer_address'),
            validated_data.get('customer_city'),
            validated_data.get('customer_state'),
            validated_data.get('customer_zip'),
            validated_data.get('customer_country'),
            validated_data.get('customer_phone'),
            now,
            validated_data.get('expected_ship_date'),
            validated_data.get('actual_ship_date'),
            validated_data.get('status', 'draft'),
            validated_data.get('notes'),
            now,
            now,
            numeric_defaults['subtotal'],
            numeric_defaults['freight'],
            numeric_defaults['misc'],
            numeric_defaults['prepaid'],
            numeric_defaults['discount'],
            numeric_defaults['grand_total'],
        ]
        
        try:
            with connection.cursor() as cursor:
                # Build the INSERT statement with all required fields
                # Django's cursor.execute() uses %s for parameterized queries (converts to ? for SQLite)
                cursor.execute("""
                    INSERT INTO erp_core_salesorder (
                        so_number, customer_id, ship_to_location_id, customer_name,
                        customer_legacy_id, customer_reference_number,
                        customer_address, customer_city, customer_state, customer_zip,
                        customer_country, customer_phone,
                        order_date, expected_ship_date, actual_ship_date,
                        status, notes, created_at, updated_at,
                        subtotal, freight, misc, prepaid, discount, grand_total
                    ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, values)
                sales_order_id = cursor.lastrowid
            
            # Get the created sales order object
            sales_order = SalesOrder.objects.get(id=sales_order_id)
        except Exception as e:
            import logging
            import traceback
            logger = logging.getLogger(__name__)
            error_msg = str(e)
            error_trace = traceback.format_exc()
            logger.error(f'Failed to create sales order: {error_msg}')
            logger.error(f'Traceback: {error_trace}')
            logger.error(f'Values count: {len(values)}, Values: {values}')
            # Return a helpful error response
            return Response(
                {
                    'error': 'Failed to create sales order',
                    'detail': error_msg,
                    'debug_info': {
                        'values_count': len(values),
                        'so_number': validated_data.get('so_number'),
                        'customer_id': customer_id,
                        'ship_to_id': ship_to_id,
                    }
                },
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        
        # Create sales order items with lot allocations
        for item_data in items_data:
            allocated_lots_data = item_data.pop('allocated_lots', [])
            
            # Extract and map fields correctly
            # SalesOrderItem only has: sales_order, item, quantity_ordered, quantity_allocated, quantity_shipped, unit_price, notes
            item_id = item_data.get('item_id') or item_data.get('item')
            if not item_id:
                return Response(
                    {'error': 'item_id is required for each sales order item'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Create the sales order item with only valid fields
            so_item = SalesOrderItem.objects.create(
                sales_order=sales_order,
                item_id=item_id,
                quantity_ordered=item_data.get('quantity_ordered', 0),
                unit_price=item_data.get('unit_price'),
                notes=item_data.get('notes')
            )
            
            # Create lot allocations (if table exists and allocations provided)
            if allocated_lots_data:
                try:
                    # Check if SalesOrderLot table exists
                    with connection.cursor() as cursor:
                        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_salesorderlot'")
                        if cursor.fetchone():
                            for lot_data in allocated_lots_data:
                                SalesOrderLot.objects.create(
                                    sales_order_item=so_item,
                                    lot_id=lot_data.get('lot_id'),
                                    quantity_allocated=lot_data.get('quantity_allocated', 0)
                                )
                                
                                # Update quantity_allocated on the SO item
                                so_item.quantity_allocated += lot_data.get('quantity_allocated', 0)
                            so_item.save()
                except Exception as e:
                    # If SalesOrderLot table doesn't exist, just skip lot allocations
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.warning(f'SalesOrderLot table does not exist, skipping lot allocations: {str(e)}')
        
        # Refresh the sales order from database to get all related items
        try:
            sales_order.refresh_from_db()
            
            # Return the created sales order with items
            serializer = self.get_serializer(sales_order)
            response_data = serializer.data
            return Response(response_data, status=status.HTTP_201_CREATED)
        except Exception as e:
            import logging
            import traceback
            logger = logging.getLogger(__name__)
            logger.error(f'Failed to serialize sales order response: {str(e)}')
            logger.error(f'Traceback: {traceback.format_exc()}')
            # Return basic response even if serialization fails
            return Response(
                {
                    'id': sales_order.id,
                    'so_number': sales_order.so_number,
                    'status': sales_order.status,
                    'message': 'Sales order created successfully'
                },
                status=status.HTTP_201_CREATED
            )
    
    @action(detail=False, methods=['post'], url_path='parse-customer-po')
    def parse_customer_po(self, request):
        """
        Accept an uploaded customer PO document (PDF or text) and return structured data
        for auto-filling the Create Sales Order form.
        """
        from .customer_po_parser import parse_customer_po as do_parse
        file_obj = request.FILES.get('file') or request.data.get('file')
        if not file_obj:
            return Response(
                {'error': 'No file provided. Upload a PDF or text file (field name: file).'},
                status=status.HTTP_400_BAD_REQUEST
            )
        try:
            file_bytes = file_obj.read()
            filename = getattr(file_obj, 'name', '') or 'document.pdf'
            result = do_parse(file_bytes, filename=filename)
            return Response(result)
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.exception('parse_customer_po failed')
            return Response(
                {'error': f'Failed to parse document: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=True, methods=['get', 'post'], url_path='customer-po')
    def customer_po(self, request, pk=None):
        """GET: serve the uploaded customer PO PDF. POST: upload and attach a customer PO PDF."""
        from django.http import FileResponse, Http404
        from io import BytesIO
        sales_order = self.get_object()
        if request.method == 'GET':
            if not sales_order.customer_po_pdf:
                raise Http404('No customer PO document attached.')
            try:
                with sales_order.customer_po_pdf.open('rb') as fh:
                    buf = BytesIO(fh.read())
                buf.seek(0)
                return FileResponse(
                    buf,
                    as_attachment=False,
                    filename=(sales_order.customer_po_pdf.name or 'customer_po.pdf').split('/')[-1],
                    content_type='application/pdf',
                )
            except Exception as e:
                import logging
                logger = logging.getLogger(__name__)
                logger.exception('Failed to serve customer PO PDF')
                return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        else:
            file_obj = request.FILES.get('file') or request.data.get('file')
            if not file_obj:
                return Response(
                    {'error': 'No file provided. Send multipart/form-data with field "file".'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            if not getattr(file_obj, 'name', '').lower().endswith('.pdf'):
                return Response(
                    {'error': 'Only PDF files are accepted for customer PO.'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            try:
                sales_order.customer_po_pdf = file_obj
                sales_order.save(update_fields=['customer_po_pdf'])
                return Response({'ok': True, 'message': 'Customer PO document saved.'})
            except Exception as e:
                import logging
                logger = logging.getLogger(__name__)
                logger.exception('Failed to save customer PO PDF')
                return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['post'])
    def allocate(self, request, pk=None):
        """Allocate lots to sales order items. Creates distributed item lots if raw materials are checked in."""
        from django.db import transaction
        
        sales_order = self.get_object()
        items_data = request.data.get('items', [])
        
        with transaction.atomic():
            for item_data in items_data:
                item_id = item_data.get('item_id')
                is_distributed = item_data.get('is_distributed', False)
                allocations = item_data.get('allocations', [])
                raw_materials = item_data.get('raw_materials', [])
                
                try:
                    so_item = SalesOrderItem.objects.get(sales_order=sales_order, item_id=item_id)
                except SalesOrderItem.DoesNotExist:
                    return Response(
                        {'error': f'Sales order item for item {item_id} not found'},
                        status=status.HTTP_404_NOT_FOUND
                    )
                
                # Delete existing allocations for this item
                SalesOrderLot.objects.filter(sales_order_item=so_item).delete()
                so_item.quantity_allocated = 0.0
                
                if is_distributed and raw_materials:
                    # For distributed items, create new lot from raw materials
                    # Only if raw materials are checked into inventory (have lots with status='accepted')
                    distributed_item = so_item.item
                    total_quantity = 0.0
                    raw_material_lots = []
                    
                    for rm_data in raw_materials:
                        lot_id = rm_data.get('lot_id')
                        quantity = float(rm_data.get('quantity', 0))
                        
                        try:
                            lot = Lot.objects.get(id=lot_id, status='accepted')
                            if lot.quantity_remaining < quantity:
                                return Response(
                                    {'error': f'Insufficient quantity in lot {lot.lot_number}. Available: {lot.quantity_remaining}, Requested: {quantity}'},
                                    status=status.HTTP_400_BAD_REQUEST
                                )
                            raw_material_lots.append((lot, quantity))
                            total_quantity += quantity
                        except Lot.DoesNotExist:
                            return Response(
                                {'error': f'Lot {lot_id} not found or not accepted'},
                                status=status.HTTP_404_NOT_FOUND
                            )
                    
                    # Create new lot for distributed item
                    new_lot_number = generate_lot_number()
                    new_lot = Lot.objects.create(
                        lot_number=new_lot_number,
                        item=distributed_item,
                        quantity=total_quantity,
                        quantity_remaining=total_quantity,
                        received_date=timezone.now(),
                        status='accepted'
                    )
                    
                    # Reduce raw material lot quantities and create transactions
                    for raw_lot, qty in raw_material_lots:
                        quantity_before = raw_lot.quantity_remaining
                        raw_lot.quantity_remaining -= qty
                        raw_lot.save()
                        
                        transaction = InventoryTransaction.objects.create(
                            transaction_type='production',
                            lot=raw_lot,
                            quantity=-qty,
                            reference_number=sales_order.so_number,
                            notes=f'Allocated to distributed item lot {new_lot_number}'
                        )
                        
                        # Log the transaction
                        log_lot_transaction(
                            lot=raw_lot,
                            quantity_before=quantity_before,
                            quantity_change=-qty,
                            transaction_type='production_input',
                            reference_number=sales_order.so_number,
                            reference_type='so_number',
                            transaction_id=transaction.id,
                            sales_order_id=sales_order.id,
                            notes=f'Used for distributed item lot {new_lot_number} in sales order {sales_order.so_number}'
                        )
                        
                        # Log depletion if lot reaches zero or below
                        log_lot_depletion(
                            lot=raw_lot,
                            quantity_before=quantity_before,
                            quantity_used=qty,
                            depletion_method='production',
                            reference_number=sales_order.so_number,
                            reference_type='so_number',
                            sales_order_id=sales_order.id,
                            transaction_id=transaction.id,
                            notes=f'Used for distributed item lot {new_lot_number} in sales order {sales_order.so_number}'
                        )
                    
                    # Create transaction for new lot
                    transaction = InventoryTransaction.objects.create(
                        transaction_type='production',
                        lot=new_lot,
                        quantity=total_quantity,
                        reference_number=sales_order.so_number,
                        notes=f'Created from raw materials for sales order {sales_order.so_number}'
                    )
                    
                    # Log the transaction
                    log_lot_transaction(
                        lot=new_lot,
                        quantity_before=0.0,  # New lot
                        quantity_change=total_quantity,
                        transaction_type='production_output',
                        reference_number=sales_order.so_number,
                        reference_type='so_number',
                        transaction_id=transaction.id,
                        sales_order_id=sales_order.id,
                        notes=f'Created from raw materials for sales order {sales_order.so_number}'
                    )
                    
                    # Allocate the new lot to sales order
                    SalesOrderLot.objects.create(
                        sales_order_item=so_item,
                        lot=new_lot,
                        quantity_allocated=total_quantity
                    )
                    so_item.quantity_allocated = total_quantity
                    
                else:
                    # Regular items - allocate from existing lots
                    for allocation in allocations:
                        lot_id = allocation.get('lot_id')
                        quantity = float(allocation.get('quantity', 0))
                        
                        try:
                            lot = Lot.objects.get(id=lot_id, item_id=item_id, status='accepted')
                            if lot.quantity_remaining < quantity:
                                return Response(
                                    {'error': f'Insufficient quantity in lot {lot.lot_number}. Available: {lot.quantity_remaining}, Requested: {quantity}'},
                                    status=status.HTTP_400_BAD_REQUEST
                                )
                            
                            SalesOrderLot.objects.create(
                                sales_order_item=so_item,
                                lot=lot,
                                quantity_allocated=quantity
                            )
                            so_item.quantity_allocated += quantity
                        except Lot.DoesNotExist:
                            return Response(
                                {'error': f'Lot {lot_id} not found for item {item_id}'},
                                status=status.HTTP_404_NOT_FOUND
                            )
                
                so_item.save()
            
            # Refresh sales order to get updated items
            sales_order.refresh_from_db()
            
            # Check if all items are fully allocated
            all_fully_allocated = all(
                item.quantity_allocated >= item.quantity_ordered 
                for item in sales_order.items.all()
            )
            
            # Update sales order status based on allocation state
            if all_fully_allocated:
                sales_order.status = 'ready_for_shipment'
            else:
                # If any allocations exist, set to 'allocated'
                # But preserve 'issued' status if the order was issued and is partially allocated
                total_allocated = sum(item.quantity_allocated for item in sales_order.items.all())
                if total_allocated > 0:
                    # If order was 'issued' and now has partial allocations, keep it as 'issued'
                    # Otherwise set to 'allocated'
                    if sales_order.status == 'issued':
                        # Keep as 'issued' for partial allocation, user can continue allocating
                        pass
                    else:
                        sales_order.status = 'allocated'
                else:
                    # No allocations - only set to 'draft' if it wasn't 'issued'
                    if sales_order.status != 'issued':
                        sales_order.status = 'draft'
            sales_order.save()
        
        serializer = self.get_serializer(sales_order)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=['post'])
    def ship(self, request, pk=None):
        """Ship the sales order (supports partial shipments), update tracking, and create invoice."""
        from django.db import transaction
        from datetime import datetime, timedelta
        import re
        from .models import Shipment, ShipmentItem
        
        sales_order = self.get_object()
        ship_date_str = request.data.get('ship_date')
        invoice_date_str = request.data.get('invoice_date', ship_date_str)
        tracking_number = request.data.get('tracking_number', '').strip()
        carrier = (request.data.get('carrier') or '').strip()
        items_to_ship = request.data.get('items', [])  # List of {item_id, quantity} for partial shipments
        
        # Validate order is issued or ready_for_shipment and has allocations
        if sales_order.status not in ('issued', 'ready_for_shipment'):
            return Response(
                {'error': f'Sales order must be issued or ready for shipment to checkout. Current status: {sales_order.status}'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Check if order has any allocations
        total_allocated = sum(item.quantity_allocated for item in sales_order.items.all())
        if total_allocated == 0:
            return Response(
                {'error': 'Sales order must have material allocated before checkout'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        if not ship_date_str:
            return Response(
                {'error': 'ship_date is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Tracking number is optional - can be added later in Finance
        # if not tracking_number:
        #     return Response(
        #         {'error': 'tracking_number is required'},
        #         status=status.HTTP_400_BAD_REQUEST
        #     )
        
        try:
            ship_date = datetime.strptime(ship_date_str, '%Y-%m-%d').date()
            invoice_date = datetime.strptime(invoice_date_str, '%Y-%m-%d').date() if invoice_date_str else ship_date
        except ValueError:
            return Response(
                {'error': 'Invalid date format. Use YYYY-MM-DD'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # If items_to_ship is provided, use it for partial shipment
        # Otherwise, ship all allocated quantities
        use_partial = bool(items_to_ship)
        
        try:
            with transaction.atomic():
                # Create shipment record (tracking number, dimensions, pieces can be set at checkout)
                ship_dt = timezone.make_aware(datetime.combine(ship_date, datetime.min.time()))
                expected_dt = None
                if request.data.get('expected_ship_date'):
                    try:
                        from django.utils.dateparse import parse_datetime
                        expected_dt = parse_datetime(request.data.get('expected_ship_date'))
                        if expected_dt and timezone.is_naive(expected_dt):
                            expected_dt = timezone.make_aware(expected_dt)
                    except Exception:
                        pass
                if expected_dt is None and sales_order.expected_ship_date:
                    expected_dt = sales_order.expected_ship_date
                shipment = Shipment.objects.create(
                    sales_order=sales_order,
                    expected_ship_date=expected_dt,
                    ship_date=ship_dt,
                    tracking_number=tracking_number or '',
                    notes=request.data.get('notes', ''),
                    dimensions=(request.data.get('dimensions') or '').strip() or None,
                    pieces=request.data.get('pieces') if request.data.get('pieces') is not None else None,
                )
                
                # Process items to ship - normalize item_id to int for dict lookup (JSON may send string)
                items_shipped_map = {}
                if use_partial:
                    for item_data in items_to_ship:
                        raw_id = item_data.get('item_id') or item_data.get('sales_order_item_id')
                        item_id = int(raw_id) if raw_id is not None else None
                        quantity_to_ship = float(item_data.get('quantity', 0))
                        if item_id is not None and quantity_to_ship > 0:
                            items_shipped_map[item_id] = quantity_to_ship
                else:
                    for so_item in sales_order.items.all():
                        if so_item.quantity_allocated > 0:
                            items_shipped_map[so_item.id] = so_item.quantity_allocated
                
                # Reduce lot quantities and create inventory transactions
                for so_item in sales_order.items.all():
                    quantity_to_ship = items_shipped_map.get(so_item.id, 0)
                    if quantity_to_ship <= 0:
                        continue
                    
                    # Validate quantity doesn't exceed allocated
                    if quantity_to_ship > so_item.quantity_allocated:
                        return Response(
                            {'error': f'Cannot ship {quantity_to_ship} of {so_item.item.name}. Only {so_item.quantity_allocated} is allocated.'},
                            status=status.HTTP_400_BAD_REQUEST
                        )
                    
                    # Ship from allocated lots proportionally or FIFO
                    remaining_to_ship = quantity_to_ship
                    allocations = SalesOrderLot.objects.filter(sales_order_item=so_item).order_by('created_at')
                    
                    for allocation in allocations:
                        if remaining_to_ship <= 0:
                            break
                        
                        lot = allocation.lot
                        # Ship from this allocation (up to the allocated amount)
                        quantity_from_allocation = min(remaining_to_ship, allocation.quantity_allocated)
                        
                        if lot.quantity_remaining < quantity_from_allocation:
                            return Response(
                                {'error': f'Insufficient quantity in lot {lot.lot_number}. Available: {lot.quantity_remaining}, Required: {quantity_from_allocation}'},
                                status=status.HTTP_400_BAD_REQUEST
                            )
                        
                        quantity_before = lot.quantity_remaining
                        
                        inv_txn = InventoryTransaction.objects.create(
                            transaction_type='adjustment',
                            lot=lot,
                            quantity=-quantity_from_allocation,
                            reference_number=sales_order.so_number,
                            notes=f'Shipped for sales order {sales_order.so_number} - Shipment {shipment.id}'
                        )
                        
                        # Log the transaction
                        log_lot_transaction(
                            lot=lot,
                            quantity_before=quantity_before,
                            quantity_change=-quantity_from_allocation,
                            transaction_type='sale',
                            reference_number=sales_order.so_number,
                            reference_type='so_number',
                            transaction_id=inv_txn.id,
                            sales_order_id=sales_order.id,
                            notes=f'Shipped for sales order {sales_order.so_number} - Shipment {shipment.id}'
                        )
                        
                        lot.quantity_remaining -= quantity_from_allocation
                        lot.save()
                        
                        # Reduce the allocation by the shipped quantity
                        allocation.quantity_allocated -= quantity_from_allocation
                        if allocation.quantity_allocated <= 0:
                            allocation.delete()
                        else:
                            allocation.save()
                        
                        # Log depletion if lot reaches zero or below
                        log_lot_depletion(
                            lot=lot,
                            quantity_before=quantity_before,
                            quantity_used=quantity_from_allocation,
                            depletion_method='sales',
                            reference_number=sales_order.so_number,
                            reference_type='so_number',
                            sales_order_id=sales_order.id,
                            transaction_id=inv_txn.id,
                            notes=f'Shipped for sales order {sales_order.so_number} - Shipment {shipment.id}'
                        )
                        
                        remaining_to_ship -= quantity_from_allocation
                    
                    # Update quantity_shipped on sales order item
                    so_item.quantity_shipped += quantity_to_ship
                    so_item.quantity_allocated -= quantity_to_ship  # Reduce allocated by shipped amount
                    so_item.save()
                    
                    # Create shipment item record
                    ShipmentItem.objects.create(
                        shipment=shipment,
                        sales_order_item=so_item,
                        quantity_shipped=quantity_to_ship
                    )
                
                # Update sales order status, tracking, and carrier
                sales_order.actual_ship_date = timezone.make_aware(datetime.combine(ship_date, datetime.min.time()))
                if not sales_order.tracking_number:
                    sales_order.tracking_number = tracking_number
                if carrier:
                    sales_order.carrier = carrier
                
                # Check if order is fully shipped
                all_fully_shipped = all(
                    item.quantity_shipped >= item.quantity_ordered 
                    for item in sales_order.items.all()
                )
                
                if all_fully_shipped:
                    sales_order.status = 'completed'
                else:
                    # Still has outstanding balance - keep as issued or ready_for_shipment
                    # Check if there are still allocations
                    total_remaining_allocated = sum(item.quantity_allocated for item in sales_order.items.all())
                    if total_remaining_allocated > 0:
                        sales_order.status = 'ready_for_shipment'
                    else:
                        sales_order.status = 'issued'  # Can allocate more if needed
                
                sales_order.save()
                
                # Create invoice for this shipment
                invoice_number = generate_invoice_number()
                
                # Calculate due date from payment terms
                due_date = invoice_date
                if sales_order.customer and sales_order.customer.payment_terms:
                    payment_terms = sales_order.customer.payment_terms
                    # Parse payment terms (e.g., "Net 30" -> 30 days)
                    match = re.search(r'(\d+)', payment_terms)
                    if match:
                        days = int(match.group(1))
                        due_date = invoice_date + timedelta(days=days)
                
                # Calculate totals from shipped quantities in this shipment
                subtotal = sum(
                    item.sales_order_item.unit_price * item.quantity_shipped 
                    for item in shipment.items.all() 
                    if item.sales_order_item.unit_price
                )
                # Get freight, discount, tax from sales order if available, otherwise 0
                freight = getattr(sales_order, 'freight', 0.0) or 0.0
                discount = getattr(sales_order, 'discount', 0.0) or 0.0
                tax = 0.0  # Calculate tax if needed
                grand_total = subtotal + freight + tax - discount
                
                # Check what columns actually exist in Invoice table and which are NOT NULL
                from django.db import connection
                available_columns = set()
                not_null_columns = set()
                try:
                    with connection.cursor() as cursor:
                        cursor.execute("PRAGMA table_info(erp_core_invoice)")
                        for row in cursor.fetchall():
                            col_name = row[1]
                            is_not_null = row[3]  # 1 if NOT NULL, 0 if nullable
                            available_columns.add(col_name)
                            if is_not_null and col_name != 'id':  # id is auto-generated
                                not_null_columns.add(col_name)
                except Exception:
                    # If we can't check, use minimal set
                    available_columns = {'id', 'invoice_number', 'invoice_date', 'created_at', 'updated_at'}
                    not_null_columns = {'invoice_number', 'invoice_date'}
                
                # If sales_order_id column doesn't exist, use raw SQL to avoid Django ORM trying to set it
                if 'sales_order_id' not in available_columns:
                    # Use raw SQL to insert invoice without sales_order_id
                    # Build column list and values dynamically based on what exists
                    now = timezone.now()
                    columns = ['invoice_number']
                    values = [invoice_number]
                    placeholders = ['?']
                    
                    # Required fields that must be included if they exist
                    # invoice_type is required (NOT NULL) - always include if column exists
                    if 'invoice_type' in available_columns:
                        columns.append('invoice_type')
                        values.append('customer')  # Default to customer invoice
                        placeholders.append('?')
                    # customer_vendor_name is required (NOT NULL) - always include if column exists
                    if 'customer_vendor_name' in available_columns:
                        columns.append('customer_vendor_name')
                        customer_name = sales_order.customer_name
                        if not customer_name and sales_order.customer:
                            customer_name = sales_order.customer.name
                        if not customer_name:
                            customer_name = 'Unknown Customer'
                        values.append(customer_name)
                        placeholders.append('?')
                    # customer_vendor_id is optional
                    if 'customer_vendor_id' in available_columns:
                        columns.append('customer_vendor_id')
                        customer_id = sales_order.customer_legacy_id
                        if not customer_id and sales_order.customer:
                            customer_id = str(sales_order.customer.id)
                        values.append(customer_id)
                        placeholders.append('?')
                    
                    # Optional fields
                    if 'invoice_date' in available_columns:
                        columns.append('invoice_date')
                        values.append(invoice_date)
                        placeholders.append('?')
                    if 'due_date' in available_columns:
                        columns.append('due_date')
                        values.append(due_date)
                        placeholders.append('?')
                    if 'status' in available_columns:
                        columns.append('status')
                        values.append('draft')
                        placeholders.append('?')
                    if 'subtotal' in available_columns:
                        columns.append('subtotal')
                        values.append(subtotal)
                        placeholders.append('?')
                    if 'freight' in available_columns:
                        columns.append('freight')
                        values.append(freight)
                        placeholders.append('?')
                    if 'tax' in available_columns:
                        columns.append('tax')
                        values.append(tax)
                        placeholders.append('?')
                    if 'tax_amount' in available_columns:
                        columns.append('tax_amount')
                        values.append(tax)
                        placeholders.append('?')
                    if 'discount' in available_columns:
                        columns.append('discount')
                        values.append(discount)
                        placeholders.append('?')
                    if 'grand_total' in available_columns:
                        columns.append('grand_total')
                        values.append(grand_total)
                        placeholders.append('?')
                    if 'total_amount' in available_columns:
                        columns.append('total_amount')
                        values.append(grand_total)
                        placeholders.append('?')
                    if 'paid_amount' in available_columns:
                        columns.append('paid_amount')
                        values.append(0.0)
                        placeholders.append('?')
                    if 'notes' in available_columns:
                        columns.append('notes')
                        # Escape % characters to avoid issues with Django's query logging
                        notes_text = f'Auto-generated from sales order {sales_order.so_number} - Shipment {shipment.id}'
                        # Replace % with %% to escape it (SQLite doesn't need this, but Django's logging does)
                        notes_text = notes_text.replace('%', '%%')
                        values.append(notes_text)
                        placeholders.append('?')
                    if 'created_at' in available_columns:
                        columns.append('created_at')
                        values.append(now)
                        placeholders.append('?')
                    if 'updated_at' in available_columns:
                        columns.append('updated_at')
                        values.append(now)
                        placeholders.append('?')
                    
                    # Ensure all NOT NULL columns are included
                    for col in not_null_columns:
                        if col not in columns and col != 'id':  # Skip id (auto-generated)
                            # Add default values for required columns we haven't handled
                            if col == 'invoice_type' and 'invoice_type' in available_columns:
                                columns.append('invoice_type')
                                values.append('customer')
                                placeholders.append('?')
                            elif col == 'customer_vendor_name' and 'customer_vendor_name' in available_columns:
                                columns.append('customer_vendor_name')
                                customer_name = sales_order.customer_name or (sales_order.customer.name if sales_order.customer else 'Unknown Customer')
                                values.append(customer_name)
                                placeholders.append('?')
                            elif col == 'invoice_date' and 'invoice_date' in available_columns:
                                columns.append('invoice_date')
                                values.append(invoice_date)
                                placeholders.append('?')
                            elif col == 'status' and 'status' in available_columns:
                                columns.append('status')
                                values.append('draft')
                                placeholders.append('?')
                            elif col == 'subtotal' and 'subtotal' in available_columns:
                                columns.append('subtotal')
                                values.append(subtotal)
                                placeholders.append('?')
                            elif col == 'tax_amount' and 'tax_amount' in available_columns:
                                columns.append('tax_amount')
                                values.append(tax)
                                placeholders.append('?')
                            elif col == 'total_amount' and 'total_amount' in available_columns:
                                columns.append('total_amount')
                                values.append(grand_total)
                                placeholders.append('?')
                            elif col == 'paid_amount' and 'paid_amount' in available_columns:
                                columns.append('paid_amount')
                                values.append(0.0)
                                placeholders.append('?')
                    
                    # Ensure we have the same number of placeholders as values
                    if len(placeholders) != len(values):
                        raise ValueError(f"Placeholder count ({len(placeholders)}) doesn't match value count ({len(values)}). Columns: {columns}")
                    
                    # Build SQL with proper parameterization
                    # Use string concatenation to build SQL to avoid f-string % formatting issues
                    columns_str = ', '.join(columns)
                    placeholders_str = ', '.join(placeholders)
                    sql = "INSERT INTO erp_core_invoice (" + columns_str + ") VALUES (" + placeholders_str + ")"
                    
                    # Execute using Django's cursor with proper parameterization
                    # The issue is Django's debug SQL formatter tries to do sql % params for logging
                    # This fails when the SQL string itself contains % characters in the notes field
                    # Solution: Disable query logging temporarily and use execute with proper parameterization
                    from django.conf import settings
                    import logging
                    logger = logging.getLogger(__name__)
                    
                    invoice_id = None
                    
                    # Use Django's connection so we stay in the same transaction (avoids "database is locked")
                    with connection.cursor() as raw_cursor:
                        raw_cursor.execute(sql, tuple(values))
                        invoice_id = raw_cursor.lastrowid
                    
                    if not invoice_id:
                        raise ValueError("Failed to create invoice - no ID returned")
                    
                    # After creating with raw SQL, update sales_order_id if column exists
                    with connection.cursor() as check_cursor:
                        check_cursor.execute("PRAGMA table_info(erp_core_invoice)")
                        columns = [row[1] for row in check_cursor.fetchall()]
                        if 'sales_order_id' in columns:
                            # Update the invoice with sales_order_id
                            check_cursor.execute(
                                "UPDATE erp_core_invoice SET sales_order_id = ? WHERE id = ?",
                                [sales_order.id, invoice_id]
                            )
                    
                    invoice = Invoice.objects.get(id=invoice_id)
                    
                    # Create AR entry when invoice is created
                    create_ar_entry_from_invoice(invoice)
                else:
                    # Use ORM - pass fields that exist on the Invoice model and any required DB columns (e.g. invoice_type from migrations)
                    invoice_data = {
                        'invoice_number': invoice_number,
                        'sales_order': sales_order,
                        'invoice_date': invoice_date,
                        'due_date': due_date,
                        'status': 'draft',
                        'subtotal': subtotal,
                        'freight': freight,
                        'tax': tax,
                        'discount': discount,
                        'grand_total': grand_total,
                        'notes': f'Auto-generated from sales order {sales_order.so_number} - Shipment {shipment.id}',
                    }
                    invoice_data['invoice_type'] = 'customer'
                    customer_name = getattr(sales_order, 'customer_name', None) or (sales_order.customer.name if sales_order.customer else None) or 'Unknown Customer'
                    invoice_data['customer_vendor_name'] = customer_name or 'Unknown Customer'
                    invoice_data['tax_amount'] = tax
                    invoice_data['total_amount'] = grand_total
                    invoice_data['paid_amount'] = 0.0
                    try:
                        invoice = Invoice.objects.create(**invoice_data)
                    except Exception as e:
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.error(f"Error creating invoice with ORM: {e}")
                        logger.error(f"Invoice data keys: {list(invoice_data.keys())}")
                        raise
                    
                    # Create AR entry when invoice is created
                    create_ar_entry_from_invoice(invoice)
                    
                    # Send invoice email if status is 'sent' (auto-send when created from shipment)
                    if invoice.status == 'sent':
                        try:
                            pdf_content = generate_invoice_pdf(invoice)
                            send_invoice_email(invoice, pdf_content)
                        except Exception as e:
                            import logging
                            logger = logging.getLogger(__name__)
                            logger.error(f"Failed to send invoice email: {str(e)}")
                            # Don't fail the request if email fails
                
                # Create invoice items from shipped quantities in this shipment
                for shipment_item in shipment.items.all():
                    if shipment_item.quantity_shipped > 0:
                        so_item = shipment_item.sales_order_item
                        line_total = (so_item.unit_price or 0.0) * shipment_item.quantity_shipped
                        InvoiceItem.objects.create(
                            invoice=invoice,
                            item=so_item.item,
                            sales_order_item=so_item,
                            description=so_item.item.name,
                            quantity=shipment_item.quantity_shipped,
                            unit_price=so_item.unit_price or 0.0,
                            total=line_total,
                            notes=''
                        )
            
            invoice_serializer = InvoiceSerializer(invoice)
            
            serializer = self.get_serializer(sales_order)
            return Response({
                'sales_order': serializer.data,
                'invoice': invoice_serializer.data,
                'shipment': {
                    'id': shipment.id,
                    'ship_date': shipment.ship_date.isoformat(),
                    'tracking_number': shipment.tracking_number
                }
            }, status=status.HTTP_200_OK)
        except Exception as e:
            import traceback
            return Response({
                'error': str(e),
                'detail': traceback.format_exc()
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=True, methods=['post'])
    def issue(self, request, pk=None):
        """Issue a sales order - changes status from 'draft' to 'issued'"""
        sales_order = self.get_object()
        
        if sales_order.status != 'draft':
            return Response(
                {'error': f'Sales order must be in draft status to issue. Current status: {sales_order.status}'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Update status to issued
        sales_order.status = 'issued'
        sales_order.save()
        
        # Generate PDF and send confirmation email
        try:
            pdf_content = generate_sales_order_pdf(sales_order)
            send_sales_order_confirmation_email(sales_order, pdf_content)
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Failed to send sales order confirmation email: {str(e)}")
            # Don't fail the request if email fails
        
        # Return updated sales order
        serializer = self.get_serializer(sales_order)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=['post'])
    def cancel(self, request, pk=None):
        """Cancel sales order and clean up allocations. Deletes distributed item lots permanently."""
        from django.db import transaction
        
        sales_order = self.get_object()
        
        with transaction.atomic():
            # Find all allocations
            for so_item in sales_order.items.all():
                allocations = SalesOrderLot.objects.filter(sales_order_item=so_item)
                
                for allocation in allocations:
                    lot = allocation.lot
                    
                    # Check if this is a distributed item lot (created for this sales order)
                    # We can identify this by checking if the lot was created recently and
                    # has a transaction referencing this sales order
                    is_distributed_lot = InventoryTransaction.objects.filter(
                        lot=lot,
                        reference_number=sales_order.so_number,
                        notes__icontains='distributed'
                    ).exists()
                    
                    if is_distributed_lot:
                        # This is a distributed item lot - delete it and reverse raw materials
                        # Find the raw material transactions
                        raw_material_transactions = InventoryTransaction.objects.filter(
                            reference_number=sales_order.so_number,
                            quantity__lt=0,  # Negative quantities are raw materials
                            notes__icontains=lot.lot_number
                        )
                        
                        # Reverse raw material quantities
                        for trans in raw_material_transactions:
                            raw_lot = trans.lot
                            raw_lot.quantity_remaining += abs(trans.quantity)
                            raw_lot.save()
                            
                            # Create reverse transaction
                            InventoryTransaction.objects.create(
                                transaction_type='adjustment',
                                lot=raw_lot,
                                quantity=abs(trans.quantity),
                                reference_number=sales_order.so_number,
                                notes=f'Reversed allocation from cancelled order {sales_order.so_number}'
                            )
                        
                        # Delete the distributed lot
                        lot.delete()
                
                # Delete all allocations
                allocations.delete()
                so_item.quantity_allocated = 0.0
                so_item.save()
            
            # Cancel any invoices
            Invoice.objects.filter(sales_order=sales_order).update(status='cancelled')
            
            # Update sales order status
            sales_order.status = 'cancelled'
            sales_order.save()
        
        serializer = self.get_serializer(sales_order)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    def destroy(self, request, *args, **kwargs):
        """Delete sales order and clean up all related data"""
        from django.db import transaction, connection
        
        sales_order = self.get_object()
        
        # Prevent deletion of shipped or cancelled orders
        if sales_order.status in ['shipped', 'cancelled']:
            return Response(
                {'error': f'Cannot delete a {sales_order.status} sales order'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            with transaction.atomic():
                # Clean up allocations and reverse inventory if needed
                for so_item in sales_order.items.all():
                    allocations = SalesOrderLot.objects.filter(sales_order_item=so_item)
                    
                    for allocation in allocations:
                        lot = allocation.lot
                        
                        # Check if this is a distributed item lot (created for this sales order)
                        is_distributed_lot = InventoryTransaction.objects.filter(
                            lot=lot,
                            reference_number=sales_order.so_number,
                            notes__icontains='distributed'
                        ).exists()
                        
                        if is_distributed_lot:
                            # This is a distributed item lot - delete it and reverse raw materials
                            raw_material_transactions = InventoryTransaction.objects.filter(
                                reference_number=sales_order.so_number,
                                quantity__lt=0,  # Negative quantities are raw materials
                                notes__icontains=lot.lot_number
                            )
                            
                            # Reverse raw material quantities
                            for trans in raw_material_transactions:
                                raw_lot = trans.lot
                                raw_lot.quantity_remaining += abs(trans.quantity)
                                raw_lot.save()
                                
                                # Create reverse transaction
                                InventoryTransaction.objects.create(
                                    transaction_type='adjustment',
                                    lot=raw_lot,
                                    quantity=abs(trans.quantity),
                                    reference_number=sales_order.so_number,
                                    notes=f'Reversed allocation from deleted order {sales_order.so_number}'
                                )
                            
                            # Delete the distributed lot
                            lot.delete()
                    
                    # Delete all allocations
                    allocations.delete()
                
                # Get sales order item IDs before any deletion
                so_item_ids = list(sales_order.items.values_list('id', flat=True))
                so_id = sales_order.id
                
                # Delete InvoiceItems first (they reference SalesOrderItems) using raw SQL
                try:
                    with connection.cursor() as cursor:
                        # Check if invoiceitem table exists and has sales_order_item_id column
                        cursor.execute("""
                            SELECT name FROM sqlite_master 
                            WHERE type='table' AND name='erp_core_invoiceitem'
                        """)
                        if cursor.fetchone() and so_item_ids:
                            cursor.execute("PRAGMA table_info(erp_core_invoiceitem)")
                            columns = [row[1] for row in cursor.fetchall()]
                            if 'sales_order_item_id' in columns:
                                # Delete invoice items that reference these sales order items
                                placeholders = ','.join(['%s'] * len(so_item_ids))
                                cursor.execute(f"""
                                    DELETE FROM erp_core_invoiceitem 
                                    WHERE sales_order_item_id IN ({placeholders})
                                """, so_item_ids)
                except Exception as e:
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.warning(f'Could not delete invoice items: {str(e)}')
                
                # Delete any invoices using raw SQL to avoid schema issues
                try:
                    with connection.cursor() as cursor:
                        # Check if invoice table exists and has sales_order_id column
                        cursor.execute("""
                            SELECT name FROM sqlite_master 
                            WHERE type='table' AND name='erp_core_invoice'
                        """)
                        if cursor.fetchone():
                            cursor.execute("PRAGMA table_info(erp_core_invoice)")
                            columns = [row[1] for row in cursor.fetchall()]
                            if 'sales_order_id' in columns:
                                # Delete invoices for this sales order
                                cursor.execute("""
                                    DELETE FROM erp_core_invoice 
                                    WHERE sales_order_id = %s
                                """, [so_id])
                except Exception as e:
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.warning(f'Could not delete invoices: {str(e)}')
                
                # Delete SalesOrderItems using raw SQL to avoid ORM checking for related InvoiceItems
                if so_item_ids:
                    with connection.cursor() as cursor:
                        placeholders = ','.join(['%s'] * len(so_item_ids))
                        cursor.execute(f"""
                            DELETE FROM erp_core_salesorderitem 
                            WHERE id IN ({placeholders})
                        """, so_item_ids)
                
                # Finally, delete the sales order itself using raw SQL
                with connection.cursor() as cursor:
                    cursor.execute("DELETE FROM erp_core_salesorder WHERE id = %s", [so_id])
            
            return Response(status=status.HTTP_204_NO_CONTENT)
            
        except Exception as e:
            import logging
            import traceback
            logger = logging.getLogger(__name__)
            logger.error(f'Failed to delete sales order {sales_order.so_number}: {str(e)}')
            logger.error(f'Traceback: {traceback.format_exc()}')
            return Response(
                {'error': f'Failed to delete sales order: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    def _get_packing_list_template_path(self):
        """Return path to Packing List template. Prefer PDF in project root (WWI ERP), then Sensitive, then backend_django; then docx."""
        project_root = Path(__file__).resolve().parent.parent.parent  # WWI ERP
        backend = Path(__file__).resolve().parent.parent  # backend_django
        pdf_names = ['Packing list template.pdf', 'Packing list template.PDF']
        docx_names = ['Packing list template.docx']
        for name in pdf_names:
            if (project_root / name).exists():
                return project_root / name, 'pdf'
        for name in pdf_names:
            if (project_root / 'Sensitive' / name).exists():
                return project_root / 'Sensitive' / name, 'pdf'
        for name in pdf_names:
            if (backend / name).exists():
                return backend / name, 'pdf'
        for name in docx_names:
            if (project_root / 'Sensitive' / name).exists():
                return project_root / 'Sensitive' / name, 'docx'
            if (project_root / name).exists():
                return project_root / name, 'docx'
        return None, None

    @action(detail=True, methods=['get'], url_path='packing-list', url_name='packing-list')
    def packing_list(self, request, pk=None):
        """Generate packing list PDF. Same flow as invoice: try HTML→PDF first, then fallback to template fill."""
        from django.http import HttpResponse

        # Load with relations needed for ship to / bill to and shipments (dimensions, pieces)
        sales_order = SalesOrder.objects.select_related(
            'ship_to_location', 'customer'
        ).prefetch_related('items__item', 'shipments').get(pk=self.get_object().pk)

        # Same flow as invoice PDF: try HTML→PDF first (no external template required)
        try:
            from .packing_list_pdf_html import generate_packing_list_pdf_from_html
            pdf_content = generate_packing_list_pdf_from_html(sales_order)
            if pdf_content:
                response = HttpResponse(pdf_content, content_type='application/pdf')
                response['Content-Disposition'] = f'inline; filename="Packing_List_{sales_order.so_number or pk}.pdf"'
                response['X-Content-Type-Options'] = 'nosniff'
                return response
        except Exception as e:
            import traceback
            logger.error('Packing list HTML PDF failed: %s', str(e))
            logger.error(traceback.format_exc())

        # Fallback: template-based (PDF or docx)
        from pathlib import Path
        import os
        import tempfile

        template_path, template_type = self._get_packing_list_template_path()
        if not template_path or not template_path.exists():
            return Response(
                {'error': 'Packing list template not found. Add "Packing list template.pdf" to the WWI ERP folder, or ensure packing_list_pdf_html is available.'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

        # Build data used by both PDF and docx
        ship_to_lines = []
        if getattr(sales_order, 'ship_to_location', None) and sales_order.ship_to_location:
            loc = sales_order.ship_to_location
            if loc.location_name:
                ship_to_lines.append(loc.location_name)
            if loc.contact_name:
                ship_to_lines.append(loc.contact_name)
            if loc.address:
                ship_to_lines.append(loc.address)
            csz = [x for x in [loc.city, loc.state, loc.zip_code] if x]
            if csz:
                ship_to_lines.append(', '.join(csz))
            if loc.country:
                ship_to_lines.append(loc.country)
        else:
            if sales_order.customer_name:
                ship_to_lines.append(sales_order.customer_name)
            if sales_order.customer_address:
                ship_to_lines.append(sales_order.customer_address)
            csz = [x for x in [sales_order.customer_city, sales_order.customer_state, sales_order.customer_zip] if x]
            if csz:
                ship_to_lines.append(', '.join(csz))
            if sales_order.customer_country:
                ship_to_lines.append(sales_order.customer_country)
        ship_to_text = '\n'.join(ship_to_lines) if ship_to_lines else (sales_order.customer_name or '')

        bill_to_lines = []
        if getattr(sales_order, 'customer', None) and sales_order.customer:
            c = sales_order.customer
            if c.name:
                bill_to_lines.append(c.name)
            if c.address:
                bill_to_lines.append(c.address)
            csz = [x for x in [c.city, c.state, c.zip_code] if x]
            if csz:
                bill_to_lines.append(', '.join(csz))
            if c.country:
                bill_to_lines.append(c.country)
        else:
            bill_to_lines.append(sales_order.customer_name or '')
            if sales_order.customer_address:
                bill_to_lines.append(sales_order.customer_address)
            csz = [x for x in [sales_order.customer_city, sales_order.customer_state, sales_order.customer_zip] if x]
            if csz:
                bill_to_lines.append(', '.join(csz))
            if sales_order.customer_country:
                bill_to_lines.append(sales_order.customer_country)
        bill_to_text = '\n'.join(bill_to_lines) if bill_to_lines else (sales_order.customer_name or '')

        from django.utils import timezone
        pack_date = timezone.now().date()
        date_str = pack_date.strftime('%Y-%m-%d')
        order_date_str = sales_order.order_date.strftime('%Y-%m-%d') if sales_order.order_date else ''
        ship_date_dt = sales_order.actual_ship_date or sales_order.expected_ship_date
        ship_date_str = ship_date_dt.strftime('%Y-%m-%d') if ship_date_dt else ''
        po_ref = (sales_order.customer_reference_number or '').strip()
        po_so_str = f"{po_ref} / {sales_order.so_number}" if po_ref and sales_order.so_number else (sales_order.so_number or po_ref or '')
        latest_shipment = sales_order.shipments.order_by('-created_at').first()
        if latest_shipment:
            dims = (latest_shipment.dimensions or '').strip()
            pcs = latest_shipment.pieces
            pallet_line = []
            if dims:
                pallet_line.append(f"Dimensions: {dims}")
            if pcs is not None:
                pallet_line.append(f"Pieces: {pcs}")
            pallet_text = ', '.join(pallet_line) if pallet_line else ''
        else:
            pallet_text = ''

        if template_type == 'pdf':
            try:
                import fitz
                from erp_core.batch_ticket_pdf import _find_text_bbox
            except ImportError:
                return Response(
                    {'error': 'PyMuPDF (fitz) is required for PDF packing list. Install with: pip install pymupdf'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
            try:
                doc = fitz.open(str(template_path))
                if len(doc) < 1:
                    doc.close()
                    return Response({'error': 'Packing list template PDF has no pages.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                page = doc[0]
                fontname = 'helv'
                fontsize = 10
                color = (0, 0, 0)
                gap = 8
                line_height = 12

                def insert_after(page, label, value, max_len=120):
                    if value is None:
                        return
                    pos = _find_text_bbox(page, label)
                    if pos:
                        x, y, fs = pos
                        x = x + gap
                        val = str(value).replace('\r', '')[:max_len]
                        if '\n' in val:
                            for i, line in enumerate(val.split('\n')):
                                page.insert_text((x, y + fs * 0.8 + i * line_height), line[:max_len], fontsize=min(fs, fontsize), fontname=fontname, color=color)
                        else:
                            page.insert_text((x, y + fs * 0.8), val, fontsize=min(fs, fontsize), fontname=fontname, color=color)

                insert_after(page, 'Ship to:', ship_to_text)
                insert_after(page, 'Ship To:', ship_to_text)
                insert_after(page, 'Bill to:', bill_to_text)
                insert_after(page, 'Bill To:', bill_to_text)
                insert_after(page, 'Date:', date_str)
                insert_after(page, 'Order date:', order_date_str)
                insert_after(page, 'Order Date:', order_date_str)
                insert_after(page, 'Ship date:', ship_date_str)
                insert_after(page, 'Ship Date:', ship_date_str)
                insert_after(page, 'Purchase order/Sales order #', po_so_str)
                insert_after(page, 'PO/SO #', po_so_str)
                insert_after(page, 'PO #', po_ref)
                insert_after(page, 'SO #', sales_order.so_number or '')
                if pallet_text:
                    insert_after(page, 'Pallet 1 dimensions', pallet_text)
                    insert_after(page, 'pallet 1 dimensions', pallet_text)
                    insert_after(page, 'Dimensions', pallet_text)

                # Items: find table header row (e.g. "Description" or "Quantity") and add item rows below
                items = list(sales_order.items.all())
                if items:
                    anchor_pos = _find_text_bbox(page, 'Quantity') or _find_text_bbox(page, 'Description') or _find_text_bbox(page, 'Customer Item')
                    if anchor_pos:
                        ax, ay, afs = anchor_pos
                        row_y = ay + afs * 1.8
                        col_sku = ax - 180 if ax > 180 else 72
                        col_desc = col_sku + 80
                        col_qty = ax - 40
                        for item in items[:20]:
                            qty = getattr(item, 'quantity_ordered', 0) or 0
                            qty_display = f"{round(qty, 2):.2f}".rstrip('0').rstrip('.') if isinstance(qty, (int, float)) else str(qty)
                            uom = (item.item.unit_of_measure if getattr(item, 'item', None) and item.item else '') or ''
                            qty_with_uom = f"{qty_display} {uom}".strip() if uom else qty_display
                            sku = (item.item.sku if getattr(item, 'item', None) and item.item else '') or ''
                            name = (item.item.name if getattr(item, 'item', None) and item.item else '') or ''
                            page.insert_text((col_sku, row_y), (sku or name)[:30], fontsize=min(afs, 9), fontname=fontname, color=color)
                            page.insert_text((col_desc, row_y), name[:40] if name else sku[:40], fontsize=min(afs, 9), fontname=fontname, color=color)
                            page.insert_text((col_qty, row_y), qty_with_uom[:20], fontsize=min(afs, 9), fontname=fontname, color=color)
                            row_y += line_height + 2

                pdf_content = doc.tobytes()
                doc.close()
                response = HttpResponse(pdf_content, content_type='application/pdf')
                response['Content-Disposition'] = f'inline; filename="Packing_List_{sales_order.so_number}.pdf"'
                response['X-Content-Type-Options'] = 'nosniff'
                return response
            except Exception as e:
                import traceback
                logger.error('Packing list PDF generation failed: %s', str(e))
                logger.error(traceback.format_exc())
                return Response(
                    {'error': f'Failed to generate packing list PDF: {str(e)}'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )

        # DOCX path (legacy)
        try:
            from docx import Document
            doc = Document(str(template_path))
            
            # Replace placeholders (ship_to_text, bill_to_text, date_str, etc. already built above)
            replacements = {
                '{so_number}': sales_order.so_number,
                '{SO_NUMBER}': sales_order.so_number,
                '{customer_name}': sales_order.customer_name or '',
                '{CUSTOMER_NAME}': sales_order.customer_name or '',
                '{customer_address}': sales_order.customer_address or '',
                '{CUSTOMER_ADDRESS}': sales_order.customer_address or '',
                '{customer_city}': sales_order.customer_city or '',
                '{CUSTOMER_CITY}': sales_order.customer_city or '',
                '{customer_state}': sales_order.customer_state or '',
                '{CUSTOMER_STATE}': sales_order.customer_state or '',
                '{customer_zip}': sales_order.customer_zip or '',
                '{CUSTOMER_ZIP}': sales_order.customer_zip or '',
                '{customer_country}': sales_order.customer_country or '',
                '{CUSTOMER_COUNTRY}': sales_order.customer_country or '',
                '{customer_phone}': sales_order.customer_phone or '',
                '{CUSTOMER_PHONE}': sales_order.customer_phone or '',
                '{ship_to}': ship_to_text,
                '{SHIP_TO}': ship_to_text,
                '{bill_to}': bill_to_text,
                '{BILL_TO}': bill_to_text,
                '{date}': date_str,
                '{DATE}': date_str,
                '{order_date}': order_date_str,
                '{ORDER_DATE}': order_date_str,
                '{ship_date}': ship_date_str,
                '{SHIP_DATE}': ship_date_str,
                '{purchase_order}': po_so_str,
                '{PURCHASE_ORDER}': po_so_str,
                '{sales_order_number}': sales_order.so_number,
                '{SALES_ORDER_NUMBER}': sales_order.so_number,
                '{po_number}': (sales_order.customer_reference_number or '').strip(),
                '{PO_NUMBER}': (sales_order.customer_reference_number or '').strip(),
                '{customer_po}': (sales_order.customer_reference_number or '').strip(),
                '{CUSTOMER_PO}': (sales_order.customer_reference_number or '').strip(),
            }
            replacements['{pallet_dimensions}'] = pallet_text
            replacements['{PALLET_DIMENSIONS}'] = pallet_text
            replacements['{dimensions}'] = (latest_shipment.dimensions or '').strip() if latest_shipment else ''
            replacements['{DIMENSIONS}'] = replacements['{dimensions}']
            replacements['{pieces}'] = str(latest_shipment.pieces) if latest_shipment and latest_shipment.pieces is not None else ''
            replacements['{PIECES}'] = replacements['{pieces}']
            # Replace literal "pallet 1 dimensions" (or similar) with the filled value
            replacements['pallet 1 dimensions'] = pallet_text
            replacements['Pallet 1 dimensions'] = pallet_text
            replacements['Pallet 1 Dimensions'] = pallet_text
            
            def get_cell_text(cell):
                """Full text of a table cell (all paragraphs)."""
                return '\n'.join(p.text for p in cell.paragraphs).strip()
            
            def set_cell_text(cell, value):
                """Set entire cell content to value (clears all paragraphs, sets first)."""
                str_val = str(value) if value else ''
                for idx, p in enumerate(cell.paragraphs):
                    p.clear()
                    p.add_run(str_val if idx == 0 else '')
                if not cell.paragraphs:
                    cell.add_paragraph(str_val)
            
            def apply_replacements_to_paragraph(paragraph):
                text = paragraph.text
                if not text:
                    return
                new_text = text
                for key, value in replacements.items():
                    new_text = new_text.replace(key, str(value))
                if new_text != text:
                    paragraph.clear()
                    paragraph.add_run(new_text)
            
            for paragraph in doc.paragraphs:
                apply_replacements_to_paragraph(paragraph)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            apply_replacements_to_paragraph(paragraph)
            
            # Headers and footers (e.g. template may put date/PO in header)
            for section in doc.sections:
                try:
                    for p in section.header.paragraphs:
                        apply_replacements_to_paragraph(p)
                except Exception:
                    pass
                try:
                    for p in section.footer.paragraphs:
                        apply_replacements_to_paragraph(p)
                except Exception:
                    pass
            
            # Label-based fill: template may have "Ship to:", "Bill to:", etc. with value in next cell or same paragraph
            # Map label (normalized, no colon) -> value to fill
            po_ref = (sales_order.customer_reference_number or '').strip()
            label_values = {
                'ship to': ship_to_text,
                'bill to': bill_to_text,
                'date': date_str,
                'order date': order_date_str,
                'ship date': ship_date_str,
                'purchase order/sales order #': po_so_str,
                'po/so #': po_so_str,
                'purchase order': po_so_str,
                'sales order #': sales_order.so_number,
                'so #': sales_order.so_number,
                'po #': po_ref,
                'so #': sales_order.so_number,
                'order #': po_so_str,
                'purchase order #': po_ref,
            }
            
            def normalize_label(s):
                return (s or '').strip().lower().replace(':', '').replace('#', '').strip()
            
            def cell_matches_label(cell_text, label_key):
                """True if cell text is exactly or starts with the label (allowing trailing space/underscores)."""
                if not cell_text or not label_key:
                    return False
                norm = normalize_label(cell_text)
                norm_no_spaces = label_key.replace(' ', '')
                if norm == label_key or norm == norm_no_spaces:
                    return True
                # Also match "Ship to:" or "Ship to: ________" - normalized start
                norm_prefix = norm.split()[0] if norm else ''
                label_first = label_key.split()[0] if label_key else ''
                if norm.startswith(label_key) or norm.startswith(norm_no_spaces):
                    return True
                if label_first and norm.startswith(label_first):
                    return True
                return False
            
            # Table: when a cell contains only or starts with a label (e.g. "Ship to:"), fill the NEXT cell with the value
            for table in doc.tables:
                for row in table.rows:
                    cells = row.cells
                    for i, cell in enumerate(cells):
                        if i + 1 >= len(cells):
                            continue
                        cell_text = get_cell_text(cell)
                        for label_key, value in label_values.items():
                            if cell_matches_label(cell_text, label_key):
                                set_cell_text(cells[i + 1], value)
                                break
            
            # Body paragraphs: if a paragraph is exactly a label or starts with one, replace with "Label:\n\n" + value
            label_display = {
                'ship to': 'Ship to', 'bill to': 'Bill to', 'date': 'Date', 'order date': 'Order date',
                'ship date': 'Ship date', 'purchase order/sales order #': 'PO/SO #', 'po/so #': 'PO/SO #',
                'purchase order': 'PO/SO #', 'sales order #': 'SO #', 'so #': 'SO #', 'po #': 'PO #', 'order #': 'PO/SO #', 'purchase order #': 'PO #',
            }
            def fill_paragraph_if_label(paragraph, *, only_exact=False):
                t = paragraph.text.strip()
                norm = normalize_label(t)
                for label_key, value in label_values.items():
                    if only_exact:
                        if norm != label_key and norm != label_key.replace(' ', ''):
                            continue
                    else:
                        if not (norm == label_key or norm == label_key.replace(' ', '') or
                                norm.startswith(label_key) or (label_key.split()[0] and norm.startswith(label_key.split()[0]))):
                            continue
                    display = label_display.get(label_key, (t.split(':')[0] if ':' in t else t))
                    paragraph.clear()
                    paragraph.add_run(f"{display}:\n\n{value}" if value else f"{display}:")
                    return True
                return False
            for paragraph in doc.paragraphs:
                fill_paragraph_if_label(paragraph, only_exact=False)
            
            # Table cell: if a cell's text is exactly/starts with a label and there's no next cell (single-column), put value in same cell
            for table in doc.tables:
                for row in table.rows:
                    cells = row.cells
                    for i, cell in enumerate(cells):
                        if i + 1 < len(cells):
                            continue  # next-cell case already handled above
                        cell_text = get_cell_text(cell)
                        for label_key, value in label_values.items():
                            if not cell_matches_label(cell_text, label_key):
                                continue
                            display = label_display.get(label_key, (cell_text.split(':')[0] if ':' in cell_text else cell_text))
                            set_cell_text(cell, f"{display}:\n\n{value}" if value else f"{display}:")
                            break
            
            # Add items to packing list (if there's a table for items)
            items_added = False
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    row_text = ' '.join([cell.text for cell in row.cells]).lower()
                    if any(keyword in row_text for keyword in ['item', 'sku', 'description', 'quantity', 'qty']):
                        start_row = i + 1
                        while len(table.rows) > start_row:
                            table._element.remove(table.rows[start_row]._element)
                        
                        # Add items: quantity rounded to 2 decimals and include UoM in quantity column (e.g. "500 kg")
                        for item in sales_order.items.all():
                            new_row = table.add_row()
                            cells = new_row.cells
                            qty = item.quantity_ordered
                            qty_display = f"{round(qty, 2):.2f}".rstrip('0').rstrip('.') if isinstance(qty, (int, float)) else str(qty)
                            uom = (item.item.unit_of_measure if item.item else '') or ''
                            qty_with_uom = f"{qty_display} {uom}".strip() if uom else qty_display
                            if len(cells) >= 1:
                                cells[0].text = item.item.sku if item.item else ''
                            if len(cells) >= 2:
                                cells[1].text = item.item.name if item.item else ''
                            if len(cells) >= 3:
                                cells[2].text = qty_with_uom
                            if len(cells) >= 4:
                                cells[3].text = item.item.unit_of_measure if item.item else ''
                            if len(cells) >= 5:
                                lot_numbers = []
                                for allocation in SalesOrderLot.objects.filter(sales_order_item=item):
                                    lot_numbers.append(allocation.lot.lot_number)
                                cells[4].text = ', '.join(lot_numbers) if lot_numbers else ''
                        items_added = True
                        break
                if items_added:
                    break
            
            # Save to temporary file
            import time
            tmp_path = None
            pdf_content = None
            is_pdf = False
            
            try:
                # Create temp file with unique name
                tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
                os.close(tmp_fd)  # Close the file descriptor so doc.save can use it
                
                # Save document
                doc.save(tmp_path)
                
                # Try to convert to PDF
                try:
                    from docx2pdf import convert
                    pdf_path = tmp_path.replace('.docx', '.pdf')
                    convert(tmp_path, pdf_path)
                    
                    # Wait a moment to ensure file is fully written and closed
                    time.sleep(0.1)
                    
                    # Read PDF content
                    with open(pdf_path, 'rb') as pdf_file:
                        pdf_content = pdf_file.read()
                    is_pdf = pdf_content.startswith(b'%PDF')
                    
                    # Clean up temp PDF file
                    try:
                        if os.path.exists(pdf_path):
                            os.unlink(pdf_path)
                    except Exception as e:
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.warning(f"Could not delete temp PDF file {pdf_path}: {e}")
                        
                except Exception as e:
                    # If PDF conversion fails, use DOCX file
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.warning(f"PDF conversion failed: {e}. Using DOCX file instead.")
                    with open(tmp_path, 'rb') as docx_file:
                        pdf_content = docx_file.read()
                    is_pdf = False
                
                # Clean up temp DOCX file
                try:
                    if tmp_path and os.path.exists(tmp_path):
                        # Wait a moment to ensure file is closed
                        time.sleep(0.1)
                        os.unlink(tmp_path)
                except Exception as e:
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.warning(f"Could not delete temp DOCX file {tmp_path}: {e}")
                    # Try again after a longer delay
                    try:
                        time.sleep(0.5)
                        if os.path.exists(tmp_path):
                            os.unlink(tmp_path)
                    except:
                        pass  # If it still fails, just continue - temp files will be cleaned up by OS
            except Exception as e:
                # If file operations fail, log and return error
                import logging
                import traceback
                logger = logging.getLogger(__name__)
                logger.error(f'Failed to save/convert packing list document: {str(e)}')
                logger.error(f'Traceback: {traceback.format_exc()}')
                return Response(
                    {'error': f'Failed to generate packing list document: {str(e)}'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
            
            # Ensure we have content before returning
            if pdf_content is None or len(pdf_content) == 0:
                return Response(
                    {'error': 'Failed to generate packing list document - no content'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
            
            # Return as HTTP response
            filename = f"Packing_List_{sales_order.so_number}.pdf" if is_pdf else f"Packing_List_{sales_order.so_number}.docx"
            content_type = 'application/pdf' if is_pdf else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
            try:
                response = HttpResponse(pdf_content, content_type=content_type)
                response['Content-Disposition'] = f'inline; filename="{filename}"'
                response['X-Content-Type-Options'] = 'nosniff'
                response['Content-Length'] = str(len(pdf_content))
                return response
            except Exception as e:
                import logging
                logger = logging.getLogger(__name__)
                logger.error(f'Failed to create HTTP response: {str(e)}')
                return Response(
                    {'error': f'Failed to create HTTP response: {str(e)}'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
            
        except ImportError:
            return Response(
                {'error': 'python-docx is not installed. Please install it with: pip install python-docx'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        except Exception as e:
            import logging
            import traceback
            logger = logging.getLogger(__name__)
            logger.error(f'Failed to generate packing list: {str(e)}')
            logger.error(f'Traceback: {traceback.format_exc()}')
            return Response(
                {'error': f'Failed to generate packing list: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class ShipmentViewSet(viewsets.ReadOnlyModelViewSet):
    """List/retrieve shipments. Packing list PDF per shipment. Filter by ?sales_order=<id>."""
    serializer_class = ShipmentSerializer
    queryset = Shipment.objects.select_related(
        'sales_order', 'sales_order__customer', 'sales_order__ship_to_location'
    ).prefetch_related('items__sales_order_item__item').all()

    def get_queryset(self):
        qs = super().get_queryset()
        so_id = self.request.query_params.get('sales_order')
        if so_id:
            qs = qs.filter(sales_order_id=so_id)
        return qs

    @action(detail=True, methods=['get'], url_path='packing-list', url_name='packing-list')
    def packing_list(self, request, pk=None):
        """Generate packing list PDF for this shipment (one PDF per release)."""
        from django.http import HttpResponse
        shipment = self.get_object()
        try:
            from .packing_list_pdf_html import generate_packing_list_pdf_from_shipment
            pdf_content = generate_packing_list_pdf_from_shipment(shipment)
            if pdf_content:
                so_num = (shipment.sales_order.so_number or pk) if shipment.sales_order_id else pk
                response = HttpResponse(pdf_content, content_type='application/pdf')
                response['Content-Disposition'] = f'inline; filename="Packing_List_{so_num}_Release_{shipment.id}.pdf"'
                response['X-Content-Type-Options'] = 'nosniff'
                return response
        except Exception as e:
            import traceback
            logger.error('Shipment packing list PDF failed: %s', str(e))
            logger.error(traceback.format_exc())
        return Response(
            {'error': 'Failed to generate packing list for this shipment.'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


class InvoiceViewSet(viewsets.ModelViewSet):
    serializer_class = InvoiceSerializer
    queryset = Invoice.objects.none()  # Will be set in get_queryset
    
    def get_queryset(self):
        from django.db import connection
        
        # Check if sales_order_id column exists
        has_sales_order_column = False
        try:
            with connection.cursor() as cursor:
                cursor.execute("PRAGMA table_info(erp_core_invoice)")
                columns = [row[1] for row in cursor.fetchall()]
                has_sales_order_column = 'sales_order_id' in columns
        except Exception:
            pass
        
        # Build queryset - use raw SQL if column doesn't exist to avoid ORM errors
        if has_sales_order_column:
            queryset = Invoice.objects.select_related(
                'sales_order', 'sales_order__customer', 'sales_order__ship_to_location'
            ).prefetch_related('items', 'sales_order__shipments').all()
        else:
            # No sales_order_id column - use raw SQL to avoid ORM trying to access it
            # We'll manually construct the queryset using raw SQL
            from django.db import models
            queryset = Invoice.objects.raw("SELECT * FROM erp_core_invoice")
            # Convert to a list and then back to a queryset-like object
            # Actually, we need to use a different approach - use only() to exclude the field
            # But Django doesn't support excluding ForeignKey fields easily
            # So we'll use raw SQL in the list method instead
            queryset = Invoice.objects.all()
            # Override the queryset to not access sales_order
            # We'll handle this in list() method
        
        status_filter = self.request.query_params.get('status', None)
        customer_id = self.request.query_params.get('customer_id', None)
        start_date = self.request.query_params.get('start_date', None)
        end_date = self.request.query_params.get('end_date', None)
        
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        if customer_id and has_sales_order_column:
            queryset = queryset.filter(sales_order__customer_id=customer_id)
        if start_date:
            queryset = queryset.filter(invoice_date__gte=start_date)
        if end_date:
            queryset = queryset.filter(invoice_date__lte=end_date)
        
        return queryset
    
    def list(self, request, *args, **kwargs):
        """Override list to handle schema mismatches"""
        from django.db import connection
        from rest_framework.response import Response
        from rest_framework import status as http_status
        from django.utils import timezone
        from datetime import date
        
        try:
            # First, automatically update sent invoices to overdue if past due date
            today = timezone.now().date()
            try:
                # Update invoices that are 'sent' and past due date to 'overdue'
                # Use raw sqlite3 to avoid Django's query logging issues
                import sqlite3
                from django.conf import settings
                db_path = settings.DATABASES['default']['NAME']
                raw_conn = sqlite3.connect(db_path)
                raw_cursor = raw_conn.cursor()
                try:
                    raw_cursor.execute(
                        "UPDATE erp_core_invoice SET status = 'overdue' WHERE status = 'sent' AND due_date < ?",
                        [today]
                    )
                    raw_conn.commit()
                finally:
                    raw_conn.close()
            except Exception as update_error:
                import logging
                logger = logging.getLogger(__name__)
                logger.warning(f"Error auto-updating overdue invoices: {update_error}")
            
            # Check what columns exist and query invoices in the same cursor context
            with connection.cursor() as cursor:
                cursor.execute("PRAGMA table_info(erp_core_invoice)")
                column_info = cursor.fetchall()
                columns = {col[1]: col for col in column_info}
                has_sales_order_column = 'sales_order_id' in columns
                has_freight = 'freight' in columns
                has_tax = 'tax' in columns
                has_grand_total = 'grand_total' in columns
                has_discount = 'discount' in columns
                
                # If model fields don't match DB, use raw SQL directly (ORM will fail)
                if not has_freight or not has_tax or not has_grand_total or not has_discount:
                    # Use raw SQL approach directly - ORM will fail because model fields don't match DB
                    # Use raw SQL to query invoices
                    query = "SELECT * FROM erp_core_invoice WHERE 1=1"
                    params = []
                    
                    # Get query params - handle both DRF Request and WSGIRequest
                    if hasattr(self.request, 'query_params'):
                        query_params = self.request.query_params
                    else:
                        query_params = self.request.GET
                    
                    status_filter = query_params.get('status', None)
                    if status_filter:
                        query += " AND status = ?"
                        params.append(status_filter)
                    
                    start_date = query_params.get('start_date', None)
                    if start_date:
                        query += " AND invoice_date >= ?"
                        params.append(start_date)
                    
                    end_date = query_params.get('end_date', None)
                    if end_date:
                        query += " AND invoice_date <= ?"
                        params.append(end_date)
                    
                    query += " ORDER BY invoice_date DESC, created_at DESC"
                    
                    # Execute query using the cursor we already have
                    # Since we're using parameterized queries with ? placeholders, 
                    # Django's query logging shouldn't have % formatting issues
                    cursor.execute(query, params)
                    rows = cursor.fetchall()
                    
                    # Get column names
                    column_names = [col[1] for col in column_info]
                    
                    # Convert rows to dicts and serialize
                    invoice_data_list = []
                    for row in rows:
                        invoice_dict = dict(zip(column_names, row))
                        # Map database field names to model field names
                        mapped_dict = {}
                        for key, value in invoice_dict.items():
                            # Map tax_amount -> tax, total_amount -> grand_total
                            if key == 'tax_amount':
                                mapped_dict['tax'] = value
                            elif key == 'total_amount':
                                mapped_dict['grand_total'] = value
                            else:
                                mapped_dict[key] = value
                        # Set defaults for missing fields
                        if 'freight' not in mapped_dict:
                            mapped_dict['freight'] = 0.0
                        if 'tax' not in mapped_dict:
                            mapped_dict['tax'] = mapped_dict.get('tax_amount', 0.0)
                        if 'discount' not in mapped_dict:
                            mapped_dict['discount'] = 0.0
                        if 'grand_total' not in mapped_dict:
                            mapped_dict['grand_total'] = mapped_dict.get('total_amount', 0.0)
                        invoice_data_list.append(mapped_dict)
                    
                    # Create Invoice objects from the mapped data
                    # We need to map database fields to model fields
                    invoices = []
                    for data in invoice_data_list:
                        invoice = Invoice()
                        # Set basic fields that exist in both DB and model
                        if 'id' in data:
                            invoice.id = data['id']
                        if 'invoice_number' in data:
                            invoice.invoice_number = data['invoice_number']
                        if 'invoice_date' in data:
                            invoice.invoice_date = data['invoice_date']
                        if 'due_date' in data:
                            invoice.due_date = data['due_date']
                        if 'status' in data:
                            invoice.status = data['status']
                        if 'subtotal' in data:
                            invoice.subtotal = data['subtotal']
                        if 'notes' in data:
                            invoice.notes = data['notes']
                        if 'created_at' in data:
                            invoice.created_at = data['created_at']
                        if 'updated_at' in data:
                            invoice.updated_at = data['updated_at']
                        
                        # Map database fields to model fields
                        # DB has tax_amount, model has tax
                        if 'tax' in data:
                            invoice.tax = data['tax']
                        elif 'tax_amount' in data:
                            invoice.tax = data.get('tax_amount', 0.0)
                        
                        # DB has total_amount, model has grand_total
                        if 'grand_total' in data:
                            invoice.grand_total = data['grand_total']
                        elif 'total_amount' in data:
                            invoice.grand_total = data.get('total_amount', 0.0)
                        
                        # Set defaults for missing fields
                        invoice.freight = data.get('freight', 0.0)
                        invoice.discount = data.get('discount', 0.0)
                        
                        # Store customer_vendor_name and customer_vendor_id if they exist (for serializer)
                        # These are database columns but not model fields
                        if 'customer_vendor_name' in data:
                            setattr(invoice, 'customer_vendor_name', data['customer_vendor_name'])
                        if 'customer_vendor_id' in data:
                            setattr(invoice, 'customer_vendor_id', data['customer_vendor_id'])
                        
                        # Handle sales_order if sales_order_id exists
                        # Set sales_order_id directly on the model (Django will handle the FK)
                        if 'sales_order_id' in data and data['sales_order_id']:
                            # Set the foreign key ID directly
                            invoice.sales_order_id = data['sales_order_id']
                            # Also store in temp attribute for serializer fallback
                            invoice._sales_order_id = data['sales_order_id']
                        else:
                            invoice.sales_order_id = None
                            invoice._sales_order_id = None
                        
                        invoice._state.adding = False
                        invoice._state.db = connection
                        invoices.append(invoice)
                    
                    # Prefetch items and sales orders for each invoice to avoid N+1 queries
                    from .models import InvoiceItem, SalesOrder
                    # Collect all sales_order_ids first for batch loading
                    sales_order_ids = []
                    for invoice in invoices:
                        if hasattr(invoice, 'sales_order_id') and invoice.sales_order_id:
                            sales_order_ids.append(invoice.sales_order_id)
                        elif hasattr(invoice, '_sales_order_id') and invoice._sales_order_id:
                            sales_order_ids.append(invoice._sales_order_id)
                    
                    # Batch load all sales orders with customers
                    sales_orders_dict = {}
                    if sales_order_ids:
                        for so in SalesOrder.objects.select_related('customer').filter(id__in=sales_order_ids):
                            sales_orders_dict[so.id] = so
                    
                    # Attach sales orders to invoices
                    for invoice in invoices:
                        try:
                            # Prefetch items for this invoice
                            invoice._prefetched_objects_cache = {
                                'items': list(InvoiceItem.objects.filter(invoice_id=invoice.id))
                            }
                            
                            # Get sales_order_id from any source
                            sales_order_id = None
                            if hasattr(invoice, 'sales_order_id') and invoice.sales_order_id:
                                sales_order_id = invoice.sales_order_id
                            elif hasattr(invoice, '_sales_order_id') and invoice._sales_order_id:
                                sales_order_id = invoice._sales_order_id
                            
                            # Attach sales order if we have it
                            if sales_order_id and sales_order_id in sales_orders_dict:
                                invoice.sales_order = sales_orders_dict[sales_order_id]
                                invoice.sales_order_id = sales_order_id
                                invoice._sales_order_id = sales_order_id
                            elif sales_order_id:
                                # Try to load it individually (fallback)
                                try:
                                    so = SalesOrder.objects.select_related('customer').get(id=sales_order_id)
                                    invoice.sales_order = so
                                    invoice.sales_order_id = sales_order_id
                                    invoice._sales_order_id = sales_order_id
                                except SalesOrder.DoesNotExist:
                                    invoice.sales_order = None
                                    invoice.sales_order_id = None
                                    invoice._sales_order_id = None
                            else:
                                invoice.sales_order = None
                                invoice.sales_order_id = None
                                invoice._sales_order_id = None
                        except Exception as e:
                            import logging
                            logger = logging.getLogger(__name__)
                            logger.error(f"Error loading sales order for invoice {invoice.id}: {e}")
                            invoice._prefetched_objects_cache = {'items': []}
                            invoice.sales_order = None
                    
                    # Serialize invoices
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.info(f"Returning {len(invoices)} invoices from raw SQL")
                    
                    # Debug: Log first invoice before serialization
                    if invoices:
                        logger.info(f"First invoice before serialization: id={invoices[0].id}, sales_order_id={getattr(invoices[0], 'sales_order_id', None)}, has_sales_order={hasattr(invoices[0], 'sales_order') and invoices[0].sales_order is not None}")
                    
                    try:
                        serializer = self.get_serializer(invoices, many=True)
                        serialized_data = serializer.data
                        
                        # Debug: Log first invoice after serialization
                        if serialized_data:
                            logger.info(f"First invoice after serialization: sales_order={serialized_data[0].get('sales_order')}, payment_terms={serialized_data[0].get('payment_terms')}")
                        
                        return Response(serialized_data)
                    except Exception as ser_error:
                        logger.error(f"Error serializing invoices: {ser_error}")
                        import traceback
                        logger.error(traceback.format_exc())
                        # Try to return at least basic data
                        basic_data = []
                        for inv in invoices:
                            basic_data.append({
                                'id': inv.id,
                                'invoice_number': inv.invoice_number,
                                'invoice_date': str(inv.invoice_date) if inv.invoice_date else None,
                                'due_date': str(inv.due_date) if inv.due_date else None,
                                'status': inv.status,
                                'subtotal': inv.subtotal,
                                'tax': getattr(inv, 'tax', 0.0),
                                'grand_total': getattr(inv, 'grand_total', 0.0),
                                'freight': getattr(inv, 'freight', 0.0),
                                'discount': getattr(inv, 'discount', 0.0),
                                'customer_vendor_name': getattr(inv, 'customer_vendor_name', None),
                                'sales_order_id': getattr(inv, 'sales_order_id', None),
                            })
                        return Response(basic_data)
                else:
                    # Schema matches - use normal queryset
                    # But we need to handle it outside the cursor context
                    pass
            
            # If we get here, schema matches - use normal queryset
            return super().list(request, *args, **kwargs)
        except Exception as e:
            import logging
            import traceback
            logger = logging.getLogger(__name__)
            logger.error(f"Error listing invoices: {e}")
            logger.error(traceback.format_exc())
            # Return more detailed error in debug mode
            error_detail = str(e)
            if hasattr(e, '__traceback__'):
                import traceback as tb
                error_detail += f"\n{tb.format_exc()}"
            return Response({
                'error': str(e),
                'detail': error_detail,
                'traceback': traceback.format_exc() if settings.DEBUG else None
            }, status=http_status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    def retrieve(self, request, *args, **kwargs):
        """Override retrieve to handle schema mismatches"""
        from django.db import connection
        from rest_framework.response import Response
        from rest_framework import status as http_status
        from django.conf import settings
        from django.utils import timezone
        from datetime import date
        
        try:
            invoice_id = kwargs.get('pk')
            
            # First, automatically update sent invoices to overdue if past due date
            today = timezone.now().date()
            try:
                # Update this specific invoice if it's 'sent' and past due date
                # Use raw sqlite3 to avoid Django's query logging issues
                import sqlite3
                from django.conf import settings
                db_path = settings.DATABASES['default']['NAME']
                raw_conn = sqlite3.connect(db_path)
                raw_cursor = raw_conn.cursor()
                try:
                    raw_cursor.execute(
                        "UPDATE erp_core_invoice SET status = 'overdue' WHERE id = ? AND status = 'sent' AND due_date < ?",
                        [invoice_id, today]
                    )
                    raw_conn.commit()
                finally:
                    raw_conn.close()
            except Exception as update_error:
                import logging
                logger = logging.getLogger(__name__)
                logger.warning(f"Error auto-updating overdue invoice: {update_error}")
            
            # Check what columns exist - use raw sqlite3 to avoid query logging issues
            import sqlite3
            from django.conf import settings
            db_path = settings.DATABASES['default']['NAME']
            raw_conn = sqlite3.connect(db_path)
            raw_cursor = raw_conn.cursor()
            
            try:
                # Get column info
                raw_cursor.execute("PRAGMA table_info(erp_core_invoice)")
                column_info = raw_cursor.fetchall()
                columns = {col[1]: col for col in column_info}
                column_names = [col[1] for col in column_info]
                has_sales_order_column = 'sales_order_id' in columns
                has_freight = 'freight' in columns
                has_tax = 'tax' in columns
                has_grand_total = 'grand_total' in columns
                has_discount = 'discount' in columns
                
                # If model fields don't match DB, use raw SQL
                if not has_freight or not has_tax or not has_grand_total or not has_discount:
                    # Use raw SQL to get invoice
                    raw_cursor.execute("SELECT * FROM erp_core_invoice WHERE id = ?", [invoice_id])
                    row = raw_cursor.fetchone()
                    
                    if not row:
                        return Response({'error': 'Invoice not found'}, status=http_status.HTTP_404_NOT_FOUND)
                    invoice_dict = dict(zip(column_names, row))
                    
                    # Map database fields to model fields
                    mapped_dict = {}
                    for key, value in invoice_dict.items():
                        if key == 'tax_amount':
                            mapped_dict['tax'] = value
                        elif key == 'total_amount':
                            mapped_dict['grand_total'] = value
                        else:
                            mapped_dict[key] = value
                    
                    # Set defaults
                    if 'freight' not in mapped_dict:
                        mapped_dict['freight'] = 0.0
                    if 'tax' not in mapped_dict:
                        mapped_dict['tax'] = mapped_dict.get('tax_amount', 0.0)
                    if 'discount' not in mapped_dict:
                        mapped_dict['discount'] = 0.0
                    if 'grand_total' not in mapped_dict:
                        mapped_dict['grand_total'] = mapped_dict.get('total_amount', 0.0)
                    
                    # Create Invoice object
                    invoice = Invoice()
                    if 'id' in mapped_dict:
                        invoice.id = mapped_dict['id']
                    if 'invoice_number' in mapped_dict:
                        invoice.invoice_number = mapped_dict['invoice_number']
                    if 'invoice_date' in mapped_dict:
                        invoice.invoice_date = mapped_dict['invoice_date']
                    if 'due_date' in mapped_dict:
                        invoice.due_date = mapped_dict['due_date']
                    if 'status' in mapped_dict:
                        invoice.status = mapped_dict['status']
                    if 'subtotal' in mapped_dict:
                        invoice.subtotal = mapped_dict['subtotal']
                    if 'notes' in mapped_dict:
                        invoice.notes = mapped_dict['notes']
                    if 'created_at' in mapped_dict:
                        invoice.created_at = mapped_dict['created_at']
                    if 'updated_at' in mapped_dict:
                        invoice.updated_at = mapped_dict['updated_at']
                    
                    invoice.tax = mapped_dict.get('tax', 0.0)
                    invoice.grand_total = mapped_dict.get('grand_total', 0.0)
                    invoice.freight = mapped_dict.get('freight', 0.0)
                    invoice.discount = mapped_dict.get('discount', 0.0)
                    
                    if 'sales_order_id' in mapped_dict and mapped_dict['sales_order_id']:
                        invoice.sales_order_id = mapped_dict['sales_order_id']
                        invoice._sales_order_id = mapped_dict['sales_order_id']
                    
                    invoice._state.adding = False
                    invoice._state.db = connection
                    
                    # Prefetch items using raw SQL
                    from .models import InvoiceItem, SalesOrder
                    try:
                        raw_cursor.execute("PRAGMA table_info(erp_core_invoiceitem)")
                        item_columns = [row[1] for row in raw_cursor.fetchall()]
                        
                        if 'invoice_id' in item_columns:
                            raw_cursor.execute("SELECT * FROM erp_core_invoiceitem WHERE invoice_id = ?", [invoice.id])
                            item_rows = raw_cursor.fetchall()
                            
                            items = []
                            for item_row in item_rows:
                                item_dict = dict(zip(item_columns, item_row))
                                item = InvoiceItem()
                                if 'id' in item_dict:
                                    item.id = item_dict['id']
                                if 'invoice_id' in item_dict:
                                    item.invoice_id = item_dict['invoice_id']
                                if 'description' in item_dict:
                                    item.description = item_dict['description']
                                if 'quantity' in item_dict:
                                    item.quantity = item_dict['quantity']
                                if 'unit_price' in item_dict:
                                    item.unit_price = item_dict['unit_price']
                                if 'total' in item_dict:
                                    item.total = item_dict['total']
                                elif 'line_total' in item_dict:
                                    item.total = item_dict['line_total']
                                if 'item_id' in item_dict:
                                    item.item_id = item_dict['item_id']
                                if 'sales_order_item_id' in item_dict:
                                    item.sales_order_item_id = item_dict['sales_order_item_id']
                                item._state.adding = False
                                item._state.db = connection
                                items.append(item)
                            
                            invoice._prefetched_objects_cache = {'items': items}
                        else:
                            invoice._prefetched_objects_cache = {'items': []}
                    except Exception as item_error:
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.warning(f"Error loading invoice items: {item_error}")
                        invoice._prefetched_objects_cache = {'items': []}
                    
                    # Load sales order if exists
                    if invoice.sales_order_id:
                        try:
                            invoice.sales_order = SalesOrder.objects.select_related('customer', 'ship_to_location').get(id=invoice.sales_order_id)
                        except SalesOrder.DoesNotExist:
                            invoice.sales_order = None
                    
                    serializer = self.get_serializer(invoice)
                    return Response(serializer.data)
                else:
                    # Schema matches - use normal queryset
                    return super().retrieve(request, *args, **kwargs)
            finally:
                raw_conn.close()
        except Exception as e:
            import logging
            import traceback
            logger = logging.getLogger(__name__)
            logger.error(f"Error retrieving invoice: {e}")
            logger.error(traceback.format_exc())
            error_detail = str(e)
            if settings.DEBUG:
                error_detail += f"\n{traceback.format_exc()}"
            return Response({
                'error': str(e),
                'detail': error_detail
            }, status=http_status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=False, methods=['get'])
    def aging_report(self, request):
        """Get aging report grouped by aging buckets"""
        invoices = self.get_queryset().exclude(status='paid')
        
        aging_buckets = {
            'current': [],  # Not yet due
            '0_30': [],     # 0-30 days overdue
            '31_60': [],    # 31-60 days overdue
            '61_90': [],    # 61-90 days overdue
            '90_plus': []   # 90+ days overdue
        }
        
        for invoice in invoices:
            days = invoice.days_aging
            if days < 0:
                aging_buckets['current'].append(invoice)
            elif days <= 30:
                aging_buckets['0_30'].append(invoice)
            elif days <= 60:
                aging_buckets['31_60'].append(invoice)
            elif days <= 90:
                aging_buckets['61_90'].append(invoice)
            else:
                aging_buckets['90_plus'].append(invoice)
        
        # Calculate totals
        totals = {}
        for bucket, invs in aging_buckets.items():
            totals[bucket] = sum(inv.grand_total for inv in invs)
        
        serializer = self.get_serializer(list(aging_buckets.values())[0] + list(aging_buckets.values())[1] + 
                                        list(aging_buckets.values())[2] + list(aging_buckets.values())[3] + 
                                        list(aging_buckets.values())[4], many=True)
        
        return Response({
            'buckets': {
                'current': {'invoices': [self.get_serializer(inv).data for inv in aging_buckets['current']], 'total': totals['current']},
                '0_30': {'invoices': [self.get_serializer(inv).data for inv in aging_buckets['0_30']], 'total': totals['0_30']},
                '31_60': {'invoices': [self.get_serializer(inv).data for inv in aging_buckets['31_60']], 'total': totals['31_60']},
                '61_90': {'invoices': [self.get_serializer(inv).data for inv in aging_buckets['61_90']], 'total': totals['61_90']},
                '90_plus': {'invoices': [self.get_serializer(inv).data for inv in aging_buckets['90_plus']], 'total': totals['90_plus']},
            },
            'grand_total': sum(totals.values())
        })
    
    def create(self, request, *args, **kwargs):
        """Create a manual invoice (not tied to a sales order)"""
        from django.db import connection
        from rest_framework.response import Response
        from rest_framework import status as http_status
        from .models import Invoice, InvoiceItem
        import logging
        logger = logging.getLogger(__name__)
        
        try:
            data = request.data.copy()
            
            # Generate invoice number
            invoice_number = generate_invoice_number()
            
            # Get invoice date
            invoice_date = data.get('invoice_date')
            if isinstance(invoice_date, str):
                from datetime import datetime
                invoice_date = datetime.strptime(invoice_date, '%Y-%m-%d').date()
            
            # Calculate due date from payment terms if provided
            due_date = data.get('due_date')
            if isinstance(due_date, str):
                from datetime import datetime
                due_date = datetime.strptime(due_date, '%Y-%m-%d').date()
            elif not due_date and invoice_date:
                # Default to 30 days if not provided
                from datetime import timedelta
                due_date = invoice_date + timedelta(days=30)
            
            # Calculate totals
            items_data = data.get('items', [])
            subtotal = sum(float(item.get('line_total', 0) or item.get('quantity', 0) * item.get('unit_price', 0)) for item in items_data)
            tax = float(data.get('tax_amount', 0) or 0)
            freight = float(data.get('freight', 0) or 0)
            discount = float(data.get('discount', 0) or 0)
            grand_total = subtotal + tax + freight - discount
            
            # Check database schema
            with connection.cursor() as cursor:
                cursor.execute("PRAGMA table_info(erp_core_invoice)")
                columns = [row[1] for row in cursor.fetchall()]
                has_sales_order_column = 'sales_order_id' in columns
                has_freight = 'freight' in columns
                has_tax = 'tax' in columns
                has_grand_total = 'grand_total' in columns
                has_discount = 'discount' in columns
            
            # Create invoice using raw SQL if schema doesn't match
            if not has_freight or not has_tax or not has_grand_total or not has_discount:
                import sqlite3
                from django.conf import settings
                db_path = settings.DATABASES['default']['NAME']
                raw_conn = sqlite3.connect(db_path)
                raw_cursor = raw_conn.cursor()
                try:
                    # Build column list based on what exists
                    col_names = ['invoice_number', 'invoice_date', 'due_date', 'status', 'subtotal', 'notes']
                    col_values = [invoice_number, invoice_date, due_date, 'draft', subtotal, data.get('notes', '')]
                    
                    if 'freight' in columns:
                        col_names.append('freight')
                        col_values.append(freight)
                    if 'tax' in columns or 'tax_amount' in columns:
                        col_name = 'tax' if 'tax' in columns else 'tax_amount'
                        col_names.append(col_name)
                        col_values.append(tax)
                    if 'discount' in columns:
                        col_names.append('discount')
                        col_values.append(discount)
                    if 'grand_total' in columns or 'total_amount' in columns:
                        col_name = 'grand_total' if 'grand_total' in columns else 'total_amount'
                        col_names.append(col_name)
                        col_values.append(grand_total)
                    if has_sales_order_column:
                        col_names.append('sales_order_id')
                        col_values.append(None)  # Manual invoice, no sales order
                    if 'customer_vendor_name' in columns:
                        col_names.append('customer_vendor_name')
                        col_values.append(data.get('customer_vendor_name', ''))
                    if 'customer_vendor_id' in columns:
                        col_names.append('customer_vendor_id')
                        col_values.append(data.get('customer_vendor_id', ''))
                    
                    placeholders = ','.join(['?' for _ in col_names])
                    insert_sql = f"INSERT INTO erp_core_invoice ({','.join(col_names)}) VALUES ({placeholders})"
                    raw_cursor.execute(insert_sql, col_values)
                    invoice_id = raw_cursor.lastrowid
                    raw_conn.commit()
                finally:
                    raw_conn.close()
            else:
                # Use ORM
                invoice_data = {
                    'invoice_number': invoice_number,
                    'invoice_date': invoice_date,
                    'due_date': due_date,
                    'status': 'draft',
                    'subtotal': subtotal,
                    'freight': freight,
                    'tax': tax,
                    'discount': discount,
                    'grand_total': grand_total,
                    'notes': data.get('notes', ''),
                }
                if has_sales_order_column:
                    invoice_data['sales_order'] = None  # Manual invoice
                if 'customer_vendor_name' in columns:
                    invoice_data['customer_vendor_name'] = data.get('customer_vendor_name', '')
                if 'customer_vendor_id' in columns:
                    invoice_data['customer_vendor_id'] = data.get('customer_vendor_id', '')
                
                invoice = Invoice.objects.create(**invoice_data)
                invoice_id = invoice.id
            
            # Create invoice items
            for item_data in items_data:
                item_id = item_data.get('item_id')
                quantity = float(item_data.get('quantity', 0))
                unit_price = float(item_data.get('unit_price', 0))
                line_total = float(item_data.get('line_total', quantity * unit_price))
                description = item_data.get('description', '')
                
                inv_item_kw = dict(
                    invoice_id=invoice_id,
                    description=description,
                    quantity=quantity,
                    unit_price=unit_price,
                    total=line_total,
                    notes=''
                )
                if item_id:
                    inv_item_kw['item_id'] = item_id
                InvoiceItem.objects.create(**inv_item_kw)
            
            # Return created invoice
            invoice = Invoice.objects.get(id=invoice_id)
            serializer = self.get_serializer(invoice)
            return Response(serializer.data, status=http_status.HTTP_201_CREATED)
            
        except Exception as e:
            logger.error(f"Error creating invoice: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return Response(
                {'error': f'Failed to create invoice: {str(e)}'},
                status=http_status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    def update(self, request, *args, **kwargs):
        """Update invoice - handle status updates and other fields"""
        from django.db import connection
        from rest_framework.response import Response
        from rest_framework import status as http_status
        from .models import Invoice
        from django.conf import settings
        import logging
        logger = logging.getLogger(__name__)
        
        try:
            partial = kwargs.pop('partial', False)
            invoice_id = kwargs.get('pk')
            data = request.data.copy()
            
            # Before moving from Draft to Sent (issued): require carrier and tracking on the sales order
            if data.get('status') == 'sent':
                try:
                    inv = Invoice.objects.select_related('sales_order').get(pk=invoice_id)
                    if inv.sales_order_id and inv.sales_order:
                        so = inv.sales_order
                        carrier = (getattr(so, 'carrier', None) or '').strip()
                        tracking = (getattr(so, 'tracking_number', None) or '').strip()
                        if not carrier or not tracking:
                            return Response(
                                {'error': 'Carrier and tracking number must be entered on the sales order before the invoice can be moved to Issued/Sent.'},
                                status=http_status.HTTP_400_BAD_REQUEST
                            )
                except Invoice.DoesNotExist:
                    pass
            
            # Check database schema
            with connection.cursor() as cursor:
                cursor.execute("PRAGMA table_info(erp_core_invoice)")
                column_info = cursor.fetchall()
                columns = {col[1]: col for col in column_info}
                column_names = [col[1] for col in column_info]
                has_sales_order_column = 'sales_order_id' in columns
                has_freight = 'freight' in columns
                has_tax = 'tax' in columns
                has_grand_total = 'grand_total' in columns
                has_discount = 'discount' in columns
                
                # If model fields don't match DB, use raw SQL for update
                if not has_freight or not has_tax or not has_grand_total or not has_discount:
                    # Use raw SQL to update
                    import sqlite3
                    from django.conf import settings
                    db_path = settings.DATABASES['default']['NAME']
                    raw_conn = sqlite3.connect(db_path)
                    raw_cursor = raw_conn.cursor()
                    try:
                        # Build UPDATE statement
                        update_parts = []
                        update_values = []
                        
                        if 'status' in data:
                            update_parts.append('status = ?')
                            update_values.append(data['status'])
                        
                        if 'invoice_date' in data:
                            from datetime import datetime
                            if isinstance(data['invoice_date'], str):
                                invoice_date = datetime.strptime(data['invoice_date'], '%Y-%m-%d').date()
                            else:
                                invoice_date = data['invoice_date']
                            update_parts.append('invoice_date = ?')
                            update_values.append(invoice_date)
                        
                        if 'due_date' in data:
                            from datetime import datetime
                            if isinstance(data['due_date'], str):
                                due_date = datetime.strptime(data['due_date'], '%Y-%m-%d').date()
                            else:
                                due_date = data['due_date']
                            update_parts.append('due_date = ?')
                            update_values.append(due_date)
                        
                        if 'notes' in data:
                            update_parts.append('notes = ?')
                            update_values.append(data['notes'])
                        
                        if 'subtotal' in data:
                            update_parts.append('subtotal = ?')
                            update_values.append(float(data['subtotal']))
                        
                        if has_freight and 'freight' in data:
                            update_parts.append('freight = ?')
                            update_values.append(float(data['freight']))
                        
                        if has_tax and 'tax' in data:
                            update_parts.append('tax = ?')
                            update_values.append(float(data['tax']))
                        elif 'tax_amount' in data and 'tax_amount' in columns:
                            update_parts.append('tax_amount = ?')
                            update_values.append(float(data['tax_amount']))
                        
                        if has_discount and 'discount' in data:
                            update_parts.append('discount = ?')
                            update_values.append(float(data['discount']))
                        
                        if has_grand_total and 'grand_total' in data:
                            update_parts.append('grand_total = ?')
                            update_values.append(float(data['grand_total']))
                        elif 'total_amount' in data and 'total_amount' in columns:
                            update_parts.append('total_amount = ?')
                            update_values.append(float(data['total_amount']))
                        
                        if 'customer_vendor_name' in data and 'customer_vendor_name' in columns:
                            update_parts.append('customer_vendor_name = ?')
                            update_values.append(data['customer_vendor_name'])
                        
                        if 'customer_vendor_id' in data and 'customer_vendor_id' in columns:
                            update_parts.append('customer_vendor_id = ?')
                            update_values.append(data['customer_vendor_id'])
                        
                        if update_parts:
                            update_values.append(invoice_id)  # WHERE clause
                            update_sql = f"UPDATE erp_core_invoice SET {', '.join(update_parts)} WHERE id = ?"
                            raw_cursor.execute(update_sql, update_values)
                            raw_conn.commit()
                        else:
                            # No fields to update
                            pass
                        
                        # Get updated invoice
                        raw_cursor.execute("SELECT * FROM erp_core_invoice WHERE id = ?", [invoice_id])
                        row = raw_cursor.fetchone()
                    finally:
                        raw_conn.close()
                    
                    if not row:
                        return Response({'error': 'Invoice not found'}, status=http_status.HTTP_404_NOT_FOUND)
                    
                    # Get column names and create invoice dict
                    invoice_dict = dict(zip(column_names, row))
                    
                    # Map database fields to model fields
                    mapped_dict = {}
                    for key, value in invoice_dict.items():
                        if key == 'tax_amount':
                            mapped_dict['tax'] = value
                        elif key == 'total_amount':
                            mapped_dict['grand_total'] = value
                        else:
                            mapped_dict[key] = value
                    
                    # Set defaults
                    if 'freight' not in mapped_dict:
                        mapped_dict['freight'] = 0.0
                    if 'tax' not in mapped_dict:
                        mapped_dict['tax'] = mapped_dict.get('tax_amount', 0.0)
                    if 'discount' not in mapped_dict:
                        mapped_dict['discount'] = 0.0
                    if 'grand_total' not in mapped_dict:
                        mapped_dict['grand_total'] = mapped_dict.get('total_amount', 0.0)
                    
                    # Create Invoice object
                    invoice = Invoice()
                    if 'id' in mapped_dict:
                        invoice.id = mapped_dict['id']
                    if 'invoice_number' in mapped_dict:
                        invoice.invoice_number = mapped_dict['invoice_number']
                    if 'invoice_date' in mapped_dict:
                        invoice.invoice_date = mapped_dict['invoice_date']
                    if 'due_date' in mapped_dict:
                        invoice.due_date = mapped_dict['due_date']
                    if 'status' in mapped_dict:
                        invoice.status = mapped_dict['status']
                    if 'subtotal' in mapped_dict:
                        invoice.subtotal = mapped_dict['subtotal']
                    if 'notes' in mapped_dict:
                        invoice.notes = mapped_dict['notes']
                    if 'created_at' in mapped_dict:
                        invoice.created_at = mapped_dict['created_at']
                    if 'updated_at' in mapped_dict:
                        invoice.updated_at = mapped_dict['updated_at']
                    
                    invoice.tax = mapped_dict.get('tax', 0.0)
                    invoice.grand_total = mapped_dict.get('grand_total', 0.0)
                    invoice.freight = mapped_dict.get('freight', 0.0)
                    invoice.discount = mapped_dict.get('discount', 0.0)
                    
                    if 'sales_order_id' in mapped_dict and mapped_dict['sales_order_id']:
                        invoice.sales_order_id = mapped_dict['sales_order_id']
                        invoice._sales_order_id = mapped_dict['sales_order_id']
                    
                    invoice._state.adding = False
                    invoice._state.db = connection
                    
                    # Check if status was changed to 'sent' and send email
                    if 'status' in data and data['status'] == 'sent':
                        try:
                            # Reload invoice with items to generate PDF
                            invoice = Invoice.objects.get(id=invoice_id)
                            pdf_content = generate_invoice_pdf(invoice)
                            send_invoice_email(invoice, pdf_content)
                        except Exception as e:
                            import logging
                            logger = logging.getLogger(__name__)
                            logger.error(f"Failed to send invoice email: {str(e)}")
                            # Don't fail the request if email fails
                    
                    # Prefetch items and sales order using raw SQL to avoid schema issues
                    from .models import InvoiceItem, SalesOrder
                    try:
                        # Try to get items using raw SQL to avoid schema mismatches
                        # Use raw sqlite3 to avoid Django's query logging issues
                        import sqlite3
                        from django.conf import settings
                        db_path = settings.DATABASES['default']['NAME']
                        raw_conn = sqlite3.connect(db_path)
                        raw_cursor = raw_conn.cursor()
                        
                        item_columns = []
                        item_rows = []
                        try:
                            raw_cursor.execute("PRAGMA table_info(erp_core_invoiceitem)")
                            item_columns = [row[1] for row in raw_cursor.fetchall()]
                            
                            if 'invoice_id' in item_columns:
                                # Use raw SQL to fetch items
                                raw_cursor.execute("SELECT * FROM erp_core_invoiceitem WHERE invoice_id = ?", [invoice.id])
                                item_rows = raw_cursor.fetchall()
                        finally:
                            raw_conn.close()
                        
                        if 'invoice_id' in item_columns and item_rows:
                            items = []
                            for item_row in item_rows:
                                item_dict = dict(zip(item_columns, item_row))
                                item = InvoiceItem()
                                if 'id' in item_dict:
                                    item.id = item_dict['id']
                                if 'invoice_id' in item_dict:
                                    item.invoice_id = item_dict['invoice_id']
                                if 'description' in item_dict:
                                    item.description = item_dict['description']
                                if 'quantity' in item_dict:
                                    item.quantity = item_dict['quantity']
                                if 'unit_price' in item_dict:
                                    item.unit_price = item_dict['unit_price']
                                if 'total' in item_dict:
                                    item.total = item_dict['total']
                                elif 'line_total' in item_dict:
                                    item.total = item_dict['line_total']
                                if 'item_id' in item_dict:
                                    item.item_id = item_dict['item_id']
                                if 'sales_order_item_id' in item_dict:
                                    item.sales_order_item_id = item_dict['sales_order_item_id']
                                item._state.adding = False
                                item._state.db = connection
                                items.append(item)
                            
                            invoice._prefetched_objects_cache = {'items': items}
                        else:
                            invoice._prefetched_objects_cache = {'items': []}
                    except Exception as item_error:
                        logger.warning(f"Error loading invoice items: {item_error}")
                        invoice._prefetched_objects_cache = {'items': []}
                    
                    if invoice.sales_order_id:
                        try:
                            invoice.sales_order = SalesOrder.objects.select_related('customer', 'ship_to_location').get(id=invoice.sales_order_id)
                        except SalesOrder.DoesNotExist:
                            invoice.sales_order = None
                    
                    # Serialize and return
                    serializer = self.get_serializer(invoice)
                    return Response(serializer.data)
                else:
                    # Schema matches - use ORM
                    try:
                        instance = Invoice.objects.get(id=invoice_id)
                    except Invoice.DoesNotExist:
                        return Response({'error': 'Invoice not found'}, status=http_status.HTTP_404_NOT_FOUND)
                    
                    # Update fields that are provided
                    update_fields = []
                    
                    if 'status' in data:
                        old_status = instance.status
                        instance.status = data['status']
                        update_fields.append('status')
                        
                        # Send email if status changed to 'sent'
                        if old_status != 'sent' and data['status'] == 'sent':
                            try:
                                pdf_content = generate_invoice_pdf(instance)
                                send_invoice_email(instance, pdf_content)
                            except Exception as e:
                                import logging
                                logger = logging.getLogger(__name__)
                                logger.error(f"Failed to send invoice email: {str(e)}")
                                # Don't fail the request if email fails
                    
                    if 'invoice_date' in data:
                        from datetime import datetime
                        if isinstance(data['invoice_date'], str):
                            instance.invoice_date = datetime.strptime(data['invoice_date'], '%Y-%m-%d').date()
                        else:
                            instance.invoice_date = data['invoice_date']
                        update_fields.append('invoice_date')
                    
                    if 'due_date' in data:
                        from datetime import datetime
                        if isinstance(data['due_date'], str):
                            instance.due_date = datetime.strptime(data['due_date'], '%Y-%m-%d').date()
                        else:
                            instance.due_date = data['due_date']
                        update_fields.append('due_date')
                    
                    if 'notes' in data:
                        instance.notes = data['notes']
                        update_fields.append('notes')
                    
                    if 'subtotal' in data:
                        instance.subtotal = float(data['subtotal'])
                        update_fields.append('subtotal')
                    
                    if has_freight and 'freight' in data:
                        instance.freight = float(data['freight'])
                        update_fields.append('freight')
                    
                    if has_tax and 'tax' in data:
                        instance.tax = float(data['tax'])
                        update_fields.append('tax')
                    
                    if has_discount and 'discount' in data:
                        instance.discount = float(data['discount'])
                        update_fields.append('discount')
                    
                    if has_grand_total and 'grand_total' in data:
                        instance.grand_total = float(data['grand_total'])
                        update_fields.append('grand_total')
                    
                    # Save the instance
                    if update_fields:
                        instance.save(update_fields=update_fields)
                    
                    # Return updated invoice
                    serializer = self.get_serializer(instance)
                    return Response(serializer.data)
            
        except Exception as e:
            logger.error(f"Error updating invoice: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return Response(
                {'error': f'Failed to update invoice: {str(e)}'},
                status=http_status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=True, methods=['get'], url_path='pdf', url_name='pdf')
    def generate_pdf(self, request, pk=None):
        """Generate a PDF invoice from template"""
        from django.http import HttpResponse
        from pathlib import Path
        import os
        import tempfile
        from rest_framework.response import Response
        from rest_framework import status as http_status
        
        # Get invoice using retrieve logic to handle schema mismatches
        invoice = None
        try:
            # Try to get invoice using the same logic as retrieve
            from django.db import connection
            
            with connection.cursor() as cursor:
                cursor.execute("PRAGMA table_info(erp_core_invoice)")
                column_info = cursor.fetchall()
                columns = {col[1]: col for col in column_info}
                has_freight = 'freight' in columns
                has_tax = 'tax' in columns
                has_grand_total = 'grand_total' in columns
                has_discount = 'discount' in columns
                
                if not has_freight or not has_tax or not has_grand_total or not has_discount:
                    # Use raw SQL - use raw sqlite3 to avoid Django's query logging issues
                    import sqlite3
                    from django.conf import settings
                    db_path = settings.DATABASES['default']['NAME']
                    raw_conn = sqlite3.connect(db_path)
                    raw_cursor = raw_conn.cursor()
                    try:
                        raw_cursor.execute("SELECT * FROM erp_core_invoice WHERE id = ?", [pk])
                        row = raw_cursor.fetchone()
                    finally:
                        raw_conn.close()
                    if not row:
                        from rest_framework.response import Response
                        from rest_framework import status as http_status
                        return Response({'error': 'Invoice not found'}, status=http_status.HTTP_404_NOT_FOUND)
                    
                    column_names = [col[1] for col in column_info]
                    invoice_dict = dict(zip(column_names, row))
                    
                    # Map database fields to model fields
                    mapped_dict = {}
                    for key, value in invoice_dict.items():
                        if key == 'tax_amount':
                            mapped_dict['tax'] = value
                        elif key == 'total_amount':
                            mapped_dict['grand_total'] = value
                        else:
                            mapped_dict[key] = value
                    
                    # Set defaults
                    if 'freight' not in mapped_dict:
                        mapped_dict['freight'] = 0.0
                    if 'tax' not in mapped_dict:
                        mapped_dict['tax'] = mapped_dict.get('tax_amount', 0.0)
                    if 'discount' not in mapped_dict:
                        mapped_dict['discount'] = 0.0
                    if 'grand_total' not in mapped_dict:
                        mapped_dict['grand_total'] = mapped_dict.get('total_amount', 0.0)
                    
                    # Create Invoice object
                    invoice = Invoice()
                    if 'id' in mapped_dict:
                        invoice.id = mapped_dict['id']
                    if 'invoice_number' in mapped_dict:
                        invoice.invoice_number = mapped_dict['invoice_number']
                    if 'invoice_date' in mapped_dict:
                        invoice.invoice_date = mapped_dict['invoice_date']
                    if 'due_date' in mapped_dict:
                        invoice.due_date = mapped_dict['due_date']
                    if 'status' in mapped_dict:
                        invoice.status = mapped_dict['status']
                    if 'subtotal' in mapped_dict:
                        invoice.subtotal = mapped_dict['subtotal']
                    if 'notes' in mapped_dict:
                        invoice.notes = mapped_dict['notes']
                    
                    invoice.tax = mapped_dict.get('tax', 0.0)
                    invoice.grand_total = mapped_dict.get('grand_total', 0.0)
                    invoice.freight = mapped_dict.get('freight', 0.0)
                    invoice.discount = mapped_dict.get('discount', 0.0)
                    
                    if 'sales_order_id' in mapped_dict and mapped_dict['sales_order_id']:
                        invoice.sales_order_id = mapped_dict['sales_order_id']
                    
                    invoice._state.adding = False
                    invoice._state.db = connection
                    
                    # Get invoice items - use raw SQL to avoid schema issues
                    from .models import InvoiceItem
                    try:
                        items = list(InvoiceItem.objects.filter(invoice_id=invoice.id))
                    except Exception:
                        # If ORM fails, use raw SQL
                        raw_conn = sqlite3.connect(db_path)
                        raw_cursor = raw_conn.cursor()
                        try:
                            raw_cursor.execute("SELECT * FROM erp_core_invoiceitem WHERE invoice_id = ?", [invoice.id])
                            item_rows = raw_cursor.fetchall()
                            raw_cursor.execute("PRAGMA table_info(erp_core_invoiceitem)")
                            item_columns = [col[1] for col in raw_cursor.fetchall()]
                            items = []
                            for item_row in item_rows:
                                item_dict = dict(zip(item_columns, item_row))
                                item = InvoiceItem()
                                item.id = item_dict.get('id')
                                item.invoice_id = invoice.id
                                item.description = item_dict.get('description', '')
                                item.quantity = item_dict.get('quantity', 0.0)
                                item.unit_price = item_dict.get('unit_price', 0.0)
                                item.total = item_dict.get('total', item_dict.get('line_total', 0.0))
                                if 'item_id' in item_dict:
                                    item.item_id = item_dict.get('item_id')
                                if 'sales_order_item_id' in item_dict:
                                    item.sales_order_item_id = item_dict.get('sales_order_item_id')
                                items.append(item)
                        finally:
                            raw_conn.close()
                    invoice._prefetched_objects_cache = {'items': items}
                    
                    # Try to get sales order if it exists
                    if invoice.sales_order_id:
                        try:
                            from .models import SalesOrder
                            invoice.sales_order = SalesOrder.objects.select_related('customer', 'ship_to_location').get(id=invoice.sales_order_id)
                        except SalesOrder.DoesNotExist:
                            invoice.sales_order = None
                else:
                    # Use normal ORM
                    invoice = self.get_object()
                    # Ensure items are prefetched
                    if not hasattr(invoice, '_prefetched_objects_cache'):
                        from .models import InvoiceItem
                        try:
                            items = list(InvoiceItem.objects.filter(invoice_id=invoice.id))
                        except Exception:
                            # If ORM fails, use raw SQL
                            import sqlite3
                            from django.conf import settings
                            db_path = settings.DATABASES['default']['NAME']
                            raw_conn = sqlite3.connect(db_path)
                            raw_cursor = raw_conn.cursor()
                            try:
                                raw_cursor.execute("SELECT * FROM erp_core_invoiceitem WHERE invoice_id = ?", [invoice.id])
                                item_rows = raw_cursor.fetchall()
                                raw_cursor.execute("PRAGMA table_info(erp_core_invoiceitem)")
                                item_columns = [col[1] for col in raw_cursor.fetchall()]
                                items = []
                                for item_row in item_rows:
                                    item_dict = dict(zip(item_columns, item_row))
                                    item = InvoiceItem()
                                    item.id = item_dict.get('id')
                                    item.invoice_id = invoice.id
                                    item.description = item_dict.get('description', '')
                                    item.quantity = item_dict.get('quantity', 0.0)
                                    item.unit_price = item_dict.get('unit_price', 0.0)
                                    item.total = item_dict.get('total', item_dict.get('line_total', 0.0))
                                    if 'item_id' in item_dict:
                                        item.item_id = item_dict.get('item_id')
                                    if 'sales_order_item_id' in item_dict:
                                        item.sales_order_item_id = item_dict.get('sales_order_item_id')
                                    items.append(item)
                            finally:
                                raw_conn.close()
                        invoice._prefetched_objects_cache = {'items': items}
        except Exception as e:
            import logging
            import traceback
            logger = logging.getLogger(__name__)
            logger.error(f"Error getting invoice for PDF: {e}")
            logger.error(traceback.format_exc())
            from rest_framework.response import Response
            from rest_framework import status as http_status
            return Response({'error': f'Error retrieving invoice: {str(e)}'}, status=http_status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        if not invoice:
            return Response({'error': 'Invoice not found'}, status=http_status.HTTP_404_NOT_FOUND)
        
        try:
            # Try HTML→PDF first; fallback to ReportLab.
            pdf_content = None
            try:
                from .invoice_pdf_html import generate_invoice_pdf_from_html
                pdf_content = generate_invoice_pdf_from_html(invoice)
            except Exception:
                pass
            if not pdf_content:
                pdf_content = generate_invoice_pdf(invoice)
            response = HttpResponse(pdf_content, content_type='application/pdf')
            response['Content-Disposition'] = f'inline; filename="Invoice_{invoice.invoice_number}.pdf"'
            response['X-Content-Type-Options'] = 'nosniff'
            return response
        except Exception as e:
            import logging
            import traceback
            logger = logging.getLogger(__name__)
            logger.error(f'Failed to generate invoice PDF: {str(e)}')
            logger.error(traceback.format_exc())
            return Response(
                {'error': f'Failed to generate invoice PDF: {str(e)}'},
                status=http_status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        
class CalendarEventsViewSet(viewsets.ViewSet):
    """Calendar events for shipments, raw materials, and production"""
    
    @action(detail=False, methods=['get'])
    def events(self, request):
        from datetime import datetime, timedelta
        
        start_date_str = request.query_params.get('start_date')
        end_date_str = request.query_params.get('end_date')
        event_types = request.query_params.get('event_types', 'shipments,raw_materials,production,receivables,payables').split(',')
        
        # Parse dates
        start_date = None
        end_date = None
        if start_date_str:
            try:
                start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
            except ValueError:
                pass
        if end_date_str:
            try:
                end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
            except ValueError:
                pass
        
        events = []
        
        # Shipment events - expected ship dates from sales orders (all non-cancelled statuses)
        if 'shipments' in event_types:
            start_dt = timezone.make_aware(datetime.combine(start_date, datetime.min.time())) if start_date else None
            end_dt = timezone.make_aware(datetime.combine(end_date, datetime.max.time())) if end_date else None
            shipments = SalesOrder.objects.filter(
                status__in=['draft', 'allocated', 'ready_for_shipment', 'issued', 'shipped', 'received', 'completed']
            ).exclude(status='cancelled').exclude(expected_ship_date__isnull=True)
            if start_dt:
                shipments = shipments.filter(
                    models.Q(actual_ship_date__gte=start_dt) | models.Q(expected_ship_date__gte=start_dt)
                )
            if end_dt:
                shipments = shipments.filter(
                    models.Q(actual_ship_date__lte=end_dt) | models.Q(expected_ship_date__lte=end_dt)
                )
            
            for so in shipments:
                # Use expected_ship_date for planning, actual_ship_date for completed
                ship_date = so.expected_ship_date.date() if so.expected_ship_date else (so.actual_ship_date.date() if so.actual_ship_date else None)
                if ship_date:
                    # Determine if this is an actual shipment or planned
                    is_actual = so.status in ['shipped', 'completed'] and so.actual_ship_date is not None
                    needs_checkout = so.status in ['ready_for_shipment']
                    needs_allocation = so.status in ['draft', 'allocated']
                    
                    if needs_checkout:
                        title = f'Check Out & Ship: {so.so_number}'
                    elif needs_allocation:
                        title = f'Allocate & Ship: {so.so_number}'
                    elif is_actual:
                        title = f'Ship: {so.so_number}'
                    else:
                        title = f'Ship (Expected): {so.so_number}'
                    
                    events.append({
                        'id': f'shipment_{so.id}',
                        'type': 'shipment',
                        'title': title,
                        'date': ship_date.isoformat(),
                        'sales_order_id': so.id,
                        'sales_order_number': so.so_number,
                        'customer_name': so.customer_name,
                        'status': so.status,
                        'is_actual': is_actual,
                        'needs_checkout': needs_checkout,
                        'needs_allocation': needs_allocation,
                    })
        
        # Raw material receipt events
        if 'raw_materials' in event_types:
            from .models import PurchaseOrder
            raw_materials = PurchaseOrder.objects.filter(
                status__in=['issued', 'received', 'completed']
            ).exclude(status='cancelled')
            
            if start_date:
                raw_materials = raw_materials.filter(
                    models.Q(expected_delivery_date__gte=start_date) if hasattr(PurchaseOrder, 'expected_delivery_date') else models.Q()
                )
            if end_date:
                raw_materials = raw_materials.filter(
                    models.Q(expected_delivery_date__lte=end_date) if hasattr(PurchaseOrder, 'expected_delivery_date') else models.Q()
                )
            
            # Also get from lots
            lots = Lot.objects.filter(status='accepted')
            if start_date:
                lots = lots.filter(received_date__gte=timezone.make_aware(datetime.combine(start_date, datetime.min.time())))
            if end_date:
                lots = lots.filter(received_date__lte=timezone.make_aware(datetime.combine(end_date, datetime.max.time())))
            
            for lot in lots:
                events.append({
                    'id': f'raw_material_{lot.id}',
                    'type': 'raw_material',
                    'title': f'Receipt: {lot.lot_number}',
                    'date': lot.received_date.date().isoformat(),
                    'lot_id': lot.id,
                    'lot_number': lot.lot_number,
                    'item_name': lot.item.name,
                    'po_number': lot.po_number,
                })
        
        # Production events - include scheduled and in-progress batches
        if 'production' in event_types:
            batches = ProductionBatch.objects.filter(
                status__in=['scheduled', 'in_progress', 'closed']
            ).select_related('finished_good_item')
            
            if start_date:
                batches = batches.filter(production_date__gte=timezone.make_aware(datetime.combine(start_date, datetime.min.time())))
            if end_date:
                batches = batches.filter(production_date__lte=timezone.make_aware(datetime.combine(end_date, datetime.max.time())))
            
            for batch in batches:
                fg_name = batch.finished_good_item.name if batch.finished_good_item else ''
                title = f'{fg_name}: {batch.batch_number}' if fg_name else batch.batch_number
                
                events.append({
                    'id': f'production_{batch.id}',
                    'type': 'production',
                    'title': title,
                    'date': batch.production_date.date().isoformat(),
                    'batch_id': batch.id,
                    'batch_number': batch.batch_number,
                    'status': batch.status,
                    'quantity_produced': float(batch.quantity_produced),
                    'is_scheduled': batch.status == 'scheduled',
                })
        
        # Receivables (AR) - due dates from customer payment terms (CRM); open/partial only
        if 'receivables' in event_types:
            ar_queryset = AccountsReceivable.objects.filter(
                status__in=['open', 'partial']
            ).exclude(due_date__isnull=True)
            if start_date:
                ar_queryset = ar_queryset.filter(due_date__gte=start_date)
            if end_date:
                ar_queryset = ar_queryset.filter(due_date__lte=end_date)
            for ar in ar_queryset:
                events.append({
                    'id': f'receivable_{ar.id}',
                    'type': 'receivable',
                    'title': f'AR due: {ar.customer_name} - ${ar.balance:,.2f}',
                    'date': ar.due_date.isoformat(),
                    'ar_id': ar.id,
                    'customer_name': ar.customer_name,
                    'balance': float(ar.balance),
                    'invoice_id': ar.invoice_id,
                    'is_overdue': ar.due_date < timezone.now().date(),
                })
        
        # Payables (AP) - due dates from vendor payment terms (Quality > Vendor approval); open/partial only
        if 'payables' in event_types:
            ap_queryset = AccountsPayable.objects.filter(
                status__in=['open', 'partial']
            ).exclude(due_date__isnull=True)
            if start_date:
                ap_queryset = ap_queryset.filter(due_date__gte=start_date)
            if end_date:
                ap_queryset = ap_queryset.filter(due_date__lte=end_date)
            for ap in ap_queryset:
                events.append({
                    'id': f'payable_{ap.id}',
                    'type': 'payable',
                    'title': f'AP due: {ap.vendor_name} - ${ap.balance:,.2f}',
                    'date': ap.due_date.isoformat(),
                    'ap_id': ap.id,
                    'vendor_name': ap.vendor_name,
                    'balance': float(ap.balance),
                    'purchase_order_id': ap.purchase_order_id,
                    'is_overdue': ap.due_date < timezone.now().date(),
                })
        
        # Sort by date
        events.sort(key=lambda x: x['date'])
        
        return Response({'events': events})


class InventoryTransactionViewSet(viewsets.ModelViewSet):
    queryset = InventoryTransaction.objects.select_related('lot__item').all()
    serializer_class = InventoryTransactionSerializer


class CustomerViewSet(viewsets.ModelViewSet):
    queryset = Customer.objects.none()  # Default to empty queryset
    serializer_class = CustomerSerializer
    
    def get_queryset(self):
        from django.db import connection
        from django.db.utils import OperationalError
        
        try:
            # Check if table exists first
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_customer'")
                table_exists = cursor.fetchone() is not None
            
            if not table_exists:
                # Table doesn't exist - return empty queryset
                return Customer.objects.none()
            
            # Table exists - try to query
            queryset = Customer.objects.all()
            is_active = self.request.query_params.get('is_active', None)
            if is_active is not None:
                queryset = queryset.filter(is_active=is_active.lower() == 'true')
            return queryset
        except OperationalError:
            # Database error - table likely doesn't exist
            return Customer.objects.none()
        except Exception as e:
            # Any other error
            import traceback
            print(f"Error in CustomerViewSet.get_queryset: {e}")
            traceback.print_exc()
            return Customer.objects.none()
    
    def list(self, request, *args, **kwargs):
        from rest_framework.response import Response
        from django.db.utils import OperationalError
        from django.db import connection
        
        try:
            # First check if table exists
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_customer'")
                table_exists = cursor.fetchone() is not None
            
            if not table_exists:
                # Table doesn't exist - return empty list
                return Response([])
            
            # Table exists - try to get queryset
            queryset = self.get_queryset()
            
            # Try to serialize
            serializer = self.get_serializer(queryset, many=True)
            return Response(serializer.data)
            
        except OperationalError as e:
            # Database error - table likely doesn't exist
            import traceback
            print(f"OperationalError in CustomerViewSet.list: {e}")
            traceback.print_exc()
            return Response([])
        except Exception as e:
            # Any other error, return empty list
            import traceback
            print(f"Error in CustomerViewSet.list: {e}")
            traceback.print_exc()
            return Response([])
    
    def create(self, request, *args, **kwargs):
        from django.db import connection
        from django.db.utils import OperationalError
        
        # Check if table exists first
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_customer'")
                table_exists = cursor.fetchone() is not None
            
            if not table_exists:
                # Table doesn't exist - try to create it using raw SQL
                try:
                    self._create_customer_table()
                    # Also create other CRM tables
                    self._create_crm_tables()
                    # Verify table was created by checking again
                    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_customer'")
                    if not cursor.fetchone():
                        return Response(
                            {'error': 'Failed to create customer table. Please run migrations.'},
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR
                        )
                except Exception as e:
                    import traceback
                    traceback.print_exc()
                    return Response(
                        {'error': f'Failed to create customer table: {str(e)}. Please run migrations.'},
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR
                    )
        except Exception as e:
            return Response(
                {'error': f'Database error: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        
        # Make a mutable copy of request.data
        data = request.data.copy()
        
        # Auto-generate customer_id if not provided or empty
        if 'customer_id' not in data or not data.get('customer_id'):
            try:
                data['customer_id'] = generate_customer_id()
            except Exception as e:
                # Fallback: generate simple ID
                import traceback
                print(f"Error generating customer ID: {e}")
                traceback.print_exc()
                # Use timestamp-based ID as fallback
                from django.utils import timezone
                data['customer_id'] = f"{int(timezone.now().timestamp()) % 100000:05d}"
        
        # Create serializer with modified data
        try:
            serializer = self.get_serializer(data=data)
            serializer.is_valid(raise_exception=True)
            self.perform_create(serializer)
            headers = self.get_success_headers(serializer.data)
            return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)
        except OperationalError as e:
            return Response(
                {'error': f'Database error: {str(e)}. Please ensure migrations have been run.'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response(
                {'error': f'Error creating customer: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    def _create_customer_table(self):
        """Create customer table if it doesn't exist (workaround for migration issues)"""
        from django.db import connection
        from django.utils import timezone
        
        with connection.cursor() as cursor:
            # Create Customer table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS erp_core_customer (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    customer_id VARCHAR(100) UNIQUE NOT NULL,
                    name VARCHAR(255) NOT NULL,
                    contact_name VARCHAR(255),
                    email VARCHAR(254),
                    phone VARCHAR(50),
                    address TEXT,
                    city VARCHAR(100),
                    state VARCHAR(50),
                    zip_code VARCHAR(20),
                    country VARCHAR(100) DEFAULT 'USA',
                    payment_terms VARCHAR(50),
                    notes TEXT,
                    is_active BOOLEAN DEFAULT 1,
                    created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
                    updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP
                )
            """)
            # Create index
            try:
                cursor.execute("CREATE INDEX IF NOT EXISTS erp_core_customer_customer_id ON erp_core_customer(customer_id)")
            except Exception:
                pass  # Index might already exist
            
            # Also create CustomerNumberSequence table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS erp_core_customernumbersequence (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    sequence_number INTEGER NOT NULL DEFAULT 0,
                    last_updated DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # Insert initial sequence record if it doesn't exist
            try:
                cursor.execute("""
                    INSERT OR IGNORE INTO erp_core_customernumbersequence (id, sequence_number, last_updated)
                    VALUES (1, 0, CURRENT_TIMESTAMP)
                """)
            except Exception:
                pass  # Record might already exist
    
    def _create_crm_tables(self):
        """Create all CRM-related tables if they don't exist"""
        from django.db import connection
        
        with connection.cursor() as cursor:
            # CustomerPricing table (if it doesn't exist - it might already exist)
            try:
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS erp_core_customerpricing (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        customer_id INTEGER NOT NULL,
                        item_id INTEGER NOT NULL,
                        unit_price REAL NOT NULL,
                        unit_of_measure VARCHAR(10) NOT NULL DEFAULT 'lbs',
                        effective_date DATE NOT NULL,
                        expiry_date DATE,
                        is_active BOOLEAN DEFAULT 1,
                        notes TEXT,
                        created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
                        updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
                        UNIQUE(customer_id, item_id, effective_date)
                    )
                """)
            except Exception:
                pass
            
            # ShipToLocation table
            try:
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS erp_core_shiptolocation (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        customer_id INTEGER NOT NULL,
                        location_name VARCHAR(255) NOT NULL,
                        contact_name VARCHAR(255),
                        email VARCHAR(254),
                        phone VARCHAR(50),
                        address TEXT NOT NULL,
                        city VARCHAR(100) NOT NULL,
                        state VARCHAR(50),
                        zip_code VARCHAR(20) NOT NULL,
                        country VARCHAR(100) DEFAULT 'USA',
                        is_default BOOLEAN DEFAULT 0,
                        is_active BOOLEAN DEFAULT 1,
                        notes TEXT,
                        created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
                        updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP
                    )
                """)
            except Exception:
                pass
            
            # CustomerContact table
            try:
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS erp_core_customercontact (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        customer_id INTEGER NOT NULL,
                        first_name VARCHAR(100) NOT NULL,
                        last_name VARCHAR(100) NOT NULL,
                        title VARCHAR(100),
                        email VARCHAR(254),
                        phone VARCHAR(50),
                        mobile VARCHAR(50),
                        is_primary BOOLEAN DEFAULT 0,
                        is_active BOOLEAN DEFAULT 1,
                        notes TEXT,
                        created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
                        updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP
                    )
                """)
            except Exception:
                pass
            
            # SalesCall table
            try:
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS erp_core_salescall (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        customer_id INTEGER NOT NULL,
                        contact_id INTEGER,
                        call_date DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
                        call_type VARCHAR(20) NOT NULL DEFAULT 'phone',
                        subject VARCHAR(255),
                        notes TEXT NOT NULL,
                        follow_up_required BOOLEAN DEFAULT 0,
                        follow_up_date DATETIME,
                        created_by VARCHAR(255),
                        created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
                        updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP
                    )
                """)
            except Exception:
                pass
            
            # CustomerForecast table
            try:
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS erp_core_customerforecast (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        customer_id INTEGER NOT NULL,
                        item_id INTEGER NOT NULL,
                        forecast_period VARCHAR(20) NOT NULL,
                        forecast_quantity REAL NOT NULL,
                        unit_of_measure VARCHAR(10) NOT NULL DEFAULT 'lbs',
                        notes TEXT,
                        created_by VARCHAR(255),
                        created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
                        updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
                        UNIQUE(customer_id, item_id, forecast_period)
                    )
                """)
            except Exception:
                pass


class CustomerPricingViewSet(viewsets.ModelViewSet):
    queryset = CustomerPricing.objects.none()
    serializer_class = CustomerPricingSerializer
    
    def get_queryset(self):
        from django.db import connection
        from django.db.utils import OperationalError
        from django.utils import timezone
        from datetime import date
        
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_customerpricing'")
                if not cursor.fetchone():
                    return CustomerPricing.objects.none()
            
            queryset = CustomerPricing.objects.select_related('customer', 'item').all()
            customer_id = self.request.query_params.get('customer_id', None)
            if customer_id:
                queryset = queryset.filter(customer_id=customer_id)
            
            # Filter by is_active if provided
            is_active = self.request.query_params.get('is_active', None)
            if is_active is not None:
                queryset = queryset.filter(is_active=is_active.lower() == 'true')
            
            # Filter by effective_date and expiry_date to show only currently active pricing
            today = date.today()
            from django.db.models import Q
            queryset = queryset.filter(
                effective_date__lte=today,
                is_active=True
            ).filter(
                Q(expiry_date__isnull=True) | Q(expiry_date__gte=today)
            )
            
            return queryset
        except Exception:
            return CustomerPricing.objects.none()
    
    def list(self, request, *args, **kwargs):
        from rest_framework.response import Response
        from django.db import connection
        
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_customerpricing'")
                if not cursor.fetchone():
                    # Try to create table
                    try:
                        CustomerViewSet()._create_crm_tables()
                    except Exception:
                        pass
                    return Response([])
            
            return super().list(request, *args, **kwargs)
        except Exception:
            return Response([])
    
    def create(self, request, *args, **kwargs):
        from django.db import connection
        import logging
        logger = logging.getLogger(__name__)
        
        # Ensure table exists
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_customerpricing'")
                if not cursor.fetchone():
                    CustomerViewSet()._create_crm_tables()
        except Exception as e:
            logger.error(f"Error ensuring table exists: {e}")
        
        try:
            logger.info(f"Creating CustomerPricing with data: {request.data}")
            return super().create(request, *args, **kwargs)
        except Exception as e:
            logger.error(f"Error creating CustomerPricing: {e}", exc_info=True)
            raise


class VendorPricingViewSet(viewsets.ModelViewSet):
    queryset = VendorPricing.objects.select_related('item').all()
    serializer_class = VendorPricingSerializer
    
    def get_queryset(self):
        queryset = VendorPricing.objects.select_related('item').all()
        vendor_name = self.request.query_params.get('vendor_name', None)
        if vendor_name:
            queryset = queryset.filter(vendor_name__icontains=vendor_name)
        item_id = self.request.query_params.get('item_id', None)
        if item_id:
            queryset = queryset.filter(item_id=item_id)
        is_active = self.request.query_params.get('is_active', None)
        if is_active is not None:
            queryset = queryset.filter(is_active=is_active.lower() == 'true')
        return queryset


class ShipToLocationViewSet(viewsets.ModelViewSet):
    queryset = ShipToLocation.objects.none()
    serializer_class = ShipToLocationSerializer
    
    def get_queryset(self):
        from django.db import connection
        
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_shiptolocation'")
                if not cursor.fetchone():
                    return ShipToLocation.objects.none()
            
            queryset = ShipToLocation.objects.select_related('customer').all()
            customer_id = self.request.query_params.get('customer_id', None)
            if customer_id:
                queryset = queryset.filter(customer_id=customer_id)
            return queryset
        except Exception:
            return ShipToLocation.objects.none()
    
    def list(self, request, *args, **kwargs):
        from rest_framework.response import Response
        from django.db import connection
        
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_shiptolocation'")
                if not cursor.fetchone():
                    try:
                        CustomerViewSet()._create_crm_tables()
                    except Exception:
                        pass
                    return Response([])
            
            return super().list(request, *args, **kwargs)
        except Exception:
            return Response([])
    
    def create(self, request, *args, **kwargs):
        from django.db import connection
        
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_shiptolocation'")
                if not cursor.fetchone():
                    CustomerViewSet()._create_crm_tables()
        except Exception:
            pass
        
        return super().create(request, *args, **kwargs)


class CustomerContactViewSet(viewsets.ModelViewSet):
    queryset = CustomerContact.objects.none()
    serializer_class = CustomerContactSerializer
    
    def get_queryset(self):
        from django.db import connection
        
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_customercontact'")
                if not cursor.fetchone():
                    return CustomerContact.objects.none()
            
            queryset = CustomerContact.objects.select_related('customer').all()
            customer_id = self.request.query_params.get('customer_id', None)
            if customer_id:
                queryset = queryset.filter(customer_id=customer_id)
            return queryset
        except Exception:
            return CustomerContact.objects.none()
    
    def list(self, request, *args, **kwargs):
        from rest_framework.response import Response
        from django.db import connection
        
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_customercontact'")
                if not cursor.fetchone():
                    try:
                        CustomerViewSet()._create_crm_tables()
                    except Exception:
                        pass
                    return Response([])
            
            return super().list(request, *args, **kwargs)
        except Exception:
            return Response([])
    
    def create(self, request, *args, **kwargs):
        from django.db import connection
        
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_customercontact'")
                if not cursor.fetchone():
                    CustomerViewSet()._create_crm_tables()
        except Exception:
            pass
        
        return super().create(request, *args, **kwargs)


class SalesCallViewSet(viewsets.ModelViewSet):
    queryset = SalesCall.objects.none()
    serializer_class = SalesCallSerializer
    
    def get_queryset(self):
        from django.db import connection
        
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_salescall'")
                if not cursor.fetchone():
                    return SalesCall.objects.none()
            
            queryset = SalesCall.objects.select_related('customer', 'contact').all()
            customer_id = self.request.query_params.get('customer_id', None)
            if customer_id:
                queryset = queryset.filter(customer_id=customer_id)
            return queryset
        except Exception:
            return SalesCall.objects.none()
    
    def list(self, request, *args, **kwargs):
        from rest_framework.response import Response
        from django.db import connection
        
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_salescall'")
                if not cursor.fetchone():
                    try:
                        CustomerViewSet()._create_crm_tables()
                    except Exception:
                        pass
                    return Response([])
            
            return super().list(request, *args, **kwargs)
        except Exception:
            return Response([])
    
    def create(self, request, *args, **kwargs):
        from django.db import connection
        
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_salescall'")
                if not cursor.fetchone():
                    CustomerViewSet()._create_crm_tables()
        except Exception:
            pass
        
        return super().create(request, *args, **kwargs)


class CustomerForecastViewSet(viewsets.ModelViewSet):
    queryset = CustomerForecast.objects.none()
    serializer_class = CustomerForecastSerializer
    
    def get_queryset(self):
        from django.db import connection
        
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_customerforecast'")
                if not cursor.fetchone():
                    return CustomerForecast.objects.none()
            
            queryset = CustomerForecast.objects.select_related('customer', 'item').all()
            customer_id = self.request.query_params.get('customer_id', None)
            if customer_id:
                queryset = queryset.filter(customer_id=customer_id)
            return queryset
        except Exception:
            return CustomerForecast.objects.none()
    
    def list(self, request, *args, **kwargs):
        from rest_framework.response import Response
        from django.db import connection
        
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_customerforecast'")
                if not cursor.fetchone():
                    try:
                        CustomerViewSet()._create_crm_tables()
                    except Exception:
                        pass
                    return Response([])
            
            return super().list(request, *args, **kwargs)
        except Exception:
            return Response([])
    
    def create(self, request, *args, **kwargs):
        from django.db import connection
        
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_customerforecast'")
                if not cursor.fetchone():
                    CustomerViewSet()._create_crm_tables()
        except Exception:
            pass
        
        return super().create(request, *args, **kwargs)


class CustomerUsageViewSet(viewsets.ViewSet):
    """ViewSet for customer usage/volume data"""
    
    def list(self, request):
        """Get usage data for a specific customer"""
        from django.db.models import Sum, Q
        from django.db import connection
        from django.db.utils import OperationalError
        from datetime import datetime
        
        customer_id = request.query_params.get('customer_id', None)
        item_id = request.query_params.get('item_id', None)
        year = request.query_params.get('year', None)
        
        if not customer_id:
            return Response({'error': 'customer_id is required'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            # Check if Customer table exists
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_customer'")
                if not cursor.fetchone():
                    return Response({
                        'customer_id': int(customer_id),
                        'customer_name': 'Unknown',
                        'usage': []
                    })
            
            customer = Customer.objects.get(id=customer_id)
        except Customer.DoesNotExist:
            return Response({'error': 'Customer not found'}, status=status.HTTP_404_NOT_FOUND)
        except OperationalError:
            return Response({
                'customer_id': int(customer_id),
                'customer_name': 'Unknown',
                'usage': []
            })
        
        # Check if SalesOrder table exists
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_salesorder'")
                if not cursor.fetchone():
                    return Response({
                        'customer_id': customer.id,
                        'customer_name': customer.name,
                        'usage': []
                    })
                # Also check for SalesOrderItem table
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='erp_core_salesorderitem'")
                if not cursor.fetchone():
                    return Response({
                        'customer_id': customer.id,
                        'customer_name': customer.name,
                        'usage': []
                    })
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({
                'customer_id': customer.id,
                'customer_name': customer.name,
                'usage': []
            })
        
        # Get all sales orders for this customer
        try:
            sales_orders = SalesOrder.objects.filter(customer=customer, status__in=['allocated', 'shipped'])
        except (OperationalError, Exception) as e:
            import traceback
            traceback.print_exc()
            return Response({
                'customer_id': customer.id,
                'customer_name': customer.name,
                'usage': []
            })
        
        # Filter by year if provided
        if year:
            sales_orders = sales_orders.filter(order_date__year=year)
        
        # Get current year for YTD calculation
        current_year = datetime.now().year
        
        # Aggregate usage by item
        usage_data = {}
        
        try:
            for so in sales_orders:
                try:
                    for item in so.items.all():
                        try:
                            if item_id and item.item_id != int(item_id):
                                continue
                            
                            item_key = item.item_id
                            if item_key not in usage_data:
                                usage_data[item_key] = {
                                    'item_id': item.item_id,
                                    'item_sku': item.item.sku if item.item else None,
                                    'item_name': item.item.name if item.item else None,
                                    'total_quantity': 0.0,
                                    'ytd_quantity': 0.0,
                                    'order_count': 0,
                                    'ytd_order_count': 0,
                                }
                            
                            quantity = item.quantity or 0.0
                            usage_data[item_key]['total_quantity'] += quantity
                            usage_data[item_key]['order_count'] += 1
                            
                            # YTD calculation
                            if so.order_date.year == current_year:
                                usage_data[item_key]['ytd_quantity'] += quantity
                                usage_data[item_key]['ytd_order_count'] += 1
                        except Exception:
                            # Skip this item if there's an error
                            continue
                except Exception:
                    # Skip this sales order if there's an error accessing items
                    continue
        except Exception as e:
            # If there's any error in the aggregation, return empty usage
            import traceback
            traceback.print_exc()
            return Response({
                'customer_id': customer.id,
                'customer_name': customer.name,
                'usage': []
            })
        
        try:
            return Response({
                'customer_id': customer.id,
                'customer_name': customer.name,
                'usage': list(usage_data.values())
            })
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({
                'customer_id': customer.id,
                'customer_name': customer.name,
                'usage': []
            })
    
    @action(detail=False, methods=['get'], url_path='customer-usage')
    def customer_usage(self, request):
        """Get usage data for a specific customer (alias for list)"""
        return self.list(request)


class VendorViewSet(viewsets.ModelViewSet):
    queryset = Vendor.objects.prefetch_related('history', 'documents', 'exceptions').all()
    serializer_class = VendorSerializer
    
    @action(detail=True, methods=['post'])
    def history(self, request, pk=None):
        vendor = self.get_object()
        serializer = VendorHistorySerializer(data={
            **request.data,
            'vendor': vendor.id
        })
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    @action(detail=True, methods=['post'])
    def approve(self, request, pk=None):
        vendor = self.get_object()
        vendor.approval_status = 'approved'
        vendor.approved_date = timezone.now()
        vendor.approved_by = request.data.get('approved_by', 'DOOF')
        vendor.save()
        return Response(VendorSerializer(vendor).data)
    
    @action(detail=True, methods=['get'])
    def items(self, request, pk=None):
        """Get approved items for this vendor with YTD usage"""
        vendor = self.get_object()
        from datetime import datetime
        from django.db.models import Sum
        
        # Get items for this vendor
        items = Item.objects.filter(vendor=vendor.name)
        
        # Calculate YTD usage from inventory transactions
        current_year = datetime.now().year
        items_data = []
        
        for item in items:
            # Get YTD usage from lots checked in this year
            ytd_usage = Lot.objects.filter(
                item=item,
                received_date__year=current_year
            ).aggregate(
                total=Sum('quantity')
            )['total'] or 0.0
            
            items_data.append({
                'id': item.id,
                'sku': item.sku,
                'name': item.name,
                'unit_of_measure': item.unit_of_measure,
                'pack_size': item.pack_size,
                'price': item.price,
                'ytd_usage': ytd_usage,
            })
        
        return Response(items_data)


class SupplierSurveyViewSet(viewsets.ModelViewSet):
    queryset = SupplierSurvey.objects.all()
    serializer_class = SupplierSurveySerializer
    
    def get_queryset(self):
        queryset = SupplierSurvey.objects.all()
        vendor_id = self.request.query_params.get('vendor', None)
        if vendor_id:
            queryset = queryset.filter(vendor_id=vendor_id)
        return queryset
    
    @action(detail=True, methods=['post'])
    def submit(self, request, pk=None):
        survey = self.get_object()
        survey.status = 'under_review'
        survey.submitted_date = timezone.now()
        survey.save()
        return Response(SupplierSurveySerializer(survey).data)
    
    @action(detail=True, methods=['post'])
    def approve(self, request, pk=None):
        survey = self.get_object()
        survey.status = 'approved'
        survey.approved_date = timezone.now()
        survey.approved_by = request.data.get('approved_by', 'DOOF')
        survey.reviewer_notes = request.data.get('reviewer_notes', '')
        survey.save()
        return Response(SupplierSurveySerializer(survey).data)


class SupplierDocumentViewSet(viewsets.ModelViewSet):
    queryset = SupplierDocument.objects.all()
    serializer_class = SupplierDocumentSerializer
    
    def get_queryset(self):
        queryset = SupplierDocument.objects.all()
        vendor_id = self.request.query_params.get('vendor', None)
        if vendor_id:
            queryset = queryset.filter(vendor_id=vendor_id)
        return queryset
    
    def create(self, request, *args, **kwargs):
        """Handle file upload for document creation"""
        # Handle multipart/form-data for file uploads
        if hasattr(request, 'FILES') and 'file' in request.FILES:
            file = request.FILES['file']
            # Read file content into binary field
            file_content = file.read()
            
            # Prepare data
            data = request.data.copy()
            data['file'] = file_content
            data['file_name'] = file.name
            data['file_size'] = file.size
            data['mime_type'] = file.content_type or 'application/pdf'
            
            serializer = self.get_serializer(data=data)
            serializer.is_valid(raise_exception=True)
            self.perform_create(serializer)
            headers = self.get_success_headers(serializer.data)
            return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)
        
        return super().create(request, *args, **kwargs)
    
    def update(self, request, *args, **kwargs):
        """Handle file upload for document replacement"""
        partial = kwargs.pop('partial', False)
        instance = self.get_object()
        
        # Handle multipart/form-data for file uploads
        if hasattr(request, 'FILES') and 'file' in request.FILES:
            file = request.FILES['file']
            # Read file content into binary field
            file_content = file.read()
            
            # Prepare data
            data = request.data.copy()
            data['file'] = file_content
            data['file_name'] = file.name
            data['file_size'] = file.size
            data['mime_type'] = file.content_type or 'application/pdf'
            
            serializer = self.get_serializer(instance, data=data, partial=partial)
            serializer.is_valid(raise_exception=True)
            self.perform_update(serializer)
            return Response(serializer.data)
        
        return super().update(request, *args, **kwargs)
    
    @action(detail=True, methods=['get'])
    def download(self, request, pk=None):
        """Download/view a supplier document"""
        from django.http import HttpResponse
        document = self.get_object()
        
        if not document.file:
            return Response(
                {'error': 'Document file not found'},
                status=status.HTTP_404_NOT_FOUND
            )
        
        response = HttpResponse(document.file, content_type=document.mime_type or 'application/pdf')
        response['Content-Disposition'] = f'inline; filename="{document.file_name or document.document_name}"'
        return response


class TemporaryExceptionViewSet(viewsets.ModelViewSet):
    queryset = TemporaryException.objects.all()
    serializer_class = TemporaryExceptionSerializer


class CostMasterViewSet(viewsets.ModelViewSet):
    queryset = CostMaster.objects.all()
    serializer_class = CostMasterSerializer
    
    def get_queryset(self):
        queryset = CostMaster.objects.all()
        product_code = self.request.query_params.get('product_code', None)
        vendor = self.request.query_params.get('vendor', None)
        if product_code:
            queryset = queryset.filter(wwi_product_code=product_code)
        if vendor:
            queryset = queryset.filter(vendor=vendor)
        return queryset
    
    @action(detail=False, methods=['get'])
    def pricing_history(self, request):
        """Get pricing history for multiple items by product codes"""
        from .serializers import CostMasterHistorySerializer
        product_codes = request.query_params.getlist('product_code')
        
        if not product_codes:
            return Response(
                {'error': 'product_code parameter is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Get CostMaster entries for these product codes
        cost_masters = CostMaster.objects.filter(wwi_product_code__in=product_codes)
        
        # Get all history records for these cost masters
        history = CostMasterHistory.objects.filter(
            cost_master__in=cost_masters
        ).select_related('cost_master').order_by('cost_master__wwi_product_code', '-effective_date')
        
        serializer = CostMasterHistorySerializer(history, many=True)
        return Response(serializer.data)
    
    @action(detail=True, methods=['get'])
    def history(self, request, pk=None):
        """Get pricing history for a specific CostMaster item"""
        from .serializers import CostMasterHistorySerializer
        cost_master = self.get_object()
        history = CostMasterHistory.objects.filter(cost_master=cost_master).order_by('-effective_date')
        serializer = CostMasterHistorySerializer(history, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['post'])
    def refresh_tariffs(self, request):
        """No-op: Flexport HTS integration was removed. Tariffs are entered manually in Cost Master."""
        return Response({
            'success': True,
            'updated': 0,
            'errors': 0,
            'message': 'Tariffs are entered manually. Flexport HTS integration has been removed.'
        }, status=status.HTTP_200_OK)


class AccountViewSet(viewsets.ModelViewSet):
    queryset = Account.objects.all()
    serializer_class = AccountSerializer


class FinishedProductSpecificationViewSet(viewsets.ModelViewSet):
    queryset = FinishedProductSpecification.objects.select_related('item').all()
    serializer_class = FinishedProductSpecificationSerializer
    
    def get_queryset(self):
        queryset = FinishedProductSpecification.objects.select_related('item').all()
        item_id = self.request.query_params.get('item', None)
        if item_id:
            queryset = queryset.filter(item_id=item_id)
        return queryset
    
    def get_serializer_context(self):
        context = super().get_serializer_context()
        context['request'] = self.request
        return context
    
    def create(self, request, *args, **kwargs):
        """Create FPS and automatically generate PDF"""
        # The serializer will handle item_id -> item mapping via PrimaryKeyRelatedField
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        fps = serializer.save()
        
        # Automatically generate PDF
        try:
            self._generate_and_save_pdf(fps)
        except Exception as e:
            import traceback
            print(f"Error generating PDF: {e}")
            traceback.print_exc()
            # Continue even if PDF generation fails
        
        headers = self.get_success_headers(serializer.data)
        return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)
    
    def update(self, request, *args, **kwargs):
        """Update FPS and regenerate PDF if needed"""
        partial = kwargs.pop('partial', False)
        instance = self.get_object()
        serializer = self.get_serializer(instance, data=request.data, partial=partial)
        serializer.is_valid(raise_exception=True)
        fps = serializer.save()
        
        # Regenerate PDF on update
        try:
            self._generate_and_save_pdf(fps)
        except Exception as e:
            import traceback
            print(f"Error regenerating PDF: {e}")
            traceback.print_exc()
        
        return Response(serializer.data)
    
    def _generate_and_save_pdf(self, fps):
        """Internal method to generate and save PDF using Word template"""
        import os
        import tempfile
        import re
        from pathlib import Path
        from docx import Document
        from io import BytesIO
        from django.core.files.base import ContentFile
        from django.conf import settings
        from .models import Formula
        
        item = fps.item
        
        # Get template path - template should be in the project root (one level up from backend_django)
        PROJECT_ROOT = Path(settings.BASE_DIR).parent.parent
        template_path = PROJECT_ROOT / 'Finished Product Specification Form (3.3-01).docx'
        
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found at {template_path}")
        
        # Load template
        doc = Document(str(template_path))
        
        # Get formula if exists
        formula_text = ''
        try:
            formula = Formula.objects.get(finished_good=item)
            formula_text = f"Version: {formula.version}\n\n"
            for ingredient in formula.ingredients.all():
                formula_text += f"{ingredient.item.name}: {ingredient.percentage}%\n"
        except Formula.DoesNotExist:
            formula_text = 'No formula assigned'
        
        # Define replacement mappings - these should match text patterns in the template
        replacements = {
            'Product Name': item.name,
            'Item number': item.sku,
            'Product Description': fps.product_description or '',
            'Color Specification': fps.color_specification or '',
            'pH': fps.ph or '',
            'Water Activity': fps.water_activity or '',
            'Microbiological Requirements': fps.microbiological_requirements or '',
            'Shelf life / Storage Requirements': fps.shelf_life_storage or '',
            'Type of Packaging': fps.packaging_type or '',
            'Additional Criteria': fps.additional_criteria or '',
            'Processing Requirements': fps.processing_requirements or '',
            'Product Formula': formula_text,
            'Name and Title of Person Completing Form': fps.completed_by_name or '',
            'Signature and Date': f"{fps.completed_by_signature or ''} {str(fps.completed_date) if fps.completed_date else ''}",
            'Test frequency': fps.test_frequency or '',
        }
        
        # Function to replace text in a paragraph (handles Word's internal structure)
        def replace_in_paragraph(paragraph, old_text, new_text):
            """Replace text in a paragraph, handling Word's run structure"""
            if old_text in paragraph.text:
                # Clear existing runs
                paragraph.clear()
                # Add new text
                run = paragraph.add_run(new_text)
                return True
            return False
        
        # Replace text in all paragraphs
        for paragraph in doc.paragraphs:
            full_text = paragraph.text
            for key, value in replacements.items():
                # Look for patterns like "Product Name:" or "Product Name" followed by empty space
                patterns = [
                    rf'{re.escape(key)}\s*:\s*\n',
                    rf'{re.escape(key)}\s*\n',
                    rf'{re.escape(key)}\s*$',
                ]
                for pattern in patterns:
                    if re.search(pattern, full_text, re.IGNORECASE):
                        # Find the position and replace
                        match = re.search(pattern, full_text, re.IGNORECASE)
                        if match:
                            # Replace the entire paragraph with the new content
                            new_text = full_text[:match.start()] + f'{key}: {value}'
                            paragraph.clear()
                            paragraph.add_run(new_text)
                            break
        
        # Handle checkboxes - look for checklist items and mark them
        checklist_items = {
            'MSDS Created': fps.msds_created,
            'Commercial Spec Created': fps.commercial_spec_created,
            'Label Template Created': fps.label_template_created,
            'Product evaluated for micro growth': fps.micro_growth_evaluated,
            'Product Added to Kosher Letter': fps.kosher_letter_added,
            'Initial HACCP Plan Created': fps.haccp_plan_created,
        }
        
        for paragraph in doc.paragraphs:
            full_text = paragraph.text
            for item_text, checked in checklist_items.items():
                if item_text in full_text:
                    # Replace checkbox symbols
                    checkmark = '☑' if checked else '☐'
                    new_text = re.sub(r'[☐☑]', checkmark, full_text, count=1)
                    if new_text != full_text:
                        paragraph.clear()
                        paragraph.add_run(new_text)
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            tmp_path = tmp_file.name
        
        try:
            # Convert DOCX to PDF using docx2pdf (requires Word or LibreOffice)
            try:
                from docx2pdf import convert
                pdf_path = tmp_path.replace('.docx', '.pdf')
                convert(tmp_path, pdf_path)
                
                # Read PDF content
                with open(pdf_path, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                
                # Clean up temp PDF file
                os.unlink(pdf_path)
            except Exception as e:
                # If conversion fails, fall back to using the DOCX file
                print(f"PDF conversion failed: {e}. Using DOCX file instead.")
                with open(tmp_path, 'rb') as docx_file:
                    pdf_content = docx_file.read()
            
            # Save PDF to model
            filename = f"FPS_{item.sku}_{item.name.replace(' ', '_').replace('/', '_')}.pdf"
            fps.fps_pdf.save(filename, ContentFile(pdf_content), save=True)
        finally:
            # Clean up temp DOCX file
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=14,
            textColor=colors.HexColor('#000000'),
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        field_label_style = ParagraphStyle(
            'FieldLabel',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#000000'),
            fontName='Helvetica-Bold',
            spaceAfter=4
        )
        
        field_value_style = ParagraphStyle(
            'FieldValue',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#000000'),
            spaceAfter=12
        )
        
        # Header
        header_data = [
            ['6431 Michels Drive', ''],
            ['Washington, MO 63090', ''],
            ['314-835-8207', '']
        ]
        header_table = Table(header_data, colWidths=[4*inch, 4*inch])
        header_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(header_table)
        story.append(Spacer(1, 0.2*inch))
        
        # Title
        story.append(Paragraph('Finished Product Specification', title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Test frequency
        if fps.test_frequency:
            story.append(Paragraph(f'<b>Test frequency:</b> {fps.test_frequency}', field_value_style))
            story.append(Spacer(1, 0.1*inch))
        
        # Main specification fields
        spec_fields = [
            ('Product Name', item.name),
            ('Item number', item.sku),
            ('Product Description<br/>(physical state, color, odor, etc.)', fps.product_description or ''),
            ('Color Specification<br/>(CV, dye %, color strength, etc.)', fps.color_specification or ''),
            ('pH', fps.ph or ''),
            ('Water Activity (aW)', fps.water_activity or ''),
            ('Microbiological Requirements<br/>(if micro testing not required, rationale must be provided)', fps.microbiological_requirements or ''),
            ('Shelf life / Storage Requirements<br/>(temperature data)<br/>Include Basis for decision and record Shelf-Life Assignment Form (Document No. 5.1.4–03)<br/>Shelf-Life Study Log (Document No. 5.1.4–02)', fps.shelf_life_storage or ''),
            ('Type of Packaging', fps.packaging_type or ''),
            ('Additional Criteria<br/>(physical parameter testing, flavor profile, customer considerations, etc.)', fps.additional_criteria or ''),
        ]
        
        for label, value in spec_fields:
            story.append(Paragraph(f'<b>{label}</b>', field_label_style))
            story.append(Paragraph(value or '', field_value_style))
        
        # Footer note
        story.append(Spacer(1, 0.2*inch))
        footer_note = 'This document contains confidential and proprietary information intended solely for the recipient. By accepting this document, you agree to maintain the confidentiality of its contents and not to disclose, distribute, or use any information herein for purposes other than those expressly authorized. Unauthorized use or disclosure may result in legal action. If you are not the intended recipient, please notify the sender immediately and delete this document from your system Finished Product Specification Form – Updated 12/10/2025 (GDM) – Reviewed by GDM – Effective Date 12/10/2025 – Doc No. 3.3 - 01.'
        story.append(Paragraph(footer_note, ParagraphStyle('Footer', parent=styles['Normal'], fontSize=7, textColor=colors.HexColor('#666666'))))
        
        story.append(PageBreak())
        
        # Checklist page
        story.append(header_table)
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph('FINISHED PRODUCT SPECIFICATION CHECKLIST', title_style))
        story.append(Spacer(1, 0.3*inch))
        
        story.append(Paragraph(f'<b>Product Name:</b> {item.name}', field_value_style))
        story.append(Spacer(1, 0.2*inch))
        
        checklist_items = [
            ('MSDS Created', fps.msds_created),
            ('Commercial Spec Created / COA', fps.commercial_spec_created),
            ('Label Template Created', fps.label_template_created),
            ('Product evaluated for micro growth', fps.micro_growth_evaluated),
            ('Product Added to Kosher Letter', fps.kosher_letter_added),
            ('Initial HACCP Plan Created', fps.haccp_plan_created),
        ]
        
        for label, checked in checklist_items:
            checkmark = '✓' if checked else '☐'
            story.append(Paragraph(f'{checkmark} {label}', field_value_style))
        
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph('<b>Processing Requirements</b><br/>(i.e. specific tank or mixer, how long should product be mixed, allergen considerations, temperature requirements)', field_label_style))
        story.append(Paragraph(fps.processing_requirements or '', field_value_style))
        
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph('<b>Product Formula</b>', field_label_style))
        # Get formula if exists
        try:
            formula = Formula.objects.get(finished_good=item)
            formula_text = f"Version: {formula.version}\n\n"
            for ingredient in formula.ingredients.all():
                formula_text += f"{ingredient.item.name}: {ingredient.percentage}%\n"
            story.append(Paragraph(formula_text.replace('\n', '<br/>'), field_value_style))
        except Formula.DoesNotExist:
            story.append(Paragraph('No formula assigned', field_value_style))
        
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(f'<b>Name and Title of Person Completing Form:</b><br/>{fps.completed_by_name or ""}', field_value_style))
        story.append(Spacer(1, 0.1*inch))
        story.append(Paragraph(f'<b>Signature and Date:</b><br/>{fps.completed_by_signature or ""} {fps.completed_date or ""}', field_value_style))
        
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(footer_note, ParagraphStyle('Footer', parent=styles['Normal'], fontSize=7, textColor=colors.HexColor('#666666'))))
        
        # Build PDF
        doc.build(story)
        
        # Get PDF content
        pdf_content = buffer.getvalue()
        buffer.close()
        
        # Save PDF to file
        filename = f"FPS_{item.sku}_{item.name.replace(' ', '_').replace('/', '_')}.pdf"
        fps.fps_pdf.save(filename, ContentFile(pdf_content), save=True)
    
    @action(detail=True, methods=['get'])
    def generate_pdf(self, request, pk=None):
        """Generate FPS PDF for a finished product specification"""
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from io import BytesIO
        from django.http import HttpResponse
        import os
        
        fps = self.get_object()
        item = fps.item
        
        # Create PDF in memory
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=14,
            textColor=colors.HexColor('#000000'),
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        header_style = ParagraphStyle(
            'Header',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#000000'),
            alignment=TA_LEFT
        )
        
        field_label_style = ParagraphStyle(
            'FieldLabel',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#000000'),
            fontName='Helvetica-Bold',
            spaceAfter=4
        )
        
        field_value_style = ParagraphStyle(
            'FieldValue',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#000000'),
            spaceAfter=12
        )
        
        # Header
        header_data = [
            ['6431 Michels Drive', ''],
            ['Washington, MO 63090', ''],
            ['314-835-8207', '']
        ]
        header_table = Table(header_data, colWidths=[4*inch, 4*inch])
        header_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(header_table)
        story.append(Spacer(1, 0.2*inch))
        
        # Title
        story.append(Paragraph('Finished Product Specification', title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Test frequency
        if fps.test_frequency:
            story.append(Paragraph(f'<b>Test frequency:</b> {fps.test_frequency}', field_value_style))
            story.append(Spacer(1, 0.1*inch))
        
        # Main specification fields
        spec_fields = [
            ('Product Name', item.name),
            ('Item number', item.sku),
            ('Product Description<br/>(physical state, color, odor, etc.)', fps.product_description or ''),
            ('Color Specification<br/>(CV, dye %, color strength, etc.)', fps.color_specification or ''),
            ('pH', fps.ph or ''),
            ('Water Activity (aW)', fps.water_activity or ''),
            ('Microbiological Requirements<br/>(if micro testing not required, rationale must be provided)', fps.microbiological_requirements or ''),
            ('Shelf life / Storage Requirements<br/>(temperature data)<br/>Include Basis for decision and record Shelf-Life Assignment Form (Document No. 5.1.4–03)<br/>Shelf-Life Study Log (Document No. 5.1.4–02)', fps.shelf_life_storage or ''),
            ('Type of Packaging', fps.packaging_type or ''),
            ('Additional Criteria<br/>(physical parameter testing, flavor profile, customer considerations, etc.)', fps.additional_criteria or ''),
        ]
        
        for label, value in spec_fields:
            story.append(Paragraph(f'<b>{label}</b>', field_label_style))
            story.append(Paragraph(value or '', field_value_style))
        
        # Footer note
        story.append(Spacer(1, 0.2*inch))
        footer_note = 'This document contains confidential and proprietary information intended solely for the recipient. By accepting this document, you agree to maintain the confidentiality of its contents and not to disclose, distribute, or use any information herein for purposes other than those expressly authorized. Unauthorized use or disclosure may result in legal action. If you are not the intended recipient, please notify the sender immediately and delete this document from your system Finished Product Specification Form – Updated 12/10/2025 (GDM) – Reviewed by GDM – Effective Date 12/10/2025 – Doc No. 3.3 - 01.'
        story.append(Paragraph(footer_note, ParagraphStyle('Footer', parent=styles['Normal'], fontSize=7, textColor=colors.HexColor('#666666'))))
        
        story.append(PageBreak())
        
        # Checklist page
        story.append(header_table)
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph('FINISHED PRODUCT SPECIFICATION CHECKLIST', title_style))
        story.append(Spacer(1, 0.3*inch))
        
        story.append(Paragraph(f'<b>Product Name:</b> {item.name}', field_value_style))
        story.append(Spacer(1, 0.2*inch))
        
        checklist_items = [
            ('MSDS Created', fps.msds_created),
            ('Commercial Spec Created / COA', fps.commercial_spec_created),
            ('Label Template Created', fps.label_template_created),
            ('Product evaluated for micro growth', fps.micro_growth_evaluated),
            ('Product Added to Kosher Letter', fps.kosher_letter_added),
            ('Initial HACCP Plan Created', fps.haccp_plan_created),
        ]
        
        for label, checked in checklist_items:
            checkmark = '✓' if checked else '☐'
            story.append(Paragraph(f'{checkmark} {label}', field_value_style))
        
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph('<b>Processing Requirements</b><br/>(i.e. specific tank or mixer, how long should product be mixed, allergen considerations, temperature requirements)', field_label_style))
        story.append(Paragraph(fps.processing_requirements or '', field_value_style))
        
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph('<b>Product Formula</b>', field_label_style))
        # Get formula if exists
        try:
            formula = Formula.objects.get(finished_good=item)
            formula_text = f"Version: {formula.version}\n\n"
            for ingredient in formula.ingredients.all():
                formula_text += f"{ingredient.item.name}: {ingredient.percentage}%\n"
            story.append(Paragraph(formula_text.replace('\n', '<br/>'), field_value_style))
        except Formula.DoesNotExist:
            story.append(Paragraph('No formula assigned', field_value_style))
        
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(f'<b>Name and Title of Person Completing Form:</b><br/>{fps.completed_by_name or ""}', field_value_style))
        story.append(Spacer(1, 0.1*inch))
        story.append(Paragraph(f'<b>Signature and Date:</b><br/>{fps.completed_by_signature or ""} {fps.completed_date or ""}', field_value_style))
        
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(footer_note, ParagraphStyle('Footer', parent=styles['Normal'], fontSize=7, textColor=colors.HexColor('#666666'))))
        
        # Build PDF
        doc.build(story)
        
        # Get PDF content
        pdf_content = buffer.getvalue()
        buffer.close()
        
        # Save PDF to file
        if not fps.fps_pdf:
            from django.core.files.base import ContentFile
            filename = f"FPS_{item.sku}_{item.name.replace(' ', '_')}.pdf"
            fps.fps_pdf.save(filename, ContentFile(pdf_content), save=True)
        
        # Return PDF as response
        response = HttpResponse(pdf_content, content_type='application/pdf')
        response['Content-Disposition'] = f'inline; filename="FPS_{item.sku}.pdf"'
        return response


class LotTransactionLogViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for viewing lot transaction logs"""
    queryset = LotTransactionLog.objects.select_related('lot', 'lot__item').all()
    serializer_class = LotTransactionLogSerializer
    
    def get_queryset(self):
        queryset = LotTransactionLog.objects.select_related('lot', 'lot__item').all()
        
        # Filter by lot_number
        lot_number = self.request.query_params.get('lot_number', None)
        if lot_number:
            queryset = queryset.filter(lot_number__icontains=lot_number)
        
        # Filter by item_sku
        sku = self.request.query_params.get('sku', None)
        if sku:
            queryset = queryset.filter(item_sku__icontains=sku)
        
        # Filter by transaction_type
        transaction_type = self.request.query_params.get('transaction_type', None)
        if transaction_type:
            queryset = queryset.filter(transaction_type=transaction_type)
        
        # Filter by reference_number
        reference_number = self.request.query_params.get('reference_number', None)
        if reference_number:
            queryset = queryset.filter(reference_number__icontains=reference_number)
        
        # Filter by date range
        date_from = self.request.query_params.get('date_from', None)
        if date_from:
            queryset = queryset.filter(logged_at__gte=date_from)
        
        date_to = self.request.query_params.get('date_to', None)
        if date_to:
            queryset = queryset.filter(logged_at__lte=date_to)
        
        return queryset.order_by('-logged_at')


class LotDepletionLogViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for viewing lot depletion logs"""
    queryset = LotDepletionLog.objects.all()
    serializer_class = LotDepletionLogSerializer
    
    def get_queryset(self):
        queryset = LotDepletionLog.objects.select_related('lot', 'lot__item').all()
        
        # Filter by lot number if provided
        lot_number = self.request.query_params.get('lot_number', None)
        if lot_number:
            queryset = queryset.filter(lot_number=lot_number)
        
        # Filter by SKU if provided
        sku = self.request.query_params.get('sku', None)
        if sku:
            queryset = queryset.filter(item_sku=sku)
        
        # Filter by depletion method if provided
        method = self.request.query_params.get('method', None)
        if method:
            queryset = queryset.filter(depletion_method=method)
        
        # Filter by date range if provided
        date_from = self.request.query_params.get('date_from', None)
        date_to = self.request.query_params.get('date_to', None)
        if date_from:
            from django.utils.dateparse import parse_datetime
            date_from_dt = parse_datetime(date_from)
            if date_from_dt:
                queryset = queryset.filter(depleted_at__gte=date_from_dt)
        if date_to:
            from django.utils.dateparse import parse_datetime
            date_to_dt = parse_datetime(date_to)
            if date_to_dt:
                queryset = queryset.filter(depleted_at__lte=date_to_dt)
        
        return queryset.order_by('-depleted_at')


class PurchaseOrderLogViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for viewing purchase order logs"""
    queryset = PurchaseOrderLog.objects.all()
    serializer_class = PurchaseOrderLogSerializer
    
    def get_queryset(self):
        queryset = PurchaseOrderLog.objects.select_related('purchase_order').all()
        
        # Filter by PO number if provided
        po_number = self.request.query_params.get('po_number', None)
        if po_number:
            queryset = queryset.filter(po_number=po_number)
        
        # Filter by vendor if provided
        vendor = self.request.query_params.get('vendor', None)
        if vendor:
            queryset = queryset.filter(vendor_name=vendor)
        
        # Filter by action if provided
        action = self.request.query_params.get('action', None)
        if action:
            queryset = queryset.filter(action=action)
        
        # Filter by lot number if provided
        lot_number = self.request.query_params.get('lot_number', None)
        if lot_number:
            queryset = queryset.filter(lot_number=lot_number)
        
        # Filter by date range if provided
        date_from = self.request.query_params.get('date_from', None)
        date_to = self.request.query_params.get('date_to', None)
        if date_from:
            from django.utils.dateparse import parse_datetime
            date_from_dt = parse_datetime(date_from)
            if date_from_dt:
                queryset = queryset.filter(logged_at__gte=date_from_dt)
        if date_to:
            from django.utils.dateparse import parse_datetime
            date_to_dt = parse_datetime(date_to)
            if date_to_dt:
                queryset = queryset.filter(logged_at__lte=date_to_dt)
        
        return queryset.order_by('-logged_at')


class FiscalPeriodViewSet(viewsets.ModelViewSet):
    """ViewSet for managing fiscal periods"""
    queryset = FiscalPeriod.objects.all()
    serializer_class = FiscalPeriodSerializer
    
    @action(detail=True, methods=['post'], url_path='close', url_name='close')
    def close_period(self, request, pk=None):
        """Mark the fiscal period as closed so no further posting is allowed."""
        period = self.get_object()
        if period.is_closed:
            return Response(
                {'detail': 'Period is already closed.', 'is_closed': True},
                status=status.HTTP_400_BAD_REQUEST
            )
        period.is_closed = True
        period.closed_date = timezone.now()
        period.closed_by = getattr(request.user, 'username', None) or 'system'
        period.save()
        return Response(self.get_serializer(period).data)
    
    def get_queryset(self):
        queryset = FiscalPeriod.objects.all()
        is_closed = self.request.query_params.get('is_closed', None)
        if is_closed is not None:
            is_closed_bool = is_closed.lower() == 'true'
            queryset = queryset.filter(is_closed=is_closed_bool)
        return queryset.order_by('-start_date')


def generate_journal_entry_number():
    """Generate a unique journal entry number in format JE-YYYYMMDD-001"""
    from django.db import transaction
    from datetime import date
    
    today = date.today()
    date_prefix = today.strftime('%Y%m%d')
    base_number = f"JE-{date_prefix}-"
    
    # Find the highest sequence number for today
    with transaction.atomic():
        existing_entries = JournalEntry.objects.filter(entry_number__startswith=base_number)
        max_sequence = 0
        
        for entry in existing_entries:
            try:
                # Extract sequence number from entry_number (e.g., "JE-20240115-001" -> 1)
                sequence_str = entry.entry_number.split('-')[-1]
                sequence = int(sequence_str)
                if sequence > max_sequence:
                    max_sequence = sequence
            except (ValueError, IndexError):
                pass
        
        next_sequence = max_sequence + 1
        entry_number = f"{base_number}{next_sequence:03d}"
        
        # Double-check uniqueness
        max_retries = 10
        retry_count = 0
        while JournalEntry.objects.filter(entry_number=entry_number).exists() and retry_count < max_retries:
            next_sequence += 1
            entry_number = f"{base_number}{next_sequence:03d}"
            retry_count += 1
    
    return entry_number


class JournalEntryViewSet(viewsets.ModelViewSet):
    """ViewSet for managing journal entries"""
    queryset = JournalEntry.objects.prefetch_related('lines', 'lines__account').all()
    serializer_class = JournalEntrySerializer
    
    def get_queryset(self):
        queryset = JournalEntry.objects.prefetch_related('lines', 'lines__account').all()
        
        # Filter by status
        status_filter = self.request.query_params.get('status', None)
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        
        # Filter by date range
        start_date = self.request.query_params.get('start_date', None)
        end_date = self.request.query_params.get('end_date', None)
        if start_date:
            queryset = queryset.filter(entry_date__gte=start_date)
        if end_date:
            queryset = queryset.filter(entry_date__lte=end_date)
        
        # Filter by reference
        reference_number = self.request.query_params.get('reference_number', None)
        if reference_number:
            queryset = queryset.filter(reference_number=reference_number)
        
        return queryset.order_by('-entry_date', '-created_at')
    
    def create(self, request, *args, **kwargs):
        """Create a journal entry with lines"""
        lines_data = request.data.pop('lines', [])
        
        # Validate that debits equal credits
        total_debits = sum(float(line.get('amount', 0)) for line in lines_data if line.get('debit_credit') == 'debit')
        total_credits = sum(float(line.get('amount', 0)) for line in lines_data if line.get('debit_credit') == 'credit')
        
        if abs(total_debits - total_credits) > 0.01:
            return Response(
                {'error': f'Journal entry must be balanced. Debits: ${total_debits:.2f}, Credits: ${total_credits:.2f}'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Generate entry number
        entry_number = generate_journal_entry_number()
        request.data['entry_number'] = entry_number
        
        # Determine fiscal period from entry_date
        from django.utils.dateparse import parse_date
        entry_date = parse_date(request.data.get('entry_date'))
        if entry_date:
            try:
                fiscal_period = FiscalPeriod.objects.filter(
                    start_date__lte=entry_date,
                    end_date__gte=entry_date
                ).first()
                if fiscal_period:
                    if fiscal_period.is_closed:
                        return Response(
                            {'error': 'Cannot create journal entry in a closed fiscal period.'},
                            status=status.HTTP_400_BAD_REQUEST
                        )
                    request.data['fiscal_period'] = fiscal_period.id
            except Exception:
                pass
        
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        journal_entry = serializer.save()
        
        # Create journal entry lines
        for line_data in lines_data:
            JournalEntryLine.objects.create(
                journal_entry=journal_entry,
                account_id=line_data['account'],
                debit_credit=line_data['debit_credit'],
                amount=float(line_data['amount']),
                description=line_data.get('description', '')
            )
        
        # Reload with lines
        journal_entry.refresh_from_db()
        return Response(self.get_serializer(journal_entry).data, status=status.HTTP_201_CREATED)
    
    @action(detail=True, methods=['post'], url_path='post', url_name='post')
    def post_entry(self, request, pk=None):
        """Post a journal entry to the general ledger"""
        journal_entry = self.get_object()
        
        if journal_entry.status == 'posted':
            return Response(
                {'error': 'Journal entry is already posted'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Validate balanced
        if not journal_entry.validate_balanced():
            return Response(
                {'error': 'Journal entry is not balanced. Cannot post.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        # Block posting to a closed period
        if journal_entry.fiscal_period and journal_entry.fiscal_period.is_closed:
            return Response(
                {'error': 'Cannot post to a closed fiscal period.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Create general ledger entries
        for line in journal_entry.lines.all():
            GeneralLedgerEntry.objects.create(
                journal_entry=journal_entry,
                account=line.account,
                fiscal_period=journal_entry.fiscal_period,
                entry_date=journal_entry.entry_date,
                description=line.description or journal_entry.description,
                debit=line.amount if line.debit_credit == 'debit' else 0.0,
                credit=line.amount if line.debit_credit == 'credit' else 0.0
            )
        
        # Update journal entry status
        journal_entry.status = 'posted'
        journal_entry.posted_by = request.user.username if hasattr(request.user, 'username') else 'system'
        journal_entry.posted_at = timezone.now()
        journal_entry.save()
        
        # Update account balances for the fiscal period
        if journal_entry.fiscal_period:
            for line in journal_entry.lines.all():
                account_balance, created = AccountBalance.objects.get_or_create(
                    account=line.account,
                    fiscal_period=journal_entry.fiscal_period,
                    defaults={
                        'opening_balance': 0.0,
                        'period_debits': 0.0,
                        'period_credits': 0.0,
                        'closing_balance': 0.0
                    }
                )
                
                if line.debit_credit == 'debit':
                    account_balance.period_debits += line.amount
                else:
                    account_balance.period_credits += line.amount
                
                # Calculate closing balance based on account type
                account_type = line.account.account_type
                if account_type in ['asset', 'expense']:
                    # Debits increase, credits decrease
                    account_balance.closing_balance = account_balance.opening_balance + account_balance.period_debits - account_balance.period_credits
                else:
                    # Credits increase, debits decrease
                    account_balance.closing_balance = account_balance.opening_balance + account_balance.period_credits - account_balance.period_debits
                
                account_balance.save()
        
        return Response(self.get_serializer(journal_entry).data)


class GeneralLedgerViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for viewing general ledger entries"""
    queryset = GeneralLedgerEntry.objects.select_related('account', 'journal_entry', 'fiscal_period').all()
    serializer_class = GeneralLedgerEntrySerializer
    
    def get_queryset(self):
        queryset = GeneralLedgerEntry.objects.select_related('account', 'journal_entry', 'fiscal_period').all()
        
        # Filter by account
        account_id = self.request.query_params.get('account_id', None)
        if account_id:
            queryset = queryset.filter(account_id=account_id)
        
        # Filter by date range
        start_date = self.request.query_params.get('start_date', None)
        end_date = self.request.query_params.get('end_date', None)
        if start_date:
            queryset = queryset.filter(entry_date__gte=start_date)
        if end_date:
            queryset = queryset.filter(entry_date__lte=end_date)
        
        # Filter by fiscal period
        fiscal_period_id = self.request.query_params.get('fiscal_period_id', None)
        if fiscal_period_id:
            queryset = queryset.filter(fiscal_period_id=fiscal_period_id)
        
        return queryset.order_by('entry_date', 'id')
    
    @action(detail=False, methods=['get'], url_path='account-balance', url_name='account-balance')
    def account_balance(self, request):
        """Get account balance for a specific account and date"""
        account_id = request.query_params.get('account_id')
        as_of_date = request.query_params.get('as_of_date')
        
        if not account_id or not as_of_date:
            return Response(
                {'error': 'account_id and as_of_date are required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        from django.utils.dateparse import parse_date
        as_of = parse_date(as_of_date)
        if not as_of:
            return Response(
                {'error': 'Invalid date format'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            account = Account.objects.get(id=account_id)
        except Account.DoesNotExist:
            return Response(
                {'error': 'Account not found'},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Calculate balance from general ledger entries up to the date
        entries = GeneralLedgerEntry.objects.filter(
            account=account,
            entry_date__lte=as_of
        )
        
        total_debits = sum(e.amount for e in entries.filter(debit_credit='debit'))
        total_credits = sum(e.amount for e in entries.filter(debit_credit='credit'))
        
        # Calculate balance based on account type
        if account.account_type in ['asset', 'expense']:
            balance = total_debits - total_credits
        else:
            balance = total_credits - total_debits
        
        return Response({
            'account_id': account.id,
            'account_number': account.account_number,
            'account_name': account.name,
            'as_of_date': as_of_date,
            'total_debits': total_debits,
            'total_credits': total_credits,
            'balance': balance
        })


class AccountBalanceViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for viewing account balances by fiscal period"""
    queryset = AccountBalance.objects.select_related('account', 'fiscal_period').all()
    serializer_class = AccountBalanceSerializer
    
    def get_queryset(self):
        queryset = AccountBalance.objects.select_related('account', 'fiscal_period').all()
        
        # Filter by fiscal period
        fiscal_period_id = self.request.query_params.get('fiscal_period_id', None)
        if fiscal_period_id:
            queryset = queryset.filter(fiscal_period_id=fiscal_period_id)
        
        # Filter by account
        account_id = self.request.query_params.get('account_id', None)
        if account_id:
            queryset = queryset.filter(account_id=account_id)
        
        return queryset.order_by('account', 'fiscal_period')


class FinancialReportsViewSet(viewsets.ViewSet):
    """ViewSet for generating financial reports"""
    
    @action(detail=False, methods=['get'], url_path='trial-balance', url_name='trial-balance')
    def trial_balance(self, request):
        """Generate trial balance report"""
        as_of_date = request.query_params.get('as_of_date', None)
        fiscal_period_id = request.query_params.get('fiscal_period_id', None)
        
        if not as_of_date and not fiscal_period_id:
            from datetime import date
            as_of_date = date.today().isoformat()
        
        accounts = Account.objects.filter(is_active=True).order_by('account_number')
        trial_balance_data = []
        
        for account in accounts:
            # Calculate balance
            if fiscal_period_id:
                try:
                    fiscal_period = FiscalPeriod.objects.get(id=fiscal_period_id)
                    account_balance = AccountBalance.objects.filter(
                        account=account,
                        fiscal_period=fiscal_period
                    ).first()
                    
                    if account_balance:
                        balance = account_balance.closing_balance
                        total_debits = account_balance.period_debits
                        total_credits = account_balance.period_credits
                    else:
                        balance = 0.0
                        total_debits = 0.0
                        total_credits = 0.0
                except FiscalPeriod.DoesNotExist:
                    balance = 0.0
                    total_debits = 0.0
                    total_credits = 0.0
            else:
                # Calculate from general ledger entries
                from django.utils.dateparse import parse_date
                as_of = parse_date(as_of_date) if as_of_date else None
                
                if as_of:
                    entries = GeneralLedgerEntry.objects.filter(
                        account=account,
                        entry_date__lte=as_of
                    )
                    total_debits = sum(e.amount for e in entries.filter(debit_credit='debit'))
                    total_credits = sum(e.amount for e in entries.filter(debit_credit='credit'))
                else:
                    entries = GeneralLedgerEntry.objects.filter(account=account)
                    total_debits = sum(e.amount for e in entries.filter(debit_credit='debit'))
                    total_credits = sum(e.amount for e in entries.filter(debit_credit='credit'))
                
                # Calculate balance based on account type
                if account.account_type in ['asset', 'expense']:
                    balance = total_debits - total_credits
                else:
                    balance = total_credits - total_debits
            
            # Only include accounts with activity or non-zero balance
            if abs(balance) > 0.01 or total_debits > 0.01 or total_credits > 0.01:
                trial_balance_data.append({
                    'account_id': account.id,
                    'account_number': account.account_number,
                    'account_name': account.name,
                    'account_type': account.account_type,
                    'debit_balance': balance if balance > 0 and account.account_type in ['asset', 'expense'] else 0.0,
                    'credit_balance': abs(balance) if balance < 0 or account.account_type in ['liability', 'equity', 'revenue'] else 0.0,
                    'total_debits': total_debits,
                    'total_credits': total_credits,
                    'balance': balance
                })
        
        # Calculate totals
        total_debits = sum(item['debit_balance'] for item in trial_balance_data)
        total_credits = sum(item['credit_balance'] for item in trial_balance_data)
        
        return Response({
            'as_of_date': as_of_date,
            'fiscal_period_id': fiscal_period_id,
            'accounts': trial_balance_data,
            'total_debits': total_debits,
            'total_credits': total_credits,
            'is_balanced': abs(total_debits - total_credits) < 0.01
        })
    
    @action(detail=False, methods=['get'], url_path='balance-sheet', url_name='balance-sheet')
    def balance_sheet(self, request):
        """Generate balance sheet report"""
        as_of_date = request.query_params.get('as_of_date', None)
        fiscal_period_id = request.query_params.get('fiscal_period_id', None)
        
        if not as_of_date and not fiscal_period_id:
            from datetime import date
            as_of_date = date.today().isoformat()
        
        # Get all accounts
        accounts = Account.objects.filter(is_active=True).order_by('account_number')
        
        assets = []
        liabilities = []
        equity = []
        
        for account in accounts:
            # Calculate balance
            if fiscal_period_id:
                try:
                    fiscal_period = FiscalPeriod.objects.get(id=fiscal_period_id)
                    account_balance = AccountBalance.objects.filter(
                        account=account,
                        fiscal_period=fiscal_period
                    ).first()
                    
                    if account_balance:
                        balance = account_balance.closing_balance
                    else:
                        balance = 0.0
                except FiscalPeriod.DoesNotExist:
                    balance = 0.0
            else:
                from django.utils.dateparse import parse_date
                as_of = parse_date(as_of_date) if as_of_date else None
                
                if as_of:
                    entries = GeneralLedgerEntry.objects.filter(
                        account=account,
                        entry_date__lte=as_of
                    )
                    total_debits = sum(e.amount for e in entries.filter(debit_credit='debit'))
                    total_credits = sum(e.amount for e in entries.filter(debit_credit='credit'))
                else:
                    entries = GeneralLedgerEntry.objects.filter(account=account)
                    total_debits = sum(e.amount for e in entries.filter(debit_credit='debit'))
                    total_credits = sum(e.amount for e in entries.filter(debit_credit='credit'))
                
                if account.account_type in ['asset', 'expense']:
                    balance = total_debits - total_credits
                else:
                    balance = total_credits - total_debits
            
            # Only include accounts with non-zero balance
            if abs(balance) > 0.01:
                account_data = {
                    'account_id': account.id,
                    'account_number': account.account_number,
                    'account_name': account.name,
                    'balance': balance
                }
                
                if account.account_type == 'asset':
                    assets.append(account_data)
                elif account.account_type == 'liability':
                    liabilities.append(account_data)
                elif account.account_type == 'equity':
                    equity.append(account_data)
        
        # Calculate totals
        total_assets = sum(item['balance'] for item in assets)
        total_liabilities = sum(item['balance'] for item in liabilities)
        total_equity = sum(item['balance'] for item in equity)
        total_liabilities_and_equity = total_liabilities + total_equity
        
        return Response({
            'as_of_date': as_of_date,
            'fiscal_period_id': fiscal_period_id,
            'assets': assets,
            'liabilities': liabilities,
            'equity': equity,
            'total_assets': total_assets,
            'total_liabilities': total_liabilities,
            'total_equity': total_equity,
            'total_liabilities_and_equity': total_liabilities_and_equity,
            'is_balanced': abs(total_assets - total_liabilities_and_equity) < 0.01
        })
    
    @action(detail=False, methods=['get'], url_path='income-statement', url_name='income-statement')
    def income_statement(self, request):
        """Generate income statement (P&L) report"""
        start_date = request.query_params.get('start_date', None)
        end_date = request.query_params.get('end_date', None)
        fiscal_period_id = request.query_params.get('fiscal_period_id', None)
        
        # If fiscal period is provided, use its dates
        if fiscal_period_id:
            try:
                fiscal_period = FiscalPeriod.objects.get(id=fiscal_period_id)
                start_date = fiscal_period.start_date.isoformat()
                end_date = fiscal_period.end_date.isoformat()
            except FiscalPeriod.DoesNotExist:
                pass
        
        if not start_date or not end_date:
            from datetime import date, timedelta
            end_date = date.today().isoformat()
            start_date = (date.today() - timedelta(days=30)).isoformat()
        
        # Get revenue and expense accounts
        revenue_accounts = Account.objects.filter(
            is_active=True,
            account_type='revenue'
        ).order_by('account_number')
        
        expense_accounts = Account.objects.filter(
            is_active=True,
            account_type='expense'
        ).order_by('account_number')
        
        revenues = []
        expenses = []
        
        from django.utils.dateparse import parse_date
        start = parse_date(start_date)
        end = parse_date(end_date)
        
        # Calculate revenue
        for account in revenue_accounts:
            if fiscal_period_id:
                try:
                    fiscal_period = FiscalPeriod.objects.get(id=fiscal_period_id)
                    account_balance = AccountBalance.objects.filter(
                        account=account,
                        fiscal_period=fiscal_period
                    ).first()
                    
                    if account_balance:
                        # For revenue accounts, credits increase, debits decrease
                        balance = account_balance.period_credits - account_balance.period_debits
                    else:
                        balance = 0.0
                except FiscalPeriod.DoesNotExist:
                    balance = 0.0
            else:
                if start and end:
                    entries = GeneralLedgerEntry.objects.filter(
                        account=account,
                        entry_date__gte=start,
                        entry_date__lte=end
                    )
                    total_debits = sum(e.amount for e in entries.filter(debit_credit='debit'))
                    total_credits = sum(e.amount for e in entries.filter(debit_credit='credit'))
                    # Revenue: credits increase, debits decrease
                    balance = total_credits - total_debits
                else:
                    balance = 0.0
            
            if abs(balance) > 0.01:
                revenues.append({
                    'account_id': account.id,
                    'account_number': account.account_number,
                    'account_name': account.name,
                    'amount': balance
                })
        
        # Calculate expenses
        for account in expense_accounts:
            if fiscal_period_id:
                try:
                    fiscal_period = FiscalPeriod.objects.get(id=fiscal_period_id)
                    account_balance = AccountBalance.objects.filter(
                        account=account,
                        fiscal_period=fiscal_period
                    ).first()
                    
                    if account_balance:
                        # For expense accounts, debits increase, credits decrease
                        balance = account_balance.period_debits - account_balance.period_credits
                    else:
                        balance = 0.0
                except FiscalPeriod.DoesNotExist:
                    balance = 0.0
            else:
                if start and end:
                    entries = GeneralLedgerEntry.objects.filter(
                        account=account,
                        entry_date__gte=start,
                        entry_date__lte=end
                    )
                    total_debits = sum(e.amount for e in entries.filter(debit_credit='debit'))
                    total_credits = sum(e.amount for e in entries.filter(debit_credit='credit'))
                    # Expenses: debits increase, credits decrease
                    balance = total_debits - total_credits
                else:
                    balance = 0.0
            
            if abs(balance) > 0.01:
                expenses.append({
                    'account_id': account.id,
                    'account_number': account.account_number,
                    'account_name': account.name,
                    'amount': balance
                })
        
        # Calculate totals
        total_revenue = sum(item['amount'] for item in revenues)
        total_expenses = sum(item['amount'] for item in expenses)
        net_income = total_revenue - total_expenses
        
        return Response({
            'start_date': start_date,
            'end_date': end_date,
            'fiscal_period_id': fiscal_period_id,
            'revenues': revenues,
            'expenses': expenses,
            'total_revenue': total_revenue,
            'total_expenses': total_expenses,
            'net_income': net_income
        })
    
    @action(detail=False, methods=['get'], url_path='cash-flow', url_name='cash-flow')
    def cash_flow(self, request):
        """Generate cash flow statement"""
        start_date = request.query_params.get('start_date', None)
        end_date = request.query_params.get('end_date', None)
        fiscal_period_id = request.query_params.get('fiscal_period_id', None)
        
        # If fiscal period is provided, use its dates
        if fiscal_period_id:
            try:
                fiscal_period = FiscalPeriod.objects.get(id=fiscal_period_id)
                start_date = fiscal_period.start_date.isoformat()
                end_date = fiscal_period.end_date.isoformat()
            except FiscalPeriod.DoesNotExist:
                pass
        
        if not start_date or not end_date:
            from datetime import date, timedelta
            end_date = date.today().isoformat()
            start_date = (date.today() - timedelta(days=30)).isoformat()
        
        from django.utils.dateparse import parse_date
        start = parse_date(start_date)
        end = parse_date(end_date)
        
        # Find cash accounts (typically asset accounts with "cash" in the name)
        cash_accounts = Account.objects.filter(
            is_active=True,
            account_type='asset'
        ).filter(
            models.Q(name__icontains='cash') | 
            models.Q(account_number__icontains='cash') |
            models.Q(name__icontains='checking') |
            models.Q(name__icontains='bank')
        )
        
        # Operating Activities - Revenue and Expense accounts
        revenue_accounts = Account.objects.filter(is_active=True, account_type='revenue')
        expense_accounts = Account.objects.filter(is_active=True, account_type='expense')
        
        operating_activities = []
        investing_activities = []
        financing_activities = []
        
        # Calculate operating activities (simplified - revenue minus expenses)
        if start and end:
            if fiscal_period_id:
                try:
                    fiscal_period = FiscalPeriod.objects.get(id=fiscal_period_id)
                    # Get net income from income statement
                    revenue_balance = 0.0
                    for account in revenue_accounts:
                        account_balance = AccountBalance.objects.filter(
                            account=account,
                            fiscal_period=fiscal_period
                        ).first()
                        if account_balance:
                            revenue_balance += account_balance.period_credits - account_balance.period_debits
                    
                    expense_balance = 0.0
                    for account in expense_accounts:
                        account_balance = AccountBalance.objects.filter(
                            account=account,
                            fiscal_period=fiscal_period
                        ).first()
                        if account_balance:
                            expense_balance += account_balance.period_debits - account_balance.period_credits
                    
                    net_income = revenue_balance - expense_balance
                    operating_activities.append({
                        'description': 'Net Income',
                        'amount': net_income
                    })
                except FiscalPeriod.DoesNotExist:
                    pass
            else:
                # Calculate from general ledger
                revenue_total = 0.0
                for account in revenue_accounts:
                    entries = GeneralLedgerEntry.objects.filter(
                        account=account,
                        entry_date__gte=start,
                        entry_date__lte=end
                    )
                    revenue_total += sum(e.amount for e in entries.filter(debit_credit='credit'))
                    revenue_total -= sum(e.amount for e in entries.filter(debit_credit='debit'))
                
                expense_total = 0.0
                for account in expense_accounts:
                    entries = GeneralLedgerEntry.objects.filter(
                        account=account,
                        entry_date__gte=start,
                        entry_date__lte=end
                    )
                    expense_total += sum(e.amount for e in entries.filter(debit_credit='debit'))
                    expense_total -= sum(e.amount for e in entries.filter(debit_credit='credit'))
                
                net_income = revenue_total - expense_total
                operating_activities.append({
                    'description': 'Net Income',
                    'amount': net_income
                })
        
        # Calculate cash flow from operations, investing, and financing
        cash_flow_operations = sum(item['amount'] for item in operating_activities)
        cash_flow_investing = sum(item['amount'] for item in investing_activities)
        cash_flow_financing = sum(item['amount'] for item in financing_activities)
        net_cash_flow = cash_flow_operations + cash_flow_investing + cash_flow_financing
        
        # Calculate beginning and ending cash
        beginning_cash = 0.0
        ending_cash = 0.0
        
        for cash_account in cash_accounts:
            if start:
                # Beginning cash balance
                entries_before = GeneralLedgerEntry.objects.filter(
                    account=cash_account,
                    entry_date__lt=start
                )
                debits_before = sum(e.amount for e in entries_before.filter(debit_credit='debit'))
                credits_before = sum(e.amount for e in entries_before.filter(debit_credit='credit'))
                beginning_cash += debits_before - credits_before
            
            if end:
                # Ending cash balance
                entries_through = GeneralLedgerEntry.objects.filter(
                    account=cash_account,
                    entry_date__lte=end
                )
                debits_through = sum(e.amount for e in entries_through.filter(debit_credit='debit'))
                credits_through = sum(e.amount for e in entries_through.filter(debit_credit='credit'))
                ending_cash += debits_through - credits_through
        
        return Response({
            'start_date': start_date,
            'end_date': end_date,
            'fiscal_period_id': fiscal_period_id,
            'operating_activities': operating_activities,
            'investing_activities': investing_activities,
            'financing_activities': financing_activities,
            'cash_flow_operations': cash_flow_operations,
            'cash_flow_investing': cash_flow_investing,
            'cash_flow_financing': cash_flow_financing,
            'net_cash_flow': net_cash_flow,
            'beginning_cash': beginning_cash,
            'ending_cash': ending_cash
        })
    
    @action(detail=False, methods=['get'], url_path='dashboard-metrics', url_name='dashboard-metrics')
    def dashboard_metrics(self, request):
        """Get aggregated financial metrics for dashboard visualization"""
        from django.db.models import Sum, Q, F
        from django.utils.dateparse import parse_date
        from datetime import datetime, timedelta
        from collections import defaultdict
        
        period_type = request.query_params.get('period_type', 'monthly')  # monthly or quarterly
        months_back = int(request.query_params.get('months_back', 12))
        
        # Calculate date range
        end_date = datetime.now().date()
        start_date = end_date - timedelta(days=months_back * 30)
        
        # Get revenue and expense accounts
        revenue_accounts = Account.objects.filter(
            is_active=True,
            account_type='revenue'
        )
        
        expense_accounts = Account.objects.filter(
            is_active=True,
            account_type='expense'
        )
        
        # Get cash accounts
        cash_accounts = Account.objects.filter(
            is_active=True,
            account_type='asset',
            account_number__startswith='1000'  # Cash accounts typically start with 1000
        )
        
        # Aggregate data by period
        periods_data = defaultdict(lambda: {
            'revenue': 0.0,
            'expenses': 0.0,
            'profit': 0.0,
            'cash_flow': 0.0
        })
        
        # Get revenue by period
        for account in revenue_accounts:
            entries = GeneralLedgerEntry.objects.filter(
                account=account,
                entry_date__gte=start_date,
                entry_date__lte=end_date
            )
            
            for entry in entries:
                period_key = self._get_period_key(entry.entry_date, period_type)
                if entry.debit_credit == 'credit':
                    periods_data[period_key]['revenue'] += entry.amount
                else:
                    periods_data[period_key]['revenue'] -= entry.amount
        
        # Get expenses by period
        for account in expense_accounts:
            entries = GeneralLedgerEntry.objects.filter(
                account=account,
                entry_date__gte=start_date,
                entry_date__lte=end_date
            )
            
            for entry in entries:
                period_key = self._get_period_key(entry.entry_date, period_type)
                if entry.debit_credit == 'debit':
                    periods_data[period_key]['expenses'] += entry.amount
                else:
                    periods_data[period_key]['expenses'] -= entry.amount
        
        # Get cash flow by period
        for account in cash_accounts:
            entries = GeneralLedgerEntry.objects.filter(
                account=account,
                entry_date__gte=start_date,
                entry_date__lte=end_date
            )
            
            for entry in entries:
                period_key = self._get_period_key(entry.entry_date, period_type)
                if entry.debit_credit == 'debit':
                    periods_data[period_key]['cash_flow'] += entry.amount
                else:
                    periods_data[period_key]['cash_flow'] -= entry.amount
        
        # Calculate profit for each period
        for period_key in periods_data:
            periods_data[period_key]['profit'] = (
                periods_data[period_key]['revenue'] - 
                periods_data[period_key]['expenses']
            )
        
        # Convert to list and sort by period
        periods_list = []
        for period_key in sorted(periods_data.keys()):
            periods_list.append({
                'period': period_key,
                **periods_data[period_key]
            })
        
        # Get current period metrics
        current_period_key = self._get_period_key(end_date, period_type)
        current_metrics = periods_data.get(current_period_key, {
            'revenue': 0.0,
            'expenses': 0.0,
            'profit': 0.0,
            'cash_flow': 0.0
        })
        
        # Get AR and AP totals
        try:
            ar_total = AccountsReceivable.objects.filter(
                status='open'
            ).aggregate(total=Sum('amount_due'))['total'] or 0.0
            
            ap_total = AccountsPayable.objects.filter(
                status='open'
            ).aggregate(total=Sum('amount_due'))['total'] or 0.0
        except Exception:
            ar_total = 0.0
            ap_total = 0.0
        
        # Get cash balance
        cash_balance = 0.0
        for account in cash_accounts:
            entries = GeneralLedgerEntry.objects.filter(
                account=account,
                entry_date__lte=end_date
            )
            for entry in entries:
                if entry.debit_credit == 'debit':
                    cash_balance += entry.amount
                else:
                    cash_balance -= entry.amount
        
        return Response({
            'period_type': period_type,
            'periods': periods_list,
            'current_metrics': current_metrics,
            'ar_total': ar_total,
            'ap_total': ap_total,
            'cash_balance': cash_balance
        })

    @action(detail=False, methods=['get'], url_path='kpis', url_name='kpis')
    def kpis(self, request):
        """Get performance KPIs (on-time shipping, etc.) for Finance dashboard and KPIs page."""
        from datetime import timedelta

        months_back = int(request.query_params.get('months_back', 12))
        end_date = timezone.now().date()
        start_date = end_date - timedelta(days=months_back * 30)

        # Shipped orders in range with both expected and actual ship date
        shipped = list(SalesOrder.objects.filter(
            status__in=['shipped', 'completed', 'received'],
            actual_ship_date__isnull=False,
            expected_ship_date__isnull=False,
            actual_ship_date__date__gte=start_date,
            actual_ship_date__date__lte=end_date,
        ).values_list('id', 'expected_ship_date', 'actual_ship_date'))
        total_shipped = len(shipped)
        on_time = 0
        days_late_list = []
        for _id, exp_dt, act_dt in shipped:
            exp = exp_dt.date() if hasattr(exp_dt, 'date') else exp_dt
            act = act_dt.date() if hasattr(act_dt, 'date') else act_dt
            if act <= exp:
                on_time += 1
            else:
                days_late_list.append((act - exp).days)
        late_count = total_shipped - on_time
        on_time_pct = round((on_time / total_shipped * 100.0), 1) if total_shipped else None
        avg_days_late = round(sum(days_late_list) / len(days_late_list), 1) if days_late_list else None

        return Response({
            'period': {'start_date': start_date.isoformat(), 'end_date': end_date.isoformat(), 'months_back': months_back},
            'shipping': {
                'on_time_shipment_pct': on_time_pct,
                'on_time_count': on_time,
                'late_count': late_count,
                'total_shipped': total_shipped,
                'avg_days_late': avg_days_late,
            },
        })

    def _get_period_key(self, date, period_type):
        """Get period key for grouping (e.g., '2025-01' for monthly, '2025-Q1' for quarterly)"""
        if period_type == 'quarterly':
            quarter = (date.month - 1) // 3 + 1
            return f"{date.year}-Q{quarter}"
        else:  # monthly
            return f"{date.year}-{date.month:02d}"


class CheckInLogViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for CheckInLog - comprehensive check-in logs"""
    from .models import CheckInLog
    from .serializers import CheckInLogSerializer
    queryset = CheckInLog.objects.select_related('lot', 'lot__item').all().order_by('-checked_in_at')
    serializer_class = CheckInLogSerializer
    
    def get_queryset(self):
        queryset = CheckInLog.objects.select_related('lot', 'lot__item').all().order_by('-checked_in_at')
        
        # Filter by item SKU
        item_sku = self.request.query_params.get('item_sku')
        if item_sku:
            queryset = queryset.filter(item_sku=item_sku)
        
        # Filter by PO number
        po_number = self.request.query_params.get('po_number')
        if po_number:
            queryset = queryset.filter(po_number=po_number)
        
        # Filter by date range
        date_from = self.request.query_params.get('date_from')
        date_to = self.request.query_params.get('date_to')
        if date_from:
            queryset = queryset.filter(checked_in_at__gte=date_from)
        if date_to:
            queryset = queryset.filter(checked_in_at__lte=date_to)
        
        return queryset


class ProductionLogViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for viewing production logs"""
    queryset = ProductionLog.objects.all()
    serializer_class = ProductionLogSerializer
    
    def get_queryset(self):
        queryset = ProductionLog.objects.select_related('batch', 'batch__finished_good_item').all()
        
        # Filter by batch number if provided
        batch_number = self.request.query_params.get('batch_number', None)
        if batch_number:
            queryset = queryset.filter(batch_number=batch_number)
        
        # Filter by SKU if provided
        sku = self.request.query_params.get('sku', None)
        if sku:
            queryset = queryset.filter(finished_good_sku=sku)
        
        # Filter by batch type if provided
        batch_type = self.request.query_params.get('batch_type', None)
        if batch_type:
            queryset = queryset.filter(batch_type=batch_type)
        
        # Filter by date range if provided
        date_from = self.request.query_params.get('date_from', None)
        date_to = self.request.query_params.get('date_to', None)
        if date_from:
            from django.utils.dateparse import parse_datetime
            date_from_dt = parse_datetime(date_from)
            if date_from_dt:
                queryset = queryset.filter(closed_date__gte=date_from_dt)
        if date_to:
            from django.utils.dateparse import parse_datetime
            date_to_dt = parse_datetime(date_to)
            if date_to_dt:
                queryset = queryset.filter(closed_date__lte=date_to_dt)
        
        return queryset.order_by('-logged_at')



class AccountsPayableViewSet(viewsets.ModelViewSet):
    """ViewSet for managing Accounts Payable"""
    queryset = AccountsPayable.objects.select_related('purchase_order', 'account', 'journal_entry').all()
    serializer_class = AccountsPayableSerializer
    
    def get_queryset(self):
        queryset = AccountsPayable.objects.select_related('purchase_order', 'account', 'journal_entry').all()
        
        # Filter by status
        status = self.request.query_params.get('status', None)
        if status:
            queryset = queryset.filter(status=status)
        
        # Filter by vendor
        vendor_name = self.request.query_params.get('vendor_name', None)
        if vendor_name:
            queryset = queryset.filter(vendor_name__icontains=vendor_name)
        
        # Filter by due date range
        due_date_from = self.request.query_params.get('due_date_from', None)
        due_date_to = self.request.query_params.get('due_date_to', None)
        if due_date_from:
            queryset = queryset.filter(due_date__gte=due_date_from)
        if due_date_to:
            queryset = queryset.filter(due_date__lte=due_date_to)
        
        return queryset.order_by('due_date', 'vendor_name')
    
    @action(detail=False, methods=['get'], url_path='aging', url_name='aging')
    def aging_report(self, request):
        """Generate AP Aging Report"""
        from django.utils import timezone
        
        today = timezone.now().date()
        
        # Get all open AP entries
        ap_entries = AccountsPayable.objects.filter(
            status__in=['open', 'partial', 'overdue']
        ).order_by('vendor_name', 'due_date')
        
        # Group by aging buckets
        aging_data = {
            'not_due': [],
            '0-30': [],
            '31-60': [],
            '61-90': [],
            'over_90': []
        }
        
        totals = {
            'not_due': 0.0,
            '0-30': 0.0,
            '31-60': 0.0,
            '61-90': 0.0,
            'over_90': 0.0,
            'total': 0.0
        }
        
        for entry in ap_entries:
            bucket = entry.aging_bucket
            aging_data[bucket].append({
                'id': entry.id,
                'vendor_name': entry.vendor_name,
                'invoice_number': entry.invoice_number,
                'invoice_date': entry.invoice_date,
                'due_date': entry.due_date,
                'original_amount': entry.original_amount,
                'amount_paid': entry.amount_paid,
                'balance': entry.balance,
                'days_aging': entry.days_aging,
                'status': entry.status,
                'po_number': entry.purchase_order.po_number if entry.purchase_order else None
            })
            totals[bucket] += entry.balance
            totals['total'] += entry.balance
        
        return Response({
            'as_of_date': today,
            'aging_data': aging_data,
            'totals': totals
        })


class AccountsReceivableViewSet(viewsets.ModelViewSet):
    """ViewSet for managing Accounts Receivable"""
    queryset = AccountsReceivable.objects.select_related('invoice', 'sales_order', 'account', 'journal_entry').all()
    serializer_class = AccountsReceivableSerializer
    
    def get_queryset(self):
        queryset = AccountsReceivable.objects.select_related('invoice', 'sales_order', 'account', 'journal_entry').all()
        
        # Filter by status
        status = self.request.query_params.get('status', None)
        if status:
            queryset = queryset.filter(status=status)
        
        # Filter by customer
        customer_name = self.request.query_params.get('customer_name', None)
        if customer_name:
            queryset = queryset.filter(customer_name__icontains=customer_name)
        
        # Filter by due date range
        due_date_from = self.request.query_params.get('due_date_from', None)
        due_date_to = self.request.query_params.get('due_date_to', None)
        if due_date_from:
            queryset = queryset.filter(due_date__gte=due_date_from)
        if due_date_to:
            queryset = queryset.filter(due_date__lte=due_date_to)
        
        return queryset.order_by('due_date', 'customer_name')
    
    @action(detail=False, methods=['get'], url_path='aging', url_name='aging')
    def aging_report(self, request):
        """Generate AR Aging Report"""
        from django.utils import timezone
        
        today = timezone.now().date()
        
        # Get all open AR entries
        ar_entries = AccountsReceivable.objects.filter(
            status__in=['open', 'partial', 'overdue']
        ).order_by('customer_name', 'due_date')
        
        # Group by aging buckets
        aging_data = {
            'not_due': [],
            '0-30': [],
            '31-60': [],
            '61-90': [],
            'over_90': []
        }
        
        totals = {
            'not_due': 0.0,
            '0-30': 0.0,
            '31-60': 0.0,
            '61-90': 0.0,
            'over_90': 0.0,
            'total': 0.0
        }
        
        for entry in ar_entries:
            bucket = entry.aging_bucket
            aging_data[bucket].append({
                'id': entry.id,
                'customer_name': entry.customer_name,
                'invoice_number': entry.invoice.invoice_number if entry.invoice else None,
                'invoice_date': entry.invoice_date,
                'due_date': entry.due_date,
                'original_amount': entry.original_amount,
                'amount_paid': entry.amount_paid,
                'balance': entry.balance,
                'days_aging': entry.days_aging,
                'status': entry.status,
                'so_number': entry.sales_order.so_number if entry.sales_order else None
            })
            totals[bucket] += entry.balance
            totals['total'] += entry.balance
        
        return Response({
            'as_of_date': today,
            'aging_data': aging_data,
            'totals': totals
        })


class PaymentViewSet(viewsets.ModelViewSet):
    """ViewSet for managing Payments"""
    queryset = Payment.objects.select_related('ap_entry', 'ar_entry', 'account', 'journal_entry').all()
    serializer_class = PaymentSerializer
    
    def get_queryset(self):
        queryset = Payment.objects.select_related('ap_entry', 'ar_entry', 'account', 'journal_entry').all()
        
        # Filter by payment type
        payment_type = self.request.query_params.get('payment_type', None)
        if payment_type:
            queryset = queryset.filter(payment_type=payment_type)
        
        # Filter by date range
        date_from = self.request.query_params.get('date_from', None)
        date_to = self.request.query_params.get('date_to', None)
        if date_from:
            queryset = queryset.filter(payment_date__gte=date_from)
        if date_to:
            queryset = queryset.filter(payment_date__lte=date_to)
        
        return queryset.order_by('-payment_date', '-created_at')
    
    def create(self, request, *args, **kwargs):
        """Create a payment and update AP/AR entry"""
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        payment = serializer.save()
        
        # Update AP or AR entry
        if payment.ap_entry:
            ap_entry = payment.ap_entry
            ap_entry.amount_paid += payment.amount
            ap_entry.balance = ap_entry.original_amount - ap_entry.amount_paid
            
            # Update status
            if ap_entry.balance <= 0.01:  # Allow for floating point precision
                ap_entry.status = 'paid'
                ap_entry.balance = 0.0
            elif ap_entry.amount_paid > 0:
                ap_entry.status = 'partial'
            
            # Check if overdue
            from django.utils import timezone
            if ap_entry.status in ['open', 'partial'] and ap_entry.due_date < timezone.now().date():
                ap_entry.status = 'overdue'
            
            ap_entry.save()
        
        elif payment.ar_entry:
            ar_entry = payment.ar_entry
            ar_entry.amount_paid += payment.amount
            ar_entry.balance = ar_entry.original_amount - ar_entry.amount_paid
            
            # Update status
            if ar_entry.balance <= 0.01:  # Allow for floating point precision
                ar_entry.status = 'paid'
                ar_entry.balance = 0.0
            elif ar_entry.amount_paid > 0:
                ar_entry.status = 'partial'
            
            # Check if overdue
            from django.utils import timezone
            if ar_entry.status in ['open', 'partial'] and ar_entry.due_date < timezone.now().date():
                ar_entry.status = 'overdue'
            
            ar_entry.save()
        
        return Response(serializer.data, status=status.HTTP_201_CREATED)


class BankReconciliationViewSet(viewsets.ModelViewSet):
    """ViewSet for bank reconciliations (statement date/balance vs GL)."""
    queryset = BankReconciliation.objects.select_related('account').all()
    serializer_class = BankReconciliationSerializer

    def get_queryset(self):
        queryset = BankReconciliation.objects.select_related('account').all()
        account_id = self.request.query_params.get('account_id', None)
        if account_id:
            queryset = queryset.filter(account_id=account_id)
        return queryset.order_by('-statement_date')
